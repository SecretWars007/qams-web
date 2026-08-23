# -*- coding: utf-8 -*-
"""
generate_perfil_proyecto.py
Perfil del Proyecto — QAMS (Quality Assurance Management System)
Maximo 15 paginas | Monografia de Diplomado
Ejecutar: python generate_perfil_proyecto.py
Salida:   c:/diplomado/perfil.docx
          c:/diplomado/qams-web/perfil.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT1 = r"c:/diplomado/qams-web/perfil.docx"
OUT2 = r"c:/diplomado/perfil.docx"

NAVY  = RGBColor(0x1E, 0x3A, 0x5F)
BLUE  = RGBColor(0x1D, 0x4E, 0x89)
GOLD  = RGBColor(0xC9, 0x9A, 0x06)
GRAY  = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _sbg(cell, hx):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hx); p.append(s)


def _sbord(cell):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for sd in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{sd}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "CCCCCC")
        b.append(e)
    p.append(b)


def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── helpers ─────────────────────────────────────────────────────────────
    def tx(text, bold=False, italic=False, sz=11,
           al=WD_ALIGN_PARAGRAPH.JUSTIFY, col=None, sa=5):
        p = doc.add_paragraph(); p.alignment = al
        p.paragraph_format.space_after = Pt(sa)
        r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.size = Pt(sz); r.font.name = "Calibri"
        if col:
            r.font.color.rgb = col
        return p

    def h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text.upper())
        r.bold = True; r.font.size = Pt(13)
        r.font.name = "Calibri"; r.font.color.rgb = NAVY
        pp = OxmlElement("w:pPr"); pb = OxmlElement("w:pBdr")
        bt = OxmlElement("w:bottom")
        bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "6")
        bt.set(qn("w:space"), "1"); bt.set(qn("w:color"), "1D4E89")
        pb.append(bt); pp.append(pb); p._p.insert(0, pp)

    def h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(11.5)
        r.font.name = "Calibri"; r.font.color.rgb = BLUE

    def bl(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10.5); r.font.name = "Calibri"

    def tb(hds, rows, cws, hc="1D4E89"):
        t = doc.add_table(rows=1, cols=len(hds))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        hr = t.rows[0]
        for i, h in enumerate(hds):
            c = hr.cells[i]; c.width = Inches(cws[i]); _sbg(c, hc)
            p2 = c.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after  = Pt(2)
            r = p2.add_run(h); r.bold = True
            r.font.size = Pt(8.5); r.font.name = "Calibri"
            r.font.color.rgb = WHITE
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for ri, rd in enumerate(rows):
            row = t.add_row()
            bg = "F2F6FC" if ri % 2 == 0 else "FFFFFF"
            for ci, v in enumerate(rd):
                c = row.cells[ci]; c.width = Inches(cws[ci])
                _sbg(c, bg); _sbord(c)
                p2 = c.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p2.paragraph_format.space_before = Pt(1)
                p2.paragraph_format.space_after  = Pt(1)
                r = p2.add_run(v)
                r.font.size = Pt(8.5); r.font.name = "Calibri"
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    def nt(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"Fuente: {text}")
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRAY

    def qb(text):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(7)
        r = p.add_run(text)
        r.italic = True; r.font.size = Pt(11); r.font.color.rgb = NAVY

    def cbar(hx, ht=20):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(" " * 100); r.font.size = Pt(ht)
        pp = OxmlElement("w:pPr"); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hx); pp.append(shd); p._p.insert(0, pp)

    # =====================================================================
    # PORTADA
    # =====================================================================
    cbar("1E3A5F", 24)
    doc.add_paragraph()

    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pi.add_run("DIPLOMADO EN ARQUITECTURA DE SOFTWARE EMPRESARIAL")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY; r.font.name = "Calibri"

    doc.add_paragraph()

    pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.paragraph_format.space_before = Pt(16)
    r = pl.add_run("PERFIL DEL PROYECTO DE MONOGRAFIA")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY; r.font.name = "Calibri"

    doc.add_paragraph()

    ps = doc.add_paragraph(); ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ps.add_run("QAMS")
    r.bold = True; r.font.size = Pt(40); r.font.color.rgb = BLUE; r.font.name = "Calibri"

    pst = doc.add_paragraph(); pst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pst.add_run("Quality Assurance Management System")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = GOLD; r.font.name = "Calibri"

    doc.add_paragraph()

    ptg = doc.add_paragraph(); ptg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ptg.add_run(
        "Plataforma Web Empresarial Fullstack para la Gestion Integral del\n"
        "Ciclo de Vida de las Pruebas de Software bajo ISTQB CTFL v4.0"
    )
    r.font.size = Pt(12); r.font.color.rgb = GRAY; r.font.name = "Calibri"

    doc.add_paragraph()

    for lb, vl in [
        ("Autor",           "Diego Fernando Sanchez"),
        ("Tutor Academico", "Nombre del Director de Monografia"),
        ("Area Tematica",   "Ingenieria de Software - Aseguramiento de Calidad"),
        ("Stack",           ".NET 9  |  Angular 19  |  PostgreSQL 16  |  Docker"),
        ("Fecha",           "Agosto 2026"),
    ]:
        pm = doc.add_paragraph(); pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_after = Pt(3)
        r1 = pm.add_run(f"{lb}: "); r1.bold = True
        r1.font.size = Pt(11); r1.font.color.rgb = NAVY; r1.font.name = "Calibri"
        r2 = pm.add_run(vl)
        r2.font.size = Pt(11); r2.font.color.rgb = GRAY; r2.font.name = "Calibri"

    doc.add_paragraph()
    cbar("1D4E89", 14)
    doc.add_page_break()
    print("  [OK] Portada")

    # =====================================================================
    # 1. RESUMEN EJECUTIVO
    # =====================================================================
    h1("1. Resumen Ejecutivo del Proyecto")
    tx("QAMS (Quality Assurance Management System) es una plataforma web empresarial "
       "fullstack de codigo abierto y autoalojable (self-hosted), disenada para centralizar "
       "y gobernar el ciclo de vida completo de las pruebas de software bajo el estandar "
       "ISTQB CTFL v4.0 e ISO/IEC/IEEE 29119. Construida sobre Monolito Modular con Clean "
       "Architecture en .NET 9 (backend) y Angular 19 con Signals (frontend), desplegada en "
       "Docker Compose con un solo comando.")
    tx("El proyecto surge como respuesta a una problematica critica: los equipos de QA en "
       "America Latina operan con herramientas fragmentadas (Excel) sin trazabilidad RTM, "
       "sin metricas formales y sin cumplimiento de estandares. Las herramientas enterprise "
       "(TestRail, HP ALM) tienen un TCO de $17,760-$108,000 USD a 5 anios para 15 usuarios. "
       "QAMS resuelve esta brecha con solo $2,400 USD de costo total a 5 anios (VPS $40/mes), "
       "cero licencias, open source.")

    tb(
        ["Indicador Clave del Proyecto", "Meta", "Componente QAMS"],
        [
            ["Defect Removal Efficiency (DRE)",       ">= 90%",   "Quality Gates + RTM + Defect Lifecycle"],
            ["Reduccion TCO vs TestRail (5 anios)",   "> 96.5%",  "Open Source Self-Hosted + Docker $40/mes"],
            ["Latencia Fast Runner (P/F/B por paso)", "< 200 ms", "Angular 19 Signals - sin Zone.js"],
            ["Cobertura RTM bidireccional",           ">= 95%",   "RequirementTestCase M:N + endpoint RTM"],
            ["Conformidad ISTQB CTFL v4.0",           "100%",     "7 modulos funcionales integrados"],
            ["Despliegue produccion completo",        "1 comando","docker compose up -d (4 servicios Alpine)"],
        ],
        [2.6, 1.1, 2.8]
    )
    print("  [OK] Seccion 1")

    # =====================================================================
    # 2. INTRODUCCION Y CONTEXTO
    # =====================================================================
    h1("2. Introduccion y Contexto")
    tx("La calidad del software ha transitado de ser una fase aislada a consolidarse como "
       "el pilar estrategico de cualquier producto digital moderno. El CISQ (2022) cuantifico "
       "en $2.41 billones de dolares anuales el costo de los fallos de software en EE.UU. "
       "En America Latina, los equipos de QA operan principalmente con Excel sin trazabilidad "
       "ni metricas formales, mientras las herramientas enterprise tienen costos prohibitivos. "
       "QAMS democratiza el acceso a herramientas ALM de clase enterprise bajo estandares internacionales.")

    tb(
        ["Fase del SDLC", "Costo Relativo", "Consecuencia del Defecto Escapado", "Mitigacion en QAMS"],
        [
            ["Requerimientos", "1x (base)",   "Ambiguedad; rediseno total",           "RTM + Static Reviews / Walkthrough"],
            ["Codificacion",   "5x - 10x",    "Refactorizacion de modulos",           "Criterios BDD / Gherkin nativos"],
            ["Pruebas",        "10x - 15x",   "Regresion completa; retraso entrega",  "Fast Runner + Defect Kanban"],
            ["Pre-Produccion", "20x - 40x",   "Riesgo despliegue; rollback necesario","Quality Gates automatizados"],
            ["Produccion",     "40x - 100x",  "Multas, perdida de datos, reputacion", "Audit Trail + RTM permanente"],
        ],
        [1.3, 0.9, 1.9, 1.9]
    )
    nt("IBM Systems Sciences Institute (2000); CISQ Report 2022; Capers Jones (2008).")
    print("  [OK] Seccion 2")

    # =====================================================================
    # 3. ANTECEDENTES
    # =====================================================================
    h1("3. Antecedentes")
    h2("3.1 Antecedentes del Objeto de Estudio")
    tx("El testing formal surgio con los trabajos de Dijkstra (1969). La IEEE 829 (1983) "
       "formalizo los artefactos del proceso: Plan de Pruebas, Casos y Reporte de Defectos. "
       "El Manifiesto Agil (2001) y DevOps (2009) transformaron el testing de fase "
       "post-desarrollo a actividad continua, invalidando herramientas ALM heredadas "
       "(HP Quality Center, IBM Rational). Nuevas plataformas (Zephyr, TestRail, Jira Xray) "
       "reprodujeron el mismo problema con modelos SaaS restrictivos que excluyen a PYMEs "
       "y universidades del acceso a tecnologia de calidad empresarial.")

    h2("3.2 Referencias Tecnicas de Trabajos Relacionados")
    tb(
        ["Ref.", "Titulo / Autor / Anio", "Objetivo General", "Diferencia con QAMS"],
        [
            ["[1]",
             "Sistema web para gestion de pruebas en DevOps\n"
             "Sanchez, D. (2023) - UPV Espania",
             "Centralizar resultados E2E vía microservicios AWS.",
             "QAMS usa Monolito Modular ACID vs. 12 microservicios con consistencia eventual. "
             "Incluye RTM, BDD y SBTM ausentes. Costo $40/mes vs. $800+/mes K8s AWS."],
            ["[2]",
             "Plataforma QA en organismos gubernamentales\n"
             "Gomez & Perez (2022) - UAEM Mexico",
             "Estandarizar ciclo de defectos con PHP/MySQL.",
             "QAMS: SPA Angular 19 Signals (sin recargas) vs. SSR PHP. "
             "Seguridad OWASP multicapa. Audit Trail vs. sin gobernanza de datos."],
            ["[3]",
             "Open Source QA para PYMEs iberoamericanas\n"
             "Martinez, A.S. (2024) - Rev. Iberoam. Ing. SW",
             "Reducir TCO con TestLink + MantisBT.",
             "QAMS unifica ambas herramientas + RTM + Dashboard en UNA plataforma nativa. "
             "UX moderna con NPS positivo vs. NPS negativo en herramientas de 2006-2010."],
        ],
        [0.4, 1.9, 1.5, 2.7]
    )
    print("  [OK] Seccion 3")

    # =====================================================================
    # 4. DESCRIPCION DEL OBJETO DE ESTUDIO
    # =====================================================================
    h1("4. Descripcion del Objeto de Estudio")
    tx("El objeto de estudio es el Proceso Formal de Gestion, Diseno, Ejecucion, "
       "Control y Gobernanza del Ciclo de Vida de las Pruebas de Software (STLC), "
       "con enfasis en su implementacion computacional bajo ISTQB CTFL v4.0. "
       "QAMS abstrae esta complejidad en un modelo de dominio rico donde cada concepto "
       "del estandar ISTQB es una Entidad de Dominio con invariantes de negocio y reglas "
       "de estado en C# .NET 9: SUT, Proyecto, TestPlan, TestSuite, TestCase, "
       "TestExecution, Defect, ReviewSession, ExploratorySession, KanbanBoard.")

    tb(
        ["#", "Principio ISTQB CTFL v4.0", "Implementacion Concreta en QAMS"],
        [
            ["P1", "Testing muestra defectos, no su ausencia",
             "Quality Gates expresan probabilidad de cobertura %, no perfeccion. "
             "Dashboard muestra zonas de riesgo no cubiertas con mapa de calor."],
            ["P2", "Pruebas exhaustivas son imposibles",
             "Priorizacion RBT (impacto x probabilidad) + suites focalizadas por modulo/riesgo."],
            ["P3", "Pruebas tempranas ahorran tiempo y dinero",
             "Modulo Static Testing (Walkthroughs) permite revisar requisitos "
             "y arquitectura ANTES de implementar codigo."],
            ["P4", "Agrupacion de defectos (Clustering)",
             "Dashboard muestra mapa de calor de defectos por modulo/SUT "
             "(Ng2-Charts) para priorizar esfuerzo de pruebas."],
            ["P5", "Paradoja del pesticida",
             "Versionado de TestCase + Sesiones SBTM-Charters inyectan "
             "aleatoriedad metodologica controlada."],
            ["P6", "El testing depende del contexto",
             "N SUTs con N estrategias independientes; Quality Gates configurables "
             "por plan de prueba con metricas propias."],
            ["P7", "Falacia de la ausencia de errores",
             "Matriz RTM vincula CADA caso de prueba con un requerimiento de "
             "negocio verificable, asegurando valor real al cliente."],
        ],
        [0.3, 1.7, 4.0]
    )
    nt("ISTQB CTFL Syllabus v4.0 (2023). Adaptacion de implementacion por el autor.")
    print("  [OK] Seccion 4")

    # =====================================================================
    # 5. IDENTIFICACION Y FORMULACION DEL PROBLEMA
    # =====================================================================
    h1("5. Identificacion y Formulacion del Problema")
    h2("5.1 Problema Identificado")
    bl("Dimension Economica: Herramientas enterprise ($15-$120+/usuario/mes) "
       "generan un TCO de $18,720-$108,000 USD a 5 anios para 15 usuarios. "
       "PYMEs y universidades quedan excluidas, gestionando calidad con Excel sin trazabilidad.")
    bl("Dimension Metodologica: Fragmentacion entre herramientas independientes rompe "
       "la RTM. DRE < 60% permite que defectos criticos lleguen a produccion "
       "con costo 40x-100x mayor de correccion. Sindrome de la Falsa Cobertura.")
    bl("Dimension Tecnologica: Interfaces SSR con > 2 segundos por interaccion; "
       "ausencia de BDD/Gherkin, SBTM y Quality Gates automaticos; "
       "vulnerabilidades OWASP no resueltas en codigo PHP heredado.")

    h2("5.2 Formulacion del Problema")
    qb("De que manera el diseno, implementacion y despliegue de una plataforma web "
       "empresarial integral, fundamentada en ISTQB CTFL v4.0 e ISO 29119, construida "
       "mediante Monolito Modular con Clean Architecture (.NET 9 / Angular 19), orquestada "
       "en contenedores Docker y asegurada criptograficamente (OWASP Top 10), permite "
       "optimizar el ciclo de vida de las pruebas de software, garantizar la trazabilidad "
       "RTM bidireccional absoluta y reducir el TCO en mas del 90% frente a las "
       "soluciones privativas dominantes del mercado?")

    h2("5.3 Hipotesis General (H1)")
    tx("La implementacion de QAMS, basada en Monolito Modular Fullstack (.NET 9 + Angular 19) "
       "con trazabilidad RTM integral bajo ISTQB CTFL v4.0, incrementara la Eficiencia de "
       "Eliminacion de Defectos (DRE >= 90%), reducira los tiempos de registro de ejecucion "
       "en mas del 60% respecto a herramientas legadas, y lograra una reduccion del TCO "
       "superior al 90% frente a las soluciones SaaS del mercado (H0: sin diferencias "
       "estadisticamente significativas en DRE, tiempo de registro ni TCO).")
    print("  [OK] Seccion 5")

    # =====================================================================
    # 6. OBJETIVOS
    # =====================================================================
    h1("6. Objetivos del Proyecto")
    h2("6.1 Objetivo General")
    tx("Desarrollar una plataforma web fullstack denominada QAMS (Quality Assurance "
       "Management System), basada en una arquitectura de Monolito Modular con Clean "
       "Architecture (.NET 9 / Angular 19), que centralice y governe el ciclo de vida "
       "completo de las pruebas de software bajo el estandar ISTQB CTFL v4.0, mediante "
       "el analisis de requisitos, el diseno del modelo de datos relacional normalizado en "
       "3FN, la construccion de los modulos funcionales de gestion de pruebas y la "
       "contenedorizacion del sistema mediante Docker para su preparacion al entorno "
       "de produccion.")

    h2("6.2 Objetivos Especificos")
    tb(
        ["OE", "Enunciado", "Entregable Principal", "Criterio de Aceptacion"],
        [
            ["OE1",
             "Disenar el modelo relacional ERD en PostgreSQL 16 con "
             "normalizacion 3FN, Soft-Delete y Audit Trail en las 23 "
             "entidades del dominio STLC.",
             "ERD + migraciones EF Core 9",
             "3FN; IAuditable + ISoftDelete en todas las entidades; "
             "integridad referencial completa."],
            ["OE2",
             "Construir la API RESTful en C#/.NET 9 con Clean Architecture "
             "(4 capas), JWT, RBAC y cifrado AES-256 cumpliendo OWASP Top 10.",
             "API REST .NET 9 en Docker Alpine",
             "JWT; BCrypt factor 11; RBAC granular; "
             "HTTP 401/403 correcto; P95 < 130ms."],
            ["OE3",
             "Construir la SPA en Angular 19 Standalone+Signals: "
             "Fast Runner P/F/B, Kanban defectos drag-and-drop "
             "y Dashboard analitico Ng2-Charts.",
             "SPA Angular/Nginx en Docker",
             "Fast Runner < 200ms; Kanban operativo; "
             "Dashboard con KPIs DDP/DRE/MTTR."],
            ["OE4",
             "Disenar e integrar los 7 modulos ISTQB CTFL v4.0: "
             "Static Testing, SBTM Charters, BDD/Gherkin nativo, "
             "RTM bidireccional automatica y Quality Gates configurables.",
             "7 modulos ISTQB funcionales",
             "Cobertura RTM >= 95%; Quality Gate con semaforo; "
             "sesiones SBTM; inspecciones formales."],
            ["OE5",
             "Configurar entorno Docker Compose multi-stage "
             "(backend .NET Alpine + frontend Nginx + PostgreSQL 16 "
             "+ Redis 7) listo para produccion.",
             "docker-compose.yml + Dockerfiles",
             "docker compose up -d levanta el stack completo; "
             "4 servicios en estado healthy."],
        ],
        [0.4, 2.4, 1.5, 2.2]
    )
    print("  [OK] Seccion 6")

    # =====================================================================
    # 7. JUSTIFICACION
    # =====================================================================
    h1("7. Justificacion del Proyecto")
    h2("7.1 Justificacion Tecnica")
    tx("El Monolito Modular supera a los microservicios para el dominio ALM: el grafo "
       "relacional denso del STLC (Defecto -> Ejecucion -> TestCase -> Requisito -> SUT) "
       "requiere transacciones ACID en una sola operacion. Angular 19 Signals es critico "
       "para la reactividad atomica del Fast Runner (200 casos en 4 horas sin recargas de "
       "pagina). La seguridad multicapa (AES-256 + JWT + BCrypt factor 11 + RBAC granular) "
       "protege datos intelectuales criticos cumpliendo OWASP Top 10 2021.")

    h2("7.2 Justificacion Social")
    tx("QAMS democratiza tecnologia ALM de elite: PYMEs, universidades y organismos "
       "estatales sin presupuesto SaaS acceden a una herramienta de clase enterprise "
       "con cero costo de licenciamiento. El codigo fuente en GitHub es un repositorio "
       "pedagogico de Clean Architecture, Angular Signals y Docker aplicados en produccion "
       "real, cerrando la brecha entre la academia y el mercado laboral de ingenieria de "
       "software.")

    h2("7.3 Justificacion Economica — Analisis TCO a 5 Anios (15 usuarios)")
    tb(
        ["Herramienta", "Costo/Usuario/Mes", "Costo Anual", "TCO 5 Anios", "Ahorro vs QAMS"],
        [
            ["TestRail Enterprise", "$74 USD",    "$13,320",   "$69,000",    "$66,600 (96.5%)"],
            ["HP ALM / OpenText",   "$120+ USD",  "$21,600+",  "$108,000+",  "$105,600+ (97.8%)"],
            ["Jira Xray + Jira",    "$18.15 USD", "$3,264",    "$18,720",    "$16,320 (85.7%)"],
            ["Zephyr + Jira",       "$19.65 USD", "$3,552",    "$20,160",    "$17,760 (86.2%)"],
            ["QAMS (Self-hosted)",  "$0 (Open Source)", "$0", "$2,400 total (VPS)", "---"],
        ],
        [1.7, 1.2, 1.1, 1.2, 1.3]
    )
    nt("Precios verificados agosto 2026. TCO incluye VPS $40/mes para QAMS.")
    print("  [OK] Seccion 7")

    # =====================================================================
    # 8. ALCANCES Y LIMITES
    # =====================================================================
    h1("8. Alcances y Limites del Proyecto")
    h2("8.1 Alcances — Modulos Funcionales v1.0")
    tb(
        ["Modulo del Sistema", "Funcionalidades Incluidas en v1.0"],
        [
            ["Dashboard Analitico",
             "KPIs ISTQB: Pass Rate, DRE, DDP, MTTR; Quality Gate Widget semaforo; "
             "graficas Ng2-Charts; 7 tipos de reportes PDF/Excel exportables."],
            ["Proyectos y SUT",
             "CRUD de Sistemas Bajo Prueba y Proyectos; asignacion de testers; "
             "estados de proyecto; estadisticas en tiempo real."],
            ["Requerimientos y RTM",
             "Requisitos funcionales/no funcionales; vinculacion M:N con TestCases; "
             "Matriz RTM automatica con calculo de cobertura %."],
            ["Planes y Suites de Prueba",
             "Planes IEEE 829 con criterios ENTRY/EXIT, hitos y riesgos (RBT); "
             "suites por nivel y tipo de prueba."],
            ["Casos de Prueba",
             "Editor Dual: Clasico (pasos estructurados) y BDD/Gherkin; "
             "versionado automatico; Risk Score (impacto x probabilidad)."],
            ["Fast Runner",
             "Motor reactivo Angular Signals; atajos P/F/B; "
             "actualizacion por paso < 200ms; adjunto de evidencias."],
            ["Defectos y Kanban",
             "Ciclo de vida Nuevo->Cerrado; Kanban drag-and-drop; "
             "trazabilidad TC+Ejecucion; metricas MTTR calculadas."],
            ["Static Testing (Cap.3 ISTQB)",
             "Sesiones de Revision (Walkthrough/Inspeccion/Revision Tecnica); "
             "hallazgos; dictamen formal (Aprobado/Rechazado)."],
            ["SBTM Exploratorio (Cap.4 ISTQB)",
             "Sesiones con Charter de mision; time-box configurable; "
             "registro de hallazgos; metricas de sesion."],
            ["Usuarios y Seguridad RBAC",
             "5 roles + 15 permisos atomicos; BCrypt + JWT; AES-256 en "
             "transito; recuperacion contrasenia; Audit Trail inmutable."],
        ],
        [2.0, 4.5]
    )

    h2("8.2 Limites — Fuera de Alcance en v1.0")
    bl("Ejecucion automatizada nativa (Selenium/Playwright/Cypress): QAMS es el "
       "ALM receptor de resultados, no el motor de automatizacion.")
    bl("Webhooks CI/CD bidireccionales nativos (GitHub Actions, Azure DevOps): "
       "integracion manual disponible via API REST en v1.0.")
    bl("Multitenant SaaS: arquitectura single-tenant self-hosted en v1.0. "
       "Multi-tenant previsto para v2.0.")
    bl("ITSM/HelpDesk: el Kanban gestiona defectos de software, "
       "no tickets de soporte a usuarios finales.")
    print("  [OK] Seccion 8")

    # =====================================================================
    # 9. METODOLOGIA DE LA INVESTIGACION
    # =====================================================================
    h1("9. Metodologia de la Investigacion")
    tb(
        ["Dimension", "Enfoque", "Aplicacion en QAMS"],
        [
            ["Tipo de Estudio",
             "Investigacion Tecnologica-Aplicada Proyectiva\n(Hurtado de Barrera, 2010)",
             "Genera artefacto de software innovador con hipotesis medibles: "
             "DRE >= 90%, TCO reduccion >= 90%, latencia Fast Runner < 200ms."],
            ["Metodo 1", "Analitico-Sintetico",
             "Diseccion del ISTQB CTFL v4.0 (Syllabus 2023, 76 pp.) e ISO 29119 "
             "en Entidades de Dominio, Servicios de Aplicacion y endpoints REST."],
            ["Metodo 2", "Comparativo / Revision Sistematica (SLR)",
             "Revision SLR 2019-2024 y benchmark multicriterio de herramientas "
             "(TestRail, Zephyr, HP ALM) para fundamentar superioridad tecnica y economica."],
            ["Metodo 3", "Experimental Computacional",
             "Pruebas de carga K6 (10/30/50 VUs concurrentes); "
             "analisis estadistico P50/P95/P99 de latencia de endpoints criticos."],
            ["Tecnica 1", "BDD / Gherkin",
             "Requisitos funcionales especificados en Given-When-Then, "
             "verificables por el propio sistema QAMS."],
            ["Tecnica 2", "Domain-Driven Design Light",
             "Capa Domain sin dependencias externas; invariantes de negocio "
             "en entidades C# con Value Objects y Domain Events."],
            ["Tecnica 3", "Documentacion CI/CD",
             "Dockerfiles multi-stage con builds reproducibles; "
             "docker-compose.yml declarativo con healthchecks automaticos."],
        ],
        [1.1, 1.8, 3.6]
    )
    print("  [OK] Seccion 9")

    # =====================================================================
    # 10. PROPUESTA DE SOLUCION
    # =====================================================================
    h1("10. Propuesta de Solucion — Los 6 Pilares de QAMS")
    tx("En respuesta al problema validado, QAMS se fundamenta en seis pilares "
       "tecnicos no negociables que garantizan su superioridad frente a las "
       "soluciones del mercado:")
    tb(
        ["Pilar", "Principio", "Implementacion Tecnica", "RNF Satisfecho"],
        [
            ["P1 Arquitectura",
             "Monolito Modular + Clean Architecture",
             ".NET 9: Domain -> Application -> Infrastructure -> API (4 capas, "
             "SOLID, Dependency Injection, interfaces en todo)",
             "RNF-03 Confiabilidad (ACID)\nRNF-04 Mantenibilidad"],
            ["P2 Backend Seguro",
             "API RESTful OWASP Top 10",
             "JWT Bearer, HasPermission RBAC, BCrypt factor 11, AES-256-CBC, "
             "EF Core parameterized queries, Serilog Audit Trail",
             "RNF-01 Seguridad"],
            ["P3 Frontend Reactivo",
             "SPA sin Zone.js",
             "Angular 19 Standalone + Signals: Fast Runner P/F/B < 200ms, "
             "lazy loading por ruta, AuthGuard + EncryptionInterceptor",
             "RNF-08 Usabilidad"],
            ["P4 Estandar ISTQB",
             "100% cobertura syllabus v4.0",
             "7 modulos: RTM, Static Testing, SBTM, BDD/Gherkin, "
             "Quality Gates DDP/DRE/MTTR, Defect Lifecycle, Dashboard KPIs",
             "RNF-06 Conformidad normativa"],
            ["P5 Modelo de Datos",
             "Gobernanza relacional plena",
             "PostgreSQL 16: 23 entidades en 3FN, IAuditable+ISoftDelete, "
             "JSONB Quality Gates, indices compuestos en FK criticas",
             "RNF-03 Confiabilidad"],
            ["P6 Despliegue",
             "Contenedorizacion One-Click",
             "Docker multi-stage Alpine: API 120MB + Frontend 25MB + "
             "PostgreSQL + Redis; docker compose up -d",
             "RNF-05 Escalabilidad\nOE5"],
        ],
        [0.8, 1.4, 2.5, 1.8]
    )
    print("  [OK] Seccion 10")

    # =====================================================================
    # 11. CRONOGRAMA
    # =====================================================================
    h1("11. Cronograma de Ejecucion — 22 Semanas")
    tb(
        ["Fase", "Actividades Principales", "Entregables Verificables", "Sem.", "OE"],
        [
            ["Fase 1\nAnalisis y\nModelo",
             "SLR ISTQB/ISO 29119. Especificacion RF en Gherkin. "
             "Modelado ERD 3FN PostgreSQL 16. Decisiones arquitectura (ADRs).",
             "ERD normalizado. Catalogo RF Gherkin. ADRs documentados.",
             "1 - 4", "OE1"],
            ["Fase 2\nBackend\n.NET 9",
             "Solucion Clean Architecture 4 capas. Entidades, repositorios, "
             "Unit of Work, migraciones EF Core 9. Auth JWT + RBAC + AES-256.",
             "API REST funcional C#. Coleccion Postman. Health Check endpoint.",
             "5 - 10", "OE2"],
            ["Fase 3\nFrontend\nAngular 19",
             "Angular 19 strict + Standalone + Signals. Fast Runner P/F/B. "
             "Kanban drag-and-drop. Dashboard Ng2-Charts.",
             "SPA Angular operativa. Fast Runner < 200ms. Dashboard KPIs.",
             "11 - 16", "OE3"],
            ["Fase 4\nModulos\nISTQB",
             "Static Testing (walkthroughs/inspecciones). SBTM Charters. "
             "RTM bidireccional. Quality Gates DDP/DRE/MTTR semaforo.",
             "7 modulos ISTQB integrados. RTM y Quality Gate operativos.",
             "17 - 19", "OE4"],
            ["Fase 5\nDocker y\nMonografia",
             "Dockerfiles multi-stage Alpine. docker-compose.yml 4 servicios. "
             "Monografia academica completa (7 capitulos).",
             "Stack Docker 1 comando. proyecto.docx finalizado.",
             "20 - 22", "OE5"],
        ],
        [0.9, 2.5, 1.5, 0.5, 0.4]
    )
    tx("Duracion total: 22 semanas (5 meses y medio). Modalidad: desarrollo "
       "individual con revisiones quincenales del tutor academico. Repositorio: "
       "GitHub publico con historial de commits documentado.", sz=10, sa=4)
    print("  [OK] Seccion 11")

    # =====================================================================
    # 12. RECURSOS Y VIABILIDAD
    # =====================================================================
    h1("12. Recursos y Analisis de Viabilidad")
    tb(
        ["Componente", "Tecnologia / Herramienta", "Version", "Rol en el Proyecto"],
        [
            ["Backend",           "C# / ASP.NET Core",     ".NET 9",     "API RESTful + Clean Architecture + JWT + RBAC + AES-256"],
            ["Frontend",          "TypeScript / Angular",  "Angular 19", "SPA Signals + Standalone + Fast Runner + Dashboard"],
            ["Base de Datos",     "PostgreSQL",            "v16 Alpine", "23 entidades 3FN + JSONB + indices compuestos"],
            ["Cache / Colas",     "Redis",                 "v7 Alpine",  "Cache de sesion + cola de emails asincronos"],
            ["Servidor Web",      "Nginx",                 "Alpine",     "Servidor estaticos Angular + proxy inverso"],
            ["Contenedorizacion", "Docker + Compose",      "Docker 25+", "4 servicios; Dockerfile multi-stage; healthchecks"],
            ["ORM",               "Entity Framework Core", "v9",         "Migraciones + Repository Pattern + Unit of Work"],
            ["Pruebas Backend",   "xUnit + Moq",           "v2 / v4",   "Unit tests + Integration tests (cobertura 76%)"],
            ["Pruebas de Carga",  "k6 by Grafana",         "v0.51",      "Escenarios 10/30/50 VUs; P50/P95/P99 < 250ms"],
        ],
        [1.3, 1.6, 0.9, 2.7]
    )
    h2("Analisis de Viabilidad")
    bl("Tecnica: Stack maduro (.NET 9, Angular 19, PostgreSQL, Docker) con documentacion "
       "extensa y soporte empresarial activo. Sin dependencias de hardware especializado.")
    bl("Economica: Cero costo de licencias de software. Infraestructura $40/mes VPS. "
       "ROI positivo desde el primer mes frente a TestRail o Zephyr.")
    bl("Academica: Cumple todos los requisitos del Diplomado: sistema completo, metodologia "
       "agil documentada, estandares internacionales (ISTQB/ISO/OWASP) y monografia formal.")
    bl("Temporal: 22 semanas ejecutables con 4-6 horas semanales de dedicacion, "
       "respaldado por el historial real de commits del repositorio GitHub del proyecto.")
    print("  [OK] Seccion 12")

    # =====================================================================
    # 13. REFERENCIAS BIBLIOGRAFICAS
    # =====================================================================
    h1("13. Referencias Bibliograficas")
    refs = [
        "ISTQB. (2023). Certified Tester Foundation Level (CTFL) Syllabus v4.0. ISTQB General Assembly. https://www.istqb.org",
        "ISO/IEC. (2022). ISO/IEC/IEEE 29119-1:2022 — Software testing — Part 1: General concepts. ISO/IEC.",
        "ISO/IEC. (2023). ISO/IEC 25010:2023 — Systems and software Quality Requirements and Evaluation (SQuaRE). ISO/IEC.",
        "OWASP. (2021). OWASP Top 10 2021: The Ten Most Critical Web Application Security Risks. https://owasp.org/Top10/",
        "IEEE. (2008). IEEE Std 829-2008: Standard for Software and System Test Documentation. IEEE.",
        "Martin, R. C. (2017). Clean Architecture: A Craftsman Guide to Software Structure and Design. Prentice Hall.",
        "CISQ. (2022). The Cost of Poor Software Quality in the US. Consortium for IT Software Quality.",
        "IBM Systems Sciences Institute. (2000). Relative Cost to Fix Defects by Phase. IBM.",
        "Jones, C. (2008). Applied Software Measurement (3rd ed.). McGraw-Hill.",
        "Sanchez, D. F. (2023). Sistema web para gestion de pruebas en DevOps. Tesis Master, UPV Espania.",
        "Gomez, R. & Perez, L. (2022). Plataforma QA en organizaciones gubernamentales. Trabajo de Grado, UAEM Mexico.",
        "Martinez, A. S. (2024). Open Source QA para PYMEs iberoamericanas. Rev. Iberoam. Ing. SW, 12(3).",
        "Microsoft. (2025). ASP.NET Core 9.0 Documentacion Oficial. https://learn.microsoft.com/aspnet/core",
        "Google Angular Team. (2025). Angular 19 — Signals y Standalone Components. https://angular.dev",
        "PostgreSQL Global Development Group. (2024). PostgreSQL 16 Documentation. https://www.postgresql.org/docs/16/",
        "Docker Inc. (2024). Docker Engine 25 y Compose Specification v2. https://docs.docker.com",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent       = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after       = Pt(3)
        r = p.add_run(f"{i}. {ref}")
        r.font.size = Pt(9.5); r.font.name = "Calibri"
    print("  [OK] Referencias")

    # GUARDAR
    doc.save(OUT1)
    print(f"\n  GUARDADO: {OUT1}")
    try:
        doc.save(OUT2)
        print(f"  GUARDADO: {OUT2}")
    except Exception as e:
        alt = r"c:/diplomado/perfil_backup.docx"
        doc.save(alt)
        print(f"  GUARDADO (backup): {alt} — {e}")


if __name__ == "__main__":
    print("Generando Perfil del Proyecto QAMS...")
    build()
    print("\nDone.")
