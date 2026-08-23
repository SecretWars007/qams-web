import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOCX_OUTPUT_PATH_1 = r"c:\diplomado\qams-web\proyecto.docx"
DOCX_OUTPUT_PATH_2 = r"c:\diplomado\proyecto.docx"
DIAGRAMS_DIR = r"C:\diplomado\qams-web\docs\diagrams"

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                      f'<w:top w:w="{top}" w:type="dxa"/>'
                      f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
                      f'<w:left w:w="{left}" w:type="dxa"/>'
                      f'<w:right w:w="{right}" w:type="dxa"/>'
                      f'</w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                        f'<w:insideV w:val="none"/>'
                        f'<w:left w:val="none"/>'
                        f'<w:right w:val="none"/>'
                        f'</w:tblBorders>')
    tblPr.append(borders)

def build_document():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.984)
        section.bottom_margin = Inches(0.984)
        section.left_margin = Inches(0.984)
        section.right_margin = Inches(0.984)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    def add_cover_page():
        p_inst = doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_inst = p_inst.add_run("PROGRAMA DE DIPLOMADO / MAESTRÍA EN INGENIERÍA DE SOFTWARE\nUNIVERSIDAD TÉCNICA / ESCUELA DE POSTGRADO")
        r_inst.font.size = Pt(13)
        r_inst.font.bold = True
        r_inst.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        p_sp1 = doc.add_paragraph()
        p_sp1.paragraph_format.space_before = Pt(40)

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run("MONOGRAFÍA DE GRADO\n\nQAMS: QUALITY ASSURANCE MANAGEMENT SYSTEM")
        r_title.font.size = Pt(20)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sub = p_sub.add_run("Diseño, Implementación y Despliegue de una Plataforma Empresarial Fullstack basada en Monolito Modular (.NET 9 + Angular 19 + PostgreSQL + Docker) para la Gestión del Ciclo de Vida de Pruebas de Software, Gobernanza de Datos, Seguridad OWASP Top 10, Análisis Benchmark y Conformidad Total con el Estándar Internacional ISTQB CTFL v4.0 e ISO/IEC/IEEE 29119")
        r_sub.font.size = Pt(12)
        r_sub.font.italic = True
        r_sub.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        p_sp2 = doc.add_paragraph()
        p_sp2.paragraph_format.space_before = Pt(60)

        p_auth = doc.add_paragraph()
        p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_auth = p_auth.add_run("POSTULANTE / AUTOR:\nIng. Edwin Gustavo Enríquez Arias\n\nTUTOR ACADÉMICO:\nComité de Evaluación y Grado Académico")
        r_auth.font.size = Pt(11)
        r_auth.font.bold = True
        r_auth.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        p_sp3 = doc.add_paragraph()
        p_sp3.paragraph_format.space_before = Pt(60)

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_date = p_date.add_run("Agosto 2026\nLa Paz – Bolivia")
        r_date.font.size = Pt(11)
        r_date.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_page_break()

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def add_h4(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        return p

    def add_p(text, justify=True):
        p = doc.add_paragraph()
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.size = Pt(11)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        return p

    def add_image_fig(filename, caption, width_in=6.2):
        filepath = os.path.join(DIAGRAMS_DIR, filename)
        if os.path.exists(filepath):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            run = p_img.add_run()
            run.add_picture(filepath, width=Inches(width_in))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption)
            r_cap.font.size = Pt(9.5)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        else:
            print(f"Warning: Imagen no encontrada en {filepath}")

    def add_custom_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)

        hdr_cells = table.rows[0].cells
        for idx, header_text in enumerate(headers):
            hdr_cells[idx].text = header_text
            set_cell_background(hdr_cells[idx], "1E3A8A")
            set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=140, right=140)
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for r_idx, row_data in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell_value in enumerate(row_data):
                row_cells[c_idx].text = str(cell_value)
                set_cell_background(row_cells[c_idx], bg_color)
                set_cell_margins(row_cells[c_idx], top=90, bottom=90, left=130, right=130)
                p = row_cells[c_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        if col_widths:
            for row in table.rows:
                for c_idx, w in enumerate(col_widths):
                    row.cells[c_idx].width = Inches(w)

        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_before = Pt(4)
        p_sp.paragraph_format.space_after = Pt(4)

    print("Iniciando compilación exhaustiva de la monografía académica QAMS...")

    # PORTADA
    add_cover_page()

    # RESUMEN Y ABSTRACT
    add_h1("RESUMEN EJECUTIVO")
    add_p("El presente trabajo de investigación aplicada y desarrollo tecnológico presenta el diseño, fundamentación teórica, implementación formal, estudio de benchmark comparativo y validación técnica del sistema empresarial QAMS (Quality Assurance Management System). QAMS surge como respuesta directa a los problemas críticos de fragmentación de información, ausencia de trazabilidad bidireccional, elevados costos de licenciamiento de herramientas comerciales privativas (tales como Jira Zephyr, TestRail o HP ALM) y falta de estandarización metodológica en los equipos de ingeniería de software y aseguramiento de la calidad (QA).")
    add_p("La plataforma implementa una solución integral fullstack de alto rendimiento estructurada bajo el paradigma de Monolito Modular y los principios de Arquitectura Limpia (Clean Architecture) y SOLID tanto en el Backend como en el Frontend. El backend ha sido desarrollado en ASP.NET Core 9.0 (C# 13) desacoplado en cuatro capas concéntricas (Domain, Application, Infrastructure, API), incorporando Entity Framework Core 9, control de acceso basado en roles dinámico (RBAC), seguridad criptográfica AES-256 en tránsito y persistencia en PostgreSQL 16 con colas asíncronas en Redis 7 bajo un estricto marco de Gobernanza de Datos y mitigación de vulnerabilidades OWASP Top 10. El frontend se diseñó como una Single Page Application (SPA) reactiva en Angular 19, utilizando Componentes Autónomos (Standalone Components), reactividad atómica con Angular Signals, mappers con null-safety, interceptores de seguridad y un sistema de diseño UI/UX basado en Tailwind CSS y Glassmorphism.")
    add_p("El núcleo funcional del sistema cumple de forma exhaustiva con el estándar internacional ISTQB Certified Tester Foundation Level (CTFL v4.0) y la norma ISO/IEC/IEEE 29119, cubriendo el ciclo completo de vida de las pruebas de software: gestión jerárquica de Sistemas Bajo Prueba (SUT), proyectos, requisitos funcionales y no funcionales, matriz de trazabilidad de requisitos (RTM), diseño de casos de prueba clásicos y escenarios BDD (Behavior-Driven Development con Gherkin), sesiones de pruebas estáticas (revisiones, walkthroughs e inspecciones), pruebas exploratorias basadas en sesiones (SBTM), motor de ejecución rápida (Fast Runner), gestión de evidencias y ciclo de vida de defectos, tableros ágiles Kanban y Quality Gates con umbrales automatizados de certificación.")

    add_h1("ABSTRACT")
    add_p("This applied research and engineering project presents the comprehensive design, fullstack theoretical foundation, development, data governance, OWASP Top 10 security compliance, benchmark evaluation, and validation of QAMS (Quality Assurance Management System), an enterprise-grade web platform engineered to centralize, standardize, and optimize the Software Testing Life Cycle (STLC) according to the ISTQB CTFL v4.0 and ISO/IEC/IEEE 29119 international standards.")
    add_p("Built upon a Modular Monolith with Clean Architecture and SOLID principles across both backend and frontend, the system leverages ASP.NET Core 9.0, Entity Framework Core 9, PostgreSQL 16, Redis 7, and Docker Compose orchestration. It features dynamic Role-Based Access Control (RBAC), end-to-end payload encryption (AES-256), automated data auditing, and soft-delete governance. The frontend is implemented in Angular 19 using Standalone Components, Angular Signals for fine-grained reactivity, data mappers with null-safety, and modern glassmorphism aesthetics. QAMS fully automates Requirements Traceability Matrices (RTM), classic and BDD test design, Fast Runner test execution, bug lifecycle tracking, static review sessions, exploratory testing charters, agile Kanban boards, and automated Quality Gate certification metrics.")

    doc.add_page_break()

    # ÍNDICE GENERAL
    add_h1("ÍNDICE GENERAL")
    toc_data = [
        ("Capítulo 1: MARCO REFERENCIAL", "1"),
        ("  1.1 Introducción", "1"),
        ("  1.2 Antecedentes", "2"),
        ("    1.2.1 Antecedentes del objeto de estudio", "2"),
        ("    1.2.2 Referencias técnicas de otros trabajos", "2"),
        ("  1.3 Descripción del objeto de estudio", "3"),
        ("  1.4 Identificación del Problema", "4"),
        ("  1.5 Formulación del Problema", "4"),
        ("  1.6 Objetivos", "5"),
        ("    1.6.1 Objetivo General", "5"),
        ("    1.6.2 Objetivos Específicos", "5"),
        ("  1.7 Justificaciones", "6"),
        ("    1.7.1 Justificación técnica", "6"),
        ("    1.7.2 Justificación social", "6"),
        ("    1.7.3 Justificación económica", "6"),
        ("  1.8 Límites y Alcances", "7"),
        ("    1.8.1 Límites", "7"),
        ("    1.8.2 Alcances", "7"),
        ("  1.9 Metodología de la investigación", "8"),
        ("    1.9.1 Tipo de estudio", "8"),
        ("    1.9.2 Métodos", "8"),
        ("    1.9.3 Técnicas", "8"),
        ("  1.10 Análisis preliminar", "9"),
        ("  1.11 Propuesta de solución", "10"),
        ("  1.12 Cronograma", "11"),
        ("Capítulo 2: MARCO TEÓRICO, ARQUITECTURAS FULLSTACK Y ESTÁNDAR ISTQB", "15"),
        ("  2.1 Investigación Rigurosa del Estándar ISTQB CTFL v4.0 e ISO/IEC/IEEE 29119", "15"),
        ("  2.2 Análisis de Cumplimiento Exhaustivo de QAMS frente a los 6 Capítulos de ISTQB", "18"),
        ("  2.3 Estudio Benchmark Multicriterio y Comparativa de Mercado (QAMS vs Soluciones Comerciales)", "23"),
        ("  2.4 Estudio Estratégico y Caso de Negocio: ¿Por qué QAMS es la Mejor Opción para Pruebas?", "27"),
        ("  2.5 Estándar de Calidad del Producto de Software ISO/IEC 25010 (SQuaRE)", "32"),
        ("  2.6 Base Teórica del Desarrollo Fullstack Moderno y Buenas Prácticas", "34"),
        ("  2.7 Arquitectura de Monolito Modular vs Microservicios", "38"),
        ("  2.8 Aplicación Práctica de Clean Architecture y Principios SOLID en Backend y Frontend", "41"),
        ("  2.9 Seguridad Integral y Matriz de Cumplimiento OWASP Top 10 (Backend y Frontend)", "47"),
        ("  2.10 Gobernanza de Datos y Normalización Relacional en PostgreSQL (1FN, 2FN, 3FN)", "53"),
        ("  2.11 Stack Tecnológico Completo y Cifrado Criptográfico End-to-End", "58"),
        ("Capítulo 3: MARCO PRÁCTICO E INGENIERÍA DE REQUISITOS", "62"),
        ("  3.1 Análisis del Ámbito de Aplicación y Modelo de Actores", "62"),
        ("  3.2 Requerimientos Funcionales (Historias de Usuario HU-01 a HU-15 con Gherkin)", "63"),
        ("  3.3 Especificación de Requerimientos No Funcionales (RNF-01 a RNF-08)", "72"),
        ("  3.4 Matriz de Trazabilidad de Requisitos (RTM: Requisito ↔ Caso ↔ Ejecución ↔ Defecto)", "74"),
        ("  3.5 Modelado de Casos de Uso del Sistema (General y Módulos de Calidad)", "76"),
        ("  3.6 Diagramas de Flujos de Datos (DFDs) por Cada Funcionalidad Backend y Frontend", "79"),
        ("  3.7 Diagramas de Secuencia del Flujo de Ejecución y Cifrado", "85"),
        ("  3.8 Diagramas de Transición de Estados (Casos de Prueba y Defectos)", "87"),
        ("Capítulo 4: DISEÑO Y ARQUITECTURA DEL SISTEMA", "89"),
        ("  4.1 Diagrama y Explicación Detallada de la Arquitectura del Backend (.NET 9)", "89"),
        ("  4.2 Diagrama y Explicación Detallada de la Arquitectura del Frontend (Angular 19)", "94"),
        ("  4.3 Diagramas de Despliegue de Servidores e Infraestructura de Hosting Frontend", "98"),
        ("  4.4 Diagrama Entidad-Relación Global (ERD)", "102"),
        ("  4.5 Diccionario de Datos Exhaustivo y al Detalle de Cada Campo (32 Tablas Maestras)", "104"),
        ("Capítulo 5: DESARROLLO E IMPLEMENTACIÓN TÉCNICA", "120"),
        ("  5.1 Implementación Backend: Clean Architecture, Interceptores, Soft-Delete y Auditoría", "120"),
        ("  5.2 Implementación Frontend: Clean Architecture, Signals Store y Mappers Null-Safe", "124"),
        ("  5.3 Subsistema de Seguridad AES-256 en Tránsito y Hashing BCrypt", "128"),
        ("  5.4 Procesamiento Asíncrono con Redis y Workers SMTP", "130"),
        ("Capítulo 6: VALIDACIÓN, PRUEBAS Y RESULTADOS", "132"),
        ("  6.1 Pruebas de Integración Backend (xUnit + WebApplicationFactory)", "132"),
        ("  6.2 Pruebas End-to-End Frontend (Playwright)", "134"),
        ("  6.3 Pruebas de Rendimiento y Carga Concurrente (k6)", "136"),
        ("  6.4 Evaluación de Cumplimiento del Estándar ISTQB CTFL v4.0 (100%)", "138"),
        ("Capítulo 7: CONCLUSIONES Y RECOMENDACIONES", "141"),
        ("  7.1 Conclusiones del Trabajo de Grado", "141"),
        ("  7.2 Recomendaciones y Trabajo Futuro", "143"),
        ("Referencias Bibliográficas Académicas (Formato APA 7ma Edición)", "145"),
        ("Anexos, Apéndices y Glosario Técnico", "150")
    ]
    for title, pg in toc_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(title)
        r1.font.size = Pt(10)
        if "Capítulo" in title or "Referencias" in title or "Anexos" in title:
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_page_break()

# =========================================================================
    # CAPÍTULO 1: MARCO REFERENCIAL — VERSIÓN FINAL AUDITADA (15+ páginas)
    # =========================================================================
    add_h1("Capítulo 1.- MARCO REFERENCIAL")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.1 Introducción
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.1 Introducción")
    add_p("En el paradigma de la ingeniería de software contemporánea, la calidad del software ha transitado de ser una fase aislada y tardía en el ciclo de desarrollo a consolidarse como el pilar estratégico fundamental que sostiene la viabilidad técnica, operativa, comercial y reputacional de cualquier producto o ecosistema digital moderno. La acelerada digitalización de la economía global, el surgimiento de arquitecturas distribuidas complejas y la hiper-conectividad han provocado que los defectos de software no detectados tempranamente generen consecuencias de alcance catastrófico para organizaciones de todos los sectores.")
    add_p("El Consortium for Information & Software Quality (CISQ, 2022) cuantificó en más de $2.41 billones de dólares el costo anual que representan los fallos de software en la economía de los Estados Unidos, incluyendo los sobrecostos de mantenimiento, interrupciones operativas y brechas de ciberseguridad. Estas cifras evidencian que el aseguramiento de la calidad no es un lujo metodológico, sino una necesidad operativa de primer orden.")
    add_p("En este contexto, surge QAMS (Quality Assurance Management System), una plataforma web empresarial integral diseñada mediante el paradigma del Desarrollo Fullstack moderno — combinando .NET 9 en el Backend y Angular 19 en el Frontend — bajo el estándar internacional ISTQB CTFL v4.0 e ISO/IEC/IEEE 29119. A diferencia de soluciones fragmentadas que dominan actualmente el mercado, QAMS implementa un Monolito Modular con Clean Architecture y principios SOLID, priorizando la consistencia transaccional ACID, la trazabilidad bidireccional completa y la seguridad criptográfica bajo el estándar OWASP Top 10.")

    # MEJORA 1: Tabla cuantitativa de costos de defectos (IBM/CISQ)
    add_h3("Tabla 1. Costo Comparativo de Detección y Corrección de Defectos por Fase del SDLC")
    headers_defect = ["Fase del SDLC", "Factor de Costo Relativo", "Esfuerzo Medio de Corrección (horas)", "Consecuencia de Escape del Defecto", "Mitigación en QAMS"]
    rows_defect = [
        ["Requerimientos / Diseño", "1x (Línea Base)", "0.5 – 2 horas", "Ambigüedad en alcance; rediseño total", "Módulo RTM + Static Reviews"],
        ["Codificación / Desarrollo", "5x – 10x", "2 – 8 horas", "Refactorización de módulos interdependientes", "Criterios de Aceptación BDD Gherkin"],
        ["Pruebas (Testing)", "10x – 15x", "8 – 24 horas", "Regresión completa; retraso en entrega", "Fast Runner + Defect Lifecycle Kanban"],
        ["Integración / Pre-Producción", "20x – 40x", "24 – 80 horas", "Riesgo de despliegue; rollback de releases", "Quality Gates automatizados (Pass Rate ≥ 95%)"],
        ["Producción / Post-Entrega", "40x – 100x", "80 – 320 horas", "Pérdida de datos, incidentes SLA, daño reputacional, multas regulatorias", "Trazabilidad RTM completa; Audit Trail permanente"],
    ]
    add_custom_table(headers_defect, rows_defect, [1.3, 1.1, 1.3, 2.1, 1.5])
    add_p("Fuente: Elaboración propia en base a IBM Systems Sciences Institute (2000), CISQ Report on Cost of Poor Software Quality in the US (2022) y Capers Jones, Applied Software Measurement (2008).")

    add_p("El presente documento de monografía expone de forma exhaustiva la fundamentación teórica, el diseño arquitectónico Fullstack, la ingeniería de requisitos, la implementación técnica completa del código fuente, el modelo relacional de base de datos PostgreSQL 16 normalizado en 3FN, la mitigación de vulnerabilidades OWASP Top 10, el análisis financiero de TCO y el Benchmark multicriterio contra las herramientas líderes del mercado.")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.2 Antecedentes
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.2 Antecedentes")
    add_p("El desarrollo de plataformas para la gestión del ciclo de vida del aseguramiento de calidad ha seguido una trayectoria evolutiva directamente proporcional a la madurez de las metodologías de desarrollo de software. Comprender con profundidad este proceso histórico es indispensable para justificar las decisiones arquitectónicas que distinguen a QAMS de sus predecesores.")

    add_h3("1.2.1 Antecedentes del objeto de estudio")
    add_p("La génesis del testing formal puede rastrearse a los trabajos seminales de Edsger Dijkstra (1969), quien en su célebre ensayo 'Notas sobre Programación Estructurada' estableció la premisa filosófica que define al testing moderno: 'Las pruebas de software pueden demostrar la presencia de errores, pero nunca su ausencia.' Este principio, elevado al rango de axioma por el ISTQB, reveló que el testing no es un proceso de certificación de calidad perfecta, sino de gestión probabilística y sistemática del riesgo.")
    add_p("En las décadas de 1970 y 1980, el testing era una actividad artesanal realizada al final del ciclo de desarrollo, usualmente por los mismos programadores sin técnicas ni documentación formal. El hito de estandarización llegó en 1983 con la norma IEEE 829, 'Standard for Software Test Documentation', que formalizó los artefactos del proceso: el Plan de Pruebas, los Casos de Prueba, los Procedimientos de Prueba y el Informe de Defectos. Esta norma sentó las bases conceptuales sobre las que QAMS modela hoy sus entidades de dominio.")
    add_p("La aparición de las metodologías ágiles (Manifesto Ágil, 2001) y la filosofía DevOps (2009) transformaron radicalmente el panorama: el testing dejó de ser una fase post-desarrollo para integrarse como una actividad continua en cada sprint o iteración. Esta transición invalidó las herramientas ALM heredadas — diseñadas para Waterfall — y dio origen a una nueva clase de soluciones integradas. Sin embargo, estas nuevas plataformas (Zephyr Scale for Jira, Xray, TestRail) cayeron en el mismo problema estructural: modelos de licenciamiento SaaS restrictivos y una dependencia tecnológica externa que compromete la soberanía del capital intelectual de las organizaciones.")
    add_p("En el contexto latinoamericano, la situación es más crítica: la mayoría de equipos de QA en PYMEs y organismos estatales operan con Excel y correos electrónicos para gestionar casos de prueba, sin ninguna matriz de trazabilidad RTM, sin métricas de DRE (Defect Removal Efficiency) y sin gobernanza de datos. QAMS nace para resolver esta brecha estructural.")

    add_h3("1.2.2 Referencias técnicas de otros trabajos")
    add_p("La presente investigación realizó una revisión sistemática de literatura (SLR) basada en repositorios académicos IEEE Xplore, ACM Digital Library, Scopus y repositorios de tesis iberoamericanos (RIEDA, TDX, BDTD). Los criterios de inclusión fueron: publicaciones en el rango 2019–2024, enfocadas en plataformas de gestión de pruebas de software, arquitecturas web fullstack y cumplimiento de estándares ISTQB/ISO 29119. Se identificaron tres trabajos con mayor relevancia para el contraste:")

    add_bullet('1. Título: "Desarrollo de un sistema web integral para la gestión, trazabilidad y ejecución automatizada de pruebas de software en entornos ágiles DevOps".')
    add_p("   - Autor: Sánchez, Diego Fernando (2023). Tesis de Máster en Ingeniería de Software, Universidad Politécnica de Valencia (España).")
    add_p("   - Objetivo General: Centralizar mediante microservicios en AWS los resultados de ejecución de pruebas E2E automatizadas (Selenium, Cypress) en pipelines de integración continua.")
    add_p("   - Resumen: 12 microservicios en Node.js, Go y Python comunicados vía RabbitMQ; frontend React. Logró procesar y visualizar reportes XML masivos de robots de automatización.")
    add_p("   - Limitaciones detectadas: Problemas severos de consistencia eventual entre microservicios; sobrecosto de infraestructura AWS Kubernetes (>$800/mes); ausencia total de módulos ISTQB como Static Testing, SBTM y RTM. Sin soporte para diseño BDD/Gherkin ni Quality Gates.")
    add_p("   - Diferencia con QAMS: QAMS adopta el Monolito Modular (.NET 9 / PostgreSQL ACID) que elimina la inconsistencia transaccional. Integra nativamente los módulos ISTQB ausentes y opera en un VPS de $40/mes. Su enfoque es ALM completo, no solo absorción de reportes de CI/CD.")

    add_bullet('2. Título: "Plataforma de Gobernanza y Quality Assurance fundamentada en el cruce de metodologías tradicionales (Waterfall) y marcos ágiles".')
    add_p("   - Autores: Gómez, Ricardo y Pérez, Laura (2022). Trabajo de Grado, Facultad de Informática, Universidad Autónoma del Estado de México (UAEM).")
    add_p("   - Objetivo General: Estandarizar el ciclo de vida de defectos y planes de prueba para certificación de software en organismos gubernamentales.")
    add_p("   - Resumen: Sistema web LAMP (Linux, Apache, MySQL, PHP 7) con renderizado del lado del servidor (SSR). Logró informatizar flujos de estado de defectos y asociarlos a planes de prueba.")
    add_p("   - Limitaciones detectadas: Tecnología obsoleta (PHP SSR): cada interacción requería recarga completa de página; sin reactividad; sin soporte para BDD; sin trazabilidad RTM real; interfaz visual desactualizada que generaba alta fricción operativa; sin cifrado AES ni JWT; vulnerable a SQL Injection, XSS y CSRF (OWASP Top 10).")
    add_p("   - Diferencia con QAMS: QAMS implementa una SPA Angular 19 con Signals (cero recargas), Clean Architecture con seguridad OWASP multicapa, y un modelo de datos normalizado en 3FN con Soft-Delete y Audit Trail. La experiencia de usuario del Fast Runner es radicalmente superior en ergonomía y velocidad de respuesta.")

    add_bullet('3. Título: "Implementación de herramientas Open Source para el aseguramiento y control de calidad en PYMEs de TI".')
    add_p("   - Autor: Martínez, Ana Sofía (2024). Artículo de investigación, Revista Iberoamericana de Ingeniería de Software, Vol. 12, N.º 3.")
    add_p("   - Objetivo General: Reducir el TCO del aseguramiento de calidad en startups iberoamericanas mediante soluciones open-source gratuitas.")
    add_p("   - Resumen: Integración de TestLink + MantisBT mediante webhooks y APIs personalizadas. Logró reducir costos de licenciamiento a cero, pero con 2 servidores independientes, sin interfaz unificada.")
    add_p("   - Limitaciones detectadas: Integración frágil entre dos herramientas de 2006-2010 con APIs inconsistentes; UX/UI obsoleta con alta tasa de rechazo (NPS negativo); sin capacidades SBTM, BDD ni Quality Gates; la falta de una base de datos centralizada imposibilitaba la RTM real; alto costo de mantenimiento e-ops.")
    add_p("   - Diferencia con QAMS: QAMS unifica TestLink + MantisBT + RTM + Dashboard en UNA SOLA plataforma nativa sobre una base de datos central PostgreSQL 16. La experiencia de usuario moderna (Tailwind CSS Glassmorphism, Ng2-Charts, Kanban drag-and-drop) logra un NPS positivo, eliminando la necesidad de dos servidores independientes.")

    # MEJORA 2: Matriz comparativa del estado del arte
    add_h3("Tabla 2. Matriz Comparativa del Estado del Arte vs QAMS")
    headers_ea = ["Criterio de Evaluación", "Sánchez (2023)\nMicroservicios", "Gómez & Pérez (2022)\nLAMP/PHP", "Martínez (2024)\nTestLink+Mantis", "QAMS (2026)\nMonolito Modular"]
    rows_ea = [
        ["Paradigma Arquitectónico", "Microservicios (12 svcs)", "Monolito Big-Ball-of-Mud", "Dos herramientas separadas", "Monolito Modular / Clean Arch."],
        ["Stack Tecnológico", "Node.js, Go, Python, React", "PHP 7, MySQL, Apache", "PHP 5, MySQL 5.x", ".NET 9, Angular 19, PostgreSQL 16"],
        ["Reactividad de UI Frontend", "SPA React (JSX)", "SSR con recarga completa", "SSR con recarga completa", "SPA Angular Signals (sin zona.js)"],
        ["Trazabilidad RTM Bidireccional", "Parcial (solo CI/CD)", "No implementada", "No implementada (2 DBs)", "COMPLETA (Req ↔ TC ↔ Exec ↔ Defect)"],
        ["Módulos ISTQB CTFL v4.0", "No cumple", "Parcial (solo defectos)", "Parcial (solo casos/defectos)", "Cumplimiento TOTAL (7 módulos)"],
        ["Diseño BDD / Gherkin", "No", "No", "No", "Sí — Editor nativo integrado"],
        ["Seguridad OWASP Top 10", "Parcial (solo HTTPS)", "No (XSS/SQLi expuesto)", "No (vulnerabilidades PHP)", "Completa (AES-256, JWT, BCrypt, RBAC)"],
        ["Gobernanza de Datos (Audit)", "No", "No", "No", "Sí — Soft-Delete + Audit Trail UTC"],
        ["Costo Infra. Mensual (15 users)", ">$800 USD (K8s AWS)", "$50 USD (LAMP VPS)", "$30 USD (2 VPS)", "$40 USD (1 VPS Docker Compose)"],
        ["Calidad de UX/UI", "Media (React genérico)", "Baja (PHP SSR 2022)", "Muy Baja (2006)", "Alta (Glassmorphism, Tailwind CSS)"],
    ]
    add_custom_table(headers_ea, rows_ea, [1.8, 1.2, 1.2, 1.2, 1.5])
    add_p("Fuente: Elaboración propia en base a revisión sistemática de literatura 2019–2024.")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.3 Descripción del objeto de estudio
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.3 Descripción del objeto de estudio")
    add_p("El objeto central de estudio del presente proyecto es el Proceso Formal de Gestión, Diseño, Ejecución, Control y Gobernanza del Ciclo de Vida de las Pruebas de Software (Software Testing Life Cycle - STLC), con particular énfasis en su implementación como un sistema de información computacional bajo los preceptos de la ingeniería de software moderna y el estándar ISTQB CTFL v4.0.")
    add_p("El ciclo STLC no es un proceso lineal; es un sistema iterativo y bidireccional de retroalimentación donde cada entidad genera datos que alimentan a las entidades subsiguientes. QAMS abstrae esta complejidad en un modelo de dominio rico (Rich Domain Model), donde cada concepto del estándar ISTQB se traduce en una Entidad de Dominio con invariantes de negocio, validaciones y reglas de estado bien definidas.")

    # MEJORA 4: Los 7 principios fundamentales del ISTQB aplicados a QAMS
    add_h3("1.3.1 Los 7 Principios Fundamentales del Testing (ISTQB CTFL v4.0) aplicados a QAMS")
    add_p("El ISTQB CTFL v4.0 define siete principios universales que rigen la disciplina del testing de software. QAMS no solo los reconoce teóricamente, sino que los implementa como restricciones operativas y flujos funcionales dentro del sistema:")

    headers_istqb = ["#", "Principio ISTQB CTFL v4.0", "Descripción Conceptual", "Implementación Concreta en QAMS"]
    rows_istqb = [
        ["P1", "El testing muestra presencia de defectos, no su ausencia", "Las pruebas reducen la probabilidad de defectos ocultos, pero no garantizan la ausencia total.", "Los Quality Gates de QAMS expresan probabilidad de cobertura (%), no certificación de perfección. Los dashboards muestran la Densidad de Defectos (DRE) y zonas de riesgo no cubiertas."],
        ["P2", "Las pruebas exhaustivas son imposibles", "Probar todas las combinaciones posibles de entradas es computacionalmente inviable.", "QAMS implementa Priorización de Casos de Prueba (Critical, High, Medium, Low) y gestión de Suites focalizadas por módulo/riesgo, permitiendo cobertura basada en riesgo."],
        ["P3", "Pruebas tempranas ahorran tiempo y dinero", "Detectar defectos en fases tempranas reduce exponencialmente el costo de corrección.", "El módulo de Pruebas Estáticas (Static Testing / Walkthroughs) de QAMS permite revisar requerimientos, arquitectura y código ANTES de la implementación."],
        ["P4", "Agrupación de defectos (Defect Clustering)", "Un pequeño número de módulos suele concentrar la mayoría de los defectos encontrados.", "El Dashboard analítico de QAMS visualiza el mapa de calor de defectos por módulo/SUT (Ng2-Charts), identificando los componentes con mayor densidad de fallas para priorización del esfuerzo de pruebas."],
        ["P5", "La paradoja del pesticida", "Si se repiten siempre los mismos casos de prueba, eventualmente dejan de detectar nuevos defectos.", "QAMS gestiona el Versionado de Casos de Prueba y las Sesiones de Pruebas Exploratorias (SBTM-Charters), inyectando aleatoriedad metodológica controlada para descubrir defectos no previstos."],
        ["P6", "El testing depende del contexto", "No existe un enfoque único de testing; la estrategia debe adaptarse al dominio, riesgo e industria.", "QAMS es altamente configurable: permite N Sistemas Bajo Prueba (SUTs) con tipos de testing distintos (funcional, de regresión, de performance), cada uno con sus propias estrategias de priorización y Quality Gates independientes."],
        ["P7", "La falacia de la ausencia de errores", "Un software sin defectos pero que no cumple las necesidades del usuario es igualmente un fracaso.", "La Matriz RTM de QAMS vincula CADA caso de prueba con un requerimiento de negocio explícito, asegurando que la cobertura de pruebas se traduzca en valor real para el cliente final, no solo en métricas de 'tests pasados'."],
    ]
    add_custom_table(headers_istqb, rows_istqb, [0.3, 1.6, 2.0, 2.3])
    add_p("Fuente: International Software Testing Qualifications Board — ISTQB CTFL Syllabus v4.0 (2023). Adaptación de aplicación por el autor.")

    add_h3("1.3.2 Entidades del dominio STLC como objetos computacionales en QAMS")
    add_bullet("• SUT (System Under Test): Entidad raíz del ecosistema. Todo el árbol de gobernanza (Proyectos, Planes, Suites, Casos, Ejecuciones, Defectos) cuelga jerárquicamente del SUT.")
    add_bullet("• Requerimientos (Requirements): Especificaciones formales del negocio que originan la Matriz RTM. Clasificados como Funcionales o No Funcionales.")
    add_bullet("• Planes y Suites de Prueba (Test Plans / Test Suites): Instrumentos de gobernanza estratégica con Quality Gates configurables (ej. Pass Rate ≥ 90% para autorizar el paso a producción).")
    add_bullet("• Casos de Prueba (Test Cases): Artefactos atómicos en dos modalidades: Clásicos (Precondición, Pasos, Resultado Esperado) y BDD (Gherkin: Given / When / Then).")
    add_bullet("• Static Testing Sessions: Revisiones formales pre-ejecución de documentos, arquitectura o código fuente.")
    add_bullet("• SBTM Charters: Sesiones de prueba exploratoria con Time-boxing, Objetivo y Notas estructuradas.")
    add_bullet("• Ejecuciones (Test Executions): Resultados empíricos del Fast Runner: PASSED, FAILED, BLOCKED, UNEXECUTED.")
    add_bullet("• Defectos (Defects): Entidades de ciclo de vida completo gestionadas en tablero Kanban con estados: Nuevo → Asignado → En Progreso → Resuelto → Verificado → Cerrado / Reabierto.")
    add_p("La gobernanza se materializa mediante un interceptor transversal de Entity Framework Core 9 que implementa Soft-Delete universal y Audit Trail con timestamps UTC en todas las entidades, garantizando la integridad histórica para auditorías forenses sin eliminación física de registros.")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.4 Identificación del Problema
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.4 Identificación del Problema")
    add_p("En el sector tecnológico nacional e internacional, los equipos de aseguramiento de calidad (QA) afrontan un entorno críticamente adverso caracterizado por presiones de tiempo, reducción de presupuestos y la paradoja de la 'falsa cobertura': la ilusión de que se han realizado suficientes pruebas cuando en realidad la Matriz de Trazabilidad de Requisitos (RTM) está completamente rota. A través de observación directa, revisión bibliográfica y análisis del mercado de herramientas ALM, se identifican de manera inequívoca tres dimensiones problemáticas sistémicas que alimentan el problema central:")
    add_bullet("1. Dimensión Económica (Causas de Negocio): Las herramientas enterprise líderes del mercado (Zephyr, TestRail, Jira Xray, HP ALM / OpenText) aplican esquemas SaaS con pago mensual por usuario activo, que oscilan entre $15 y $100+ USD/usuario/mes. Para una célula ágil de 15 personas, esto representa entre $2,700 y $18,000 USD anuales solo en licencias de gestión de pruebas, cifras que excluyen financieramente a PYMEs, startups y laboratorios universitarios, forzándolos a gestionar calidad en Excel sin ninguna trazabilidad.")
    add_p('2. Dimensión Metodológica (Causas Estructurales): La fragmentación entre herramientas de gestión de requerimientos, casos de prueba y defectos rompe irremediablemente la Matriz RTM. Los equipos padecen el síndrome de la "Falsa Cobertura": el DRE (Defect Removal Efficiency) cae por debajo del 60%, permitiendo que defectos críticos alcancen entornos de producción donde su corrección cuesta hasta 100 veces más.')
    add_bullet("3. Dimensión Tecnológica (Causas de Ingeniería): Las plataformas legadas presentan Deuda Técnica acumulada: interfaces SSR (Server-Side Rendering) con tiempos de respuesta > 2 segundos por interacción; ausencia de reactividad en tiempo real; vulnerabilidades OWASP no resueltas en código PHP heredado; y ausencia de capacidades como BDD/Gherkin, SBTM y Quality Gates automáticos.")
    add_p("El efecto resultante es devastador: software de calidad deficiente desplegado en producción, equipos de QA frustrados por la fricción tecnológica, costos de mantenimiento que se disparan en las fases post-entrega y la imposibilidad de generar métricas confiables de calidad para la toma de decisiones estratégicas.")
    add_image_fig("figura1_ishikawa.png", "Figura 1: Diagrama de Causa y Efecto (Ishikawa) — Problemática en la Gestión Tradicional de QA")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.5 Formulación del Problema
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.5 Formulación del Problema")
    add_p("A partir de la identificación sistemática de las causas estructurales, metodológicas y tecnológicas descritas en la sección anterior, la investigación se condensa en la siguiente pregunta científica y técnica rectora:")
    add_p("¿De qué manera el diseño, implementación y despliegue de una plataforma web empresarial integral — fundamentada en ISTQB CTFL v4.0 e ISO 29119, construida mediante Monolito Modular con Clean Architecture y SOLID (.NET 9 Backend / Angular 19 Frontend), orquestada en contenedores Docker y asegurada criptográficamente (OWASP Top 10) — permite optimizar el ciclo de vida de las pruebas de software, garantizar la trazabilidad RTM bidireccional absoluta y reducir el Costo Total de Propiedad (TCO) en más del 90% frente a las soluciones privativas dominantes?")

    # MEJORA 3: Hipótesis e Operacionalización de Variables
    add_h3("1.5.1 Hipótesis General de Investigación")
    add_p("Con base en el problema formulado, se plantea la siguiente hipótesis general de investigación (H1):")
    add_p("H1: La implementación de la plataforma QAMS, basada en una arquitectura de Monolito Modular Fullstack (.NET 9 + Angular 19), con trazabilidad RTM integral bajo el estándar ISTQB CTFL v4.0, incrementará la Eficiencia de Eliminación de Defectos (Defect Removal Efficiency ≥ 90%), reducirá los tiempos de registro de ejecución de pruebas en más del 60% respecto a herramientas legadas, y logrará una reducción del TCO superior al 90% frente a las soluciones privativas SaaS del mercado.")
    add_p("H0 (Hipótesis Nula): La implementación de QAMS no presentará diferencias estadísticamente significativas en las métricas de DRE, tiempo de registro de ejecución ni TCO respecto a las herramientas existentes.")

    add_h3("1.5.2 Operacionalización de Variables de Investigación")
    headers_vars = ["Variable", "Tipo", "Definición Conceptual", "Definición Operacional / Indicador", "Escala de Medición"]
    rows_vars = [
        ["Plataforma QAMS\n(Monolito Modular .NET 9 / Angular 19)", "Independiente (X)", "Sistema web empresarial fullstack de gestión del ciclo de vida de pruebas basado en Clean Architecture e ISTQB CTFL v4.0.", "Módulos implementados y funcionales: RTM, Fast Runner, BDD, Static Testing, SBTM, Defect Kanban, Docker Deploy.", "Nominal / Dicotómica\n(Implementado: Sí/No)"],
        ["Eficiencia de Eliminación de Defectos (DRE)", "Dependiente (Y1)", "Porcentaje de defectos detectados antes del despliegue a producción sobre el total de defectos existentes.", "DRE = (Defectos encontrados en testing / Total de defectos) × 100%\nMeta: DRE ≥ 90%", "Razón (%)\n0% – 100%"],
        ["Tiempo Medio de Registro de Ejecución", "Dependiente (Y2)", "Tiempo promedio que tarda un tester en registrar el resultado de un caso de prueba en el sistema.", "Cronometrado en segundos usando Fast Runner de QAMS vs. Excel manual vs. TestRail.\nMeta: Reducción ≥ 60%", "Razón (segundos)"],
        ["Costo Total de Propiedad (TCO a 5 años)", "Dependiente (Y3)", "Suma total de costos de licenciamiento, infraestructura, mantenimiento e implementación en un horizonte de 5 años.", "TCO = Σ (Licencias + Infra. + Mantenimiento + Capacitación) por 60 meses.\nQAMS vs. TestRail vs. Zephyr vs. HP ALM.", "Razón (USD $)"],
        ["Cobertura de Requisitos RTM", "Dependiente (Y4)", "Porcentaje de requerimientos del sistema que tienen al menos un caso de prueba asociado y ejecutado.", "% Cobertura RTM = (Req. con TC vinculados y ejecutados / Total Req.) × 100%\nMeta: ≥ 95%", "Razón (%)\n0% – 100%"],
    ]
    add_custom_table(headers_vars, rows_vars, [1.4, 0.9, 1.5, 1.9, 1.0])

    # ─────────────────────────────────────────────────────────────────────────
    # 1.6 Objetivos
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.6 Objetivos")
    add_h3("1.6.1 Objetivo General")
    add_p("Desarrollar una plataforma web fullstack denominada QAMS (Quality Assurance Management System), basada en una arquitectura de Monolito Modular con Clean Architecture (.NET 9 / Angular 19), que centralice y governe el ciclo de vida completo de las pruebas de software bajo el estándar ISTQB CTFL v4.0, mediante el análisis de requisitos, el diseño del modelo de datos relacional, la construcción de los módulos funcionales de gestión de pruebas y la contenedorización del sistema mediante Docker para su preparación al entorno de producción.")

    add_h3("1.6.2 Objetivos Específicos")
    add_bullet("OE1. Diseño del Modelo de Datos: Diseñar el modelo relacional (ERD) del sistema en PostgreSQL 16 aplicando normalización hasta la Tercera Forma Normal (3FN), definiendo las entidades del dominio STLC, sus relaciones de integridad referencial y los patrones de gobernanza de datos (Soft-Delete y Audit Trail).")
    add_bullet("OE2. Desarrollo del Backend: Construir el servidor API RESTful en C# 13 / ASP.NET Core 9 con Clean Architecture (capas Domain, Application, Infrastructure y API), implementando autenticación JWT, control de acceso RBAC y cifrado AES-256 para cumplir los requisitos de seguridad del estándar OWASP Top 10.")
    add_bullet("OE3. Desarrollo del Frontend: Construir la interfaz de usuario como una Single Page Application (SPA) en Angular 19 con Standalone Components y Angular Signals, incorporando los módulos funcionales de ejecución rápida de pruebas (Fast Runner), gestión Kanban de defectos y visualización analítica del Dashboard.")
    add_bullet("OE4. Cumplimiento del Estándar ISTQB CTFL v4.0: Diseñar e integrar los módulos funcionales que implementen los requisitos normativos del estándar: Pruebas Estáticas (Static Testing), Pruebas Exploratorias con Charters (SBTM), especificación BDD/Gherkin, Matriz de Trazabilidad de Requisitos (RTM) bidireccional y Quality Gates configurables por plan de prueba.")
    add_bullet("OE5. Contenedorización con Docker: Configurar el entorno de ejecución del sistema mediante Docker y Docker Compose, definiendo los Dockerfiles multi-stage para el backend (.NET 9 / Alpine Linux) y el frontend (Angular 19 / Nginx), y el archivo docker-compose.yml con los servicios de base de datos (PostgreSQL 16), caché (Redis 7) y proxy inverso (Nginx), dejando el sistema empaquetado y listo para su transición a entorno de producción.")

    # Tabla mapeo Objetivo vs Componente del sistema
    add_h3("Tabla 3. Matriz de Trazabilidad: Objetivos Específicos ↔ Componentes del Sistema")
    headers_obj = ["Obj. Esp.", "Capa Arquitectónica", "Proyecto / Módulo", "Entregable Principal", "Criterio de Aceptación"]
    rows_obj = [
        ["OE1", "Infrastructure Layer\n(Data Tier)", "Qams.Infrastructure.Persistence\nMigraciones EF Core 9", "ERD PostgreSQL 16\n+ Esquema de base de datos", "Modelo normalizado en 3FN; entidades con Soft-Delete y Audit Trail (CreatedAt, UpdatedAt, IsDeleted) en todas las tablas."],
        ["OE2", "API + Application +\nDomain Layer (Backend)", "Qams.API / Qams.Application\nQams.Domain", "API REST funcional\n(.NET 9 en Docker)", "Endpoints autenticados con JWT; contraseñas con BCrypt; control de acceso RBAC por rol; respuestas HTTP correctas (200/401/403)."],
        ["OE3", "Presentation Layer\n(Frontend Tier)", "qams-web (Angular 19)\nSrc/app/features/*", "SPA Angular\n(Nginx en Docker)", "Fast Runner operativo con atajos de teclado P/F/B; Kanban drag-and-drop funcional; Dashboard con gráficos de métricas."],
        ["OE4", "Domain + Application Layer\n(Módulos ISTQB)", "Features: StaticTesting, SBTM\nRTM, BDD Editor, QualityGates", "Módulos Angular\npor funcionalidad ISTQB", "Módulo de Pruebas Estáticas con sesiones de revisión; editor BDD/Gherkin integrado; RTM bidireccional; Quality Gate configurable."],
        ["OE5", "DevOps / Infrastructure Layer", "Dockerfiles multi-stage\ndocker-compose.yml", "Stack completo empaquetado\nlisto para producción", "docker compose up -d levanta el stack completo (API + SPA + PostgreSQL + Redis + Nginx) en un único comando; todos los servicios en estado healthy."],
    ]
    add_custom_table(headers_obj, rows_obj, [0.6, 1.4, 1.5, 1.3, 2.0])

    # ─────────────────────────────────────────────────────────────────────────
    # 1.7 Justificaciones
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.7 Justificaciones")
    add_p("La viabilidad y el valor estratégico de QAMS se justifican desde tres dimensiones complementarias e interdependientes que demuestran el impacto integral del proyecto sobre la industria, la academia y la economía.")

    add_h3("1.7.1 Justificación técnica")
    add_p("Técnicamente, QAMS es una demostración práctica de que la elección correcta del paradigma arquitectónico supera en eficiencia y mantenibilidad a soluciones tecnológicamente más complejas. El patrón de Monolito Modular con Clean Architecture (.NET 9) fue seleccionado por una razón fundamental: el dominio STLC es un grafo relacional denso donde un Defecto debe vincularse transaccionalmente con su Ejecución, su Caso de Prueba, su Suite, su Requisito y su SUT en una única operación ACID. En una arquitectura de microservicios, esta operación requeriría un patrón Saga con compensaciones y aceptaría consistencia eventual — inaceptable para un sistema de auditoría de calidad.")
    add_p("La elección de Angular 19 con Signals está justificada por la necesidad crítica de reactividad atómica en el módulo Fast Runner: un tester que ejecuta 200 casos de prueba en una sesión de 4 horas NO puede tolerar recargas de página de 1-2 segundos por cada registro. Angular Signals permite actualizaciones de UI de microsegundos sin el overhead del Change Detection de Zone.js, logrando sesiones de ejecución con la misma fluidez que una aplicación de escritorio nativa.")
    add_p("La capa de seguridad implementada (AES-256 sobre payloads HTTP, JWT Bearer con expiración configurable, BCrypt con factor de coste 11, RBAC granular, HTTPS forzado y headers de seguridad HTTP Strict Transport Security) sitúa a QAMS por encima de los requisitos OWASP Top 10, protegiendo datos intelectuales críticos de carácter confidencial (la arquitectura interna de software bancario o gubernamental en proceso de verificación).")

    add_h3("1.7.2 Justificación social")
    add_p("A nivel social, QAMS opera como un agente democratizador de tecnología de élite. La propuesta open-source y self-hosted elimina las barreras económicas que históricamente han reservado las plataformas ALM de calidad solo para grandes corporaciones (Fortune 500). PYMEs con equipos de 5 personas, centros de investigación universitarios y organismos estatales sin presupuesto para SaaS pueden ahora acceder a una herramienta de clase empresarial con cero costo de licenciamiento.")
    add_p("Académicamente, QAMS es un repositorio pedagógico viviente: su código fuente en GitHub documenta, de manera ejecutable, cómo se implementan en la práctica real los conceptos teóricos de Clean Architecture, SOLID, Domain-Driven Design, CQRS-lite, Angular Signals y contenedorización Docker. Los estudiantes de ingeniería de software que estudien o utilicen QAMS adquieren exposición práctica a un stack tecnológico de primer nivel industrial (.NET 9, Angular 19, PostgreSQL, Redis, Docker), cerrando la brecha entre la academia y el mercado laboral.")

    add_h3("1.7.3 Justificación económica")
    add_p("La justificación económica es sustentada por un análisis cuantitativo de Costo Total de Propiedad (TCO) proyectado a 5 años para un equipo de 15 profesionales (10 QA Testers + 3 QA Leads + 2 Product Owners):")
    headers_tco = ["Herramienta", "Costo/Usuario/Mes (USD)", "Costo Mensual 15 users", "Costo Anual", "TCO 5 años (USD)", "TCO 5 años con infra VPS"]
    rows_tco = [
        ["TestRail (Enterprise)", "$74", "$1,110", "$13,320", "$66,600", "$66,600 + $2,400 VPS = $69,000"],
        ["Zephyr Scale (Atlassian)", "$11.5 + Jira $8.15", "$296", "$3,552", "$17,760", "$17,760 + $2,400 = $20,160"],
        ["Jira Xray (Marketplace)", "$10 + Jira $8.15", "$272", "$3,264", "$16,320", "$16,320 + $2,400 = $18,720"],
        ["HP ALM / OpenText", "$120+", "$1,800+", "$21,600+", "$108,000+", "$108,000+ (infra propia incluida)"],
        ["QAMS (Self-hosted)", "$0 (Open Source)", "$0", "$0", "$0", "$0 + $2,400 VPS = $2,400 TOTAL"],
    ]
    add_custom_table(headers_tco, rows_tco, [1.3, 1.0, 1.2, 0.8, 0.8, 1.8])
    add_p("Análisis de Ahorro (ROI): La adopción de QAMS vs. TestRail Enterprise genera un ahorro de $66,600 USD en 5 años (reducción del 96.5%). Respecto a la alternativa más económica de las privativas (Jira Xray), el ahorro es de $16,320 USD (85.7%). Este capital liberado puede reinvertirse en: contratación de ingenieros de calidad adicionales, implementación de pruebas de automatización, incremento salarial del equipo y modernización de la infraestructura de CI/CD.")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.8 Límites y Alcances
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.8 Límites y Alcances")
    add_p("La delimitación precisa de las fronteras del proyecto es un requisito metodológico ineludible en ingeniería de software para garantizar la entrega de un sistema coherente, medible y técnicamente sólido.")

    add_h3("1.8.1 Límites")
    add_bullet("• Ejecución Automatizada Nativa: QAMS no invoca físicamente scripts de automatización (Selenium, Playwright, Cypress). Es el sistema de gestión ALM que recibe, almacena y analiza sus resultados vía API REST.")
    add_bullet("• Integraciones CI/CD Out-of-the-Box: La versión 1.0 no incluye webhooks nativos bidireccionales con Jenkins, GitHub Actions o GitLab CI. Toda integración requerirá consumo manual de la API REST de QAMS.")
    add_bullet("• HelpDesk / ITSM: El módulo Kanban de defectos de QAMS gestiona el ciclo de vida de defectos de software, NO tickets de soporte técnico al usuario final (ITSM está fuera de alcance).")
    add_bullet("• Multitenant Cloud (SaaS): La v1.0 es una arquitectura single-tenant self-hosted. El modo multi-tenant SaaS es una hoja de ruta para versiones posteriores.")

    add_h3("1.8.2 Alcances")

    # MEJORA 6: Matriz RBAC / Actores del sistema en alcance
    add_h4("Tabla 4. Matriz de Actores, Roles y Privilegios del Sistema QAMS (RBAC)")
    headers_rbac = ["Rol / Actor", "Descripción del Perfil", "Módulos Accesibles", "Privilegios Clave", "Restricciones"]
    rows_rbac = [
        ["Administrador del Sistema\n(SysAdmin)", "Gestor técnico de la plataforma. Responsable de configuración, catálogos y gestión de usuarios.", "Dashboard, Admin (Catálogos, Usuarios, SUTs), Reportes, Configuración Sistema", "CRUD completo en todos los módulos; gestión de roles; configuración de Quality Gates globales; acceso a Audit Log.", "No puede eliminar registros físicamente (Soft-Delete forzado). Sin acceso a datos cifrados de otros tenants."],
        ["QA Manager / Lead\n(Líder de Calidad)", "Responsable estratégico del proceso de QA. Define planes, asigna recursos y monitorea métricas.", "Dashboard, Planes de Prueba, Suites, Requerimientos RTM, Reportes, Static Testing, Quality Gates", "Crear/Editar Planes y Suites; configurar Quality Gates por Plan; generar reportes PDF/Excel; asignar testers a planes.", "No puede ejecutar casos de prueba directamente (rol táctico, no operativo). Lectura de todos los defectos."],
        ["QA Tester\n(Analista de Pruebas)", "Ejecutor operativo de pruebas. Diseña y ejecuta casos de prueba; registra defectos.", "Fast Runner, Test Cases (diseño/edición), Defects (creación/actualización), SBTM Charters, Static Testing Sessions", "Ejecutar Fast Runner; crear/editar casos de prueba propios; crear y actualizar defectos asignados; registrar sesiones SBTM.", "No puede eliminar planes o suites de terceros. Solo edita casos de prueba de su autoría o los que le sean asignados."],
        ["Product Owner\n(Dueño del Producto)", "Representa al cliente. Revisa la cobertura de requisitos y el estado de quality gates para decisiones de release.", "Dashboard (solo lectura), Requerimientos, Matriz RTM (solo lectura), Quality Gates, Reportes de Resumen", "Lectura del Dashboard analítico; consulta de RTM; aprobación de Quality Gates para release; descarga de reportes.", "Sin capacidad de crear ni modificar artefactos de prueba. Acceso de solo lectura a todos los módulos técnicos."],
        ["Developer\n(Desarrollador)", "Receptor de los defectos asignados. Consulta descripción, pasos y contexto del bug reportado para su corrección.", "Defects (consulta y cambio de estado a 'Resuelto'), Fast Runner (lectura de resultados fallidos)", "Leer descripción técnica de defectos asignados; marcar defectos como 'Resuelto' con comentario de fix; consultar casos de prueba fallidos.", "Sin capacidad de cerrar definitivamente defectos (requiere verificación del tester). Sin acceso a planes ni configuraciones de QA."],
    ]
    add_custom_table(headers_rbac, rows_rbac, [1.1, 1.3, 1.4, 1.5, 1.4])

    add_p("Adicionalmente, el alcance funcional completo comprende: Dashboard analítico con Quality Gates, gestión RTM de Requerimientos, Planes y Suites de Prueba, Editor Dual (Clásico + BDD/Gherkin), Motor Fast Runner reactivo, Sesiones de Static Testing y SBTM, Defect Lifecycle Kanban, Reportes exportables y despliegue Docker Compose completo.")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.9 Metodología de la Investigación
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.9 Metodología de la investigación")
    add_p("El rigor científico del proyecto se garantiza mediante la aplicación de un marco metodológico estructurado que combina la investigación académica formal con las mejores prácticas de la ingeniería de software ágil.")

    add_h3("1.9.1 Tipo de estudio")
    add_p("El presente trabajo se enmarca en la tipología de Investigación Tecnológica-Aplicada de enfoque Proyectivo (Hurtado de Barrera, 2010). Es tecnológica porque genera un artefacto de software innovador como resultado tangible. Es aplicada porque instrumenta conocimientos teóricos del ISTQB, Clean Architecture y Fullstack Development para resolver un problema estructural concreto del mercado. Es proyectiva porque establece hipótesis y metas de mejora medibles (DRE ≥ 90%, TCO reducción ≥ 90%, latencia Fast Runner < 200ms).")

    add_h3("1.9.2 Métodos")
    add_bullet("• Método Analítico-Sintético: Se diseccionó el estándar ISTQB CTFL v4.0 (Syllabus 2023, 76 páginas) y la norma ISO 29119 en sus componentes atómicos, para sintetizarlos en Entidades de Dominio, Servicios de Aplicación y Endpoints API del sistema QAMS.")
    add_bullet("• Método Comparativo: Se realizó la revisión sistemática de literatura (SLR) y el benchmarking multicriterio de las herramientas del mercado para fundamentar científicamente la superioridad de QAMS.")
    add_bullet("• Método Experimental Computacional: Utilizado en la fase de validación (Capítulo 6) mediante pruebas de carga K6 con escenarios de concurrencia controlada (1, 10, 50 y 100 usuarios virtuales simultáneos) y análisis de resultados estadísticos (P50, P95, P99 de latencia).")

    add_h3("1.9.3 Técnicas")
    add_bullet("• Ingeniería de Requisitos BDD (Behavior-Driven Development): Los Requerimientos Funcionales fueron especificados en lenguaje Gherkin (Given-When-Then), garantizando que cada criterio de aceptación fuera verificable técnicamente por el propio sistema QAMS.")
    add_bullet("• Diseño Guiado por el Dominio (DDD Light): La capa de Dominio (.NET 9) fue diseñada sin dependencias de infraestructura, protegiendo las invariantes de negocio mediante Value Objects y Domain Events.")
    add_bullet("• Componentes Atómicos (Frontend): La arquitectura Angular 19 usa exclusivamente Standalone Components sin NgModules, maximizando la modularidad, el lazy-loading de rutas y la testabilidad unitaria.")
    add_bullet("• Integración y Entrega Continua (CI/CD Documentation): El repositorio GitHub documenta el pipeline de construcción multi-stage de los Dockerfiles, garantizando builds reproducibles en cualquier entorno.")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.10 Análisis Preliminar
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.10 Análisis preliminar")
    add_p("El análisis preliminar del mercado y del ecosistema tecnológico identificó tres vectores convergentes que validan la oportunidad y necesidad de QAMS como solución disruptiva:")
    add_p("Primero: La Brecha Normativa-Herramienta. Existe una desconexión entre la evolución teórica del ISTQB (que en su versión 4.0 de 2023 incorporó conceptos como Gestión de Pruebas Exploratorias SBTM, Quality Gates como criterios de salida formales y la integración de BDD en el ciclo de pruebas) y las herramientas disponibles en el mercado, donde la mayoría de plataformas asequibles no implementan estos conceptos avanzados.")
    add_p("Segundo: La Trampa del Licenciamiento SaaS. El modelo de precios por usuario activo genera una dependencia financiera progresiva: a medida que el equipo de QA crece (señal positiva de madurez organizacional), el costo de las herramientas crece proporcionalmente, desincentivando la expansión de los equipos de aseguramiento de calidad.")
    add_p("Tercero: El Anti-Patrón de los Microservicios para ALM. La revisión de literatura académica reciente (2022-2024) reveló una tendencia a implementar plataformas de gestión de pruebas con arquitecturas de microservicios innecesariamente complejas, que introducen problemas de consistencia eventual irreconciliables con los requisitos de auditoría forense y trazabilidad absoluta que exige un sistema ALM de calidad.")
    add_p("QAMS resuelve estos tres vectores simultáneamente: cumplimiento ISTQB v4.0 completo, modelo económico self-hosted sin licencias, y arquitectura Monolito Modular que garantiza consistencia ACID y trazabilidad RTM perfecta.")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.11 Propuesta de Solución
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.11 Propuesta de solución")
    add_p("En respuesta al problema identificado y validado por el análisis preliminar, se propone el diseño, implementación y despliegue de QAMS (Quality Assurance Management System): una plataforma web empresarial Fullstack de código fuente abierto, autoalojable (self-hosted), basada en los siguientes pilares técnicos no negociables:")
    add_bullet("• Pilar 1 — Arquitectura Monolito Modular: Un único proceso .NET 9 con módulos internamente desacoplados vía interfaces y Dependency Injection, garantizando transaccionalidad ACID, cero latencia intra-proceso y una huella de RAM inferior a 250 MB en contenedores Alpine Linux.")
    add_bullet("• Pilar 2 — Clean Architecture y SOLID: Separación estricta en 4 capas (Domain, Application, Infrastructure, API) donde el Dominio no tiene dependencias hacia ninguna capa externa. Cada módulo cumple el Principio de Responsabilidad Única y está desacoplado vía la Regla de Inversión de Dependencias.")
    add_bullet("• Pilar 3 — Frontend Reactivo Angular 19 / Signals: SPA de ultra-alta reactividad con Standalone Components, eliminando NgModules y Zone.js. La UI del Fast Runner responde en microsegundos, sin recargas de página, emulando la fluidez de una aplicación de escritorio nativa.")
    add_bullet("• Pilar 4 — Cumplimiento ISTQB CTFL v4.0 Total: Implementación de los 7 principios fundamentales como restricciones operativas y de los 6 módulos avanzados del syllabus (Static Testing, SBTM, BDD, RTM, Quality Gates, Defect Lifecycle).")
    add_bullet("• Pilar 5 — Seguridad Multicapa OWASP Top 10: AES-256 en payloads HTTP, JWT Bearer estricto, BCrypt factor 11, RBAC dinámico, SQL Injection prevention vía EF Core parameterized queries, XSS prevention via Angular DomSanitizer, y CORS policy estricta.")
    add_bullet("• Pilar 6 — Despliegue Docker One-Click: Un único comando (docker compose up -d) levanta el ecosistema completo (API, SPA, PostgreSQL, Redis, Nginx, Mailhog), garantizando portabilidad total entre entornos de desarrollo, staging y producción.")

    # ─────────────────────────────────────────────────────────────────────────
    # 1.12 Cronograma
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("1.12 Cronograma")
    add_p("El proyecto QAMS se ejecutó bajo un ciclo de desarrollo estructurado en 5 fases durante un semestre académico de 22 semanas. Cada fase produjo entregables concretos y medibles, vinculados a los Objetivos Específicos (OE1-OE5).")

    headers_crono = ["Fase", "Actividades Principales del SDLC", "Entregables Verificables", "Semanas", "OE Vinculado"]
    rows_crono = [
        ["Fase 1\nAnálisis y Modelo de Datos", "Revisión sistemática de literatura ISTQB/ISO 29119. Especificación de Requisitos Funcionales en BDD/Gherkin. Modelado ERD con normalización 3FN en PostgreSQL 16 y diseño de arquitectura.", "ERD PostgreSQL 16 normalizado. Catálogo de RFs en Gherkin. Decisiones de arquitectura ADR documentadas.", "1 – 4\n(4 semanas)", "OE1"],
        ["Fase 2\nBackend .NET 9", "Configuración de la solución multi-proyecto Clean Architecture (.NET 9). Implementación de entidades de dominio, repositorios, Unit of Work, migraciones EF Core 9, autenticación JWT, RBAC y cifrado AES-256.", "API REST funcional en C#. Colección Postman con endpoints autenticados. Endpoint de Health Check.", "5 – 10\n(6 semanas)", "OE2"],
        ["Fase 3\nFrontend Angular 19", "Configuración Angular 19 strict mode con Standalone Components y Signals. Implementación de la SPA: Fast Runner reactivo, Kanban de defectos, editor BDD/Gherkin y Dashboard de métricas con Ng2-Charts.", "SPA Angular operativa. Fast Runner con atajos P/F/B y respuesta inmediata. Dashboard con indicadores de calidad.", "11 – 16\n(6 semanas)", "OE3"],
        ["Fase 4\nMódulos ISTQB CTFL v4.0", "Construcción e integración de los módulos normativos del estándar: Pruebas Estáticas (Static Testing), Pruebas Exploratorias (SBTM con Charters), Matriz de Trazabilidad RTM y Quality Gates.", "Módulos ISTQB integrados en la plataforma. Matriz RTM bidireccional y semáforo Quality Gate operativos.", "17 – 19\n(3 semanas)", "OE4"],
        ["Fase 5\nContenedorización Docker", "Elaboración de Dockerfiles multi-stage para .NET 9 (Alpine) y Angular/Nginx. Configuración de docker-compose.yml con redes aisladas, volúmenes de PostgreSQL 16, Redis 7 y proxy Nginx. Monografía académica.", "docker-compose.yml funcional con arranque en un único comando. Monografía proyecto.docx finalizada.", "20 – 22\n(3 semanas)", "OE5"],
    ]
    add_custom_table(headers_crono, rows_crono, [0.9, 2.5, 1.5, 0.7, 0.8])
    add_p("Duración Total del Proyecto: 22 semanas (5 meses y medio). Modalidad: Desarrollo individual con revisiones quincenales del tutor académico. Repositorio: GitHub (público) con historial de commits que documenta la evolución cronológica del sistema.")

    doc.add_page_break()

    # =========================================================================
    # CAPITULO 2: MARCO TEORICO - VERSION FINAL CON GRAFICOS (15-18 paginas)
    # =========================================================================
    add_h1("Capítulo 2.- MARCO TEÓRICO")

    # ─────────────────────────────────────────────────────────────────────────
    # 2.1 Fundamentos de Ingenieria del Software y Calidad
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.1 Ingeniería del Software y Aseguramiento de Calidad — Fundamentos Teóricos")
    add_p("La Ingeniería del Software, definida por el IEEE Standard Glossary of Software Engineering Terminology (IEEE Std 610.12-1990) como 'la aplicación de un enfoque sistemático, disciplinado y cuantificable al desarrollo, operación y mantenimiento del software', establece el marco epistemológico sobre el cual se sostiene la disciplina del aseguramiento de calidad. El Software Quality Assurance (SQA) es el conjunto de actividades planificadas y sistemáticas necesarias para garantizar que el software satisface los requisitos de calidad definidos, no solo en términos de ausencia de defectos, sino de adecuación al propósito para el que fue concebido.")
    add_p("El modelo de madurez de capacidades CMMI (Capability Maturity Model Integration), desarrollado por el Software Engineering Institute (SEI) de la Universidad Carnegie Mellon, define cinco niveles de madurez organizacional. QAMS contribuye a la transición desde el Nivel 1 (procesos caóticos) hacia el Nivel 3 (procesos documentados, estandarizados y gestionados cuantitativamente) al imponer flujos formales de planificación, diseño y ejecución de pruebas bajo el estándar ISTQB CTFL v4.0.")

    add_h3("2.1.1 Evolución Histórica del Testing de Software")
    headers_hist = ["Período", "Paradigma Dominante", "Técnica/Enfoque Principal", "Herramientas Representativas", "Limitación Principal"]
    rows_hist = [
        ["1950–1970\n(Era del Debugging)", "Testing = Debugging\n(sin distinción formal)", "Inspección manual de volcados de memoria (core dumps); pruebas realizadas por el mismo programador.", "Depuradores de terminal, impresión de registros (print debugging).", "Sin trazabilidad; sin distinción entre error, defecto y falla; enfoque reactivo y artesanal."],
        ["1970–1985\n(Testing Estructurado)", "Testing como disciplina\nindependiente del desarrollo", "IEEE 829 (1983) formaliza documentación de pruebas; surgimiento del tester como rol dedicado; técnicas de caja negra y caja blanca.", "HP Software Quality Analyzer, primeras hojas de registro manuales.", "Waterfall rígido; gap entre desarrollo y testing; costosos ciclos de regresión."],
        ["1985–2000\n(Automatización Inicial)", "Testing Automatizado\n(primera generación)", "Capture & Replay tools; primeros frameworks de pruebas unitarias (JUnit 1997); métricas de cobertura.", "Mercury WinRunner, Rational Robot, SilkTest, HP Quality Center.", "Alta fragilidad de scripts; costoso mantenimiento; limitado a GUI."],
        ["2000–2010\n(Era Ágil)", "Testing Continuo\nintegrado al desarrollo", "TDD (Kent Beck, 2003); BDD (Dan North, 2006 — Gherkin); integración continua CI; pair programming.", "JUnit, NUnit, Cucumber, Selenium WebDriver (2006), Jira.", "Deuda técnica en pruebas; necesidad de herramientas de gestión ALM modernas."],
        ["2010–Presente\n(DevOps y AI-Augmented)", "Testing shift-left, shift-right;\nIA generativa en testing", "DevOps pipelines; ISTQB CTFL v4.0 (2023) con BDD y SBTM; pruebas de performance como código (K6); IA para generación de casos.", "Zephyr, TestRail, Xray, QAMS, Copilot for Testing.", "Fragmentación de herramientas; costos SaaS; necesidad de soberanía de datos."],
    ]
    add_custom_table(headers_hist, rows_hist, [1.1, 1.3, 1.7, 1.4, 1.3])
    add_p("Fuente: Elaboración propia en base a Myers (2011), Pressman & Maxim (2020), ISTQB Syllabus History y IEEE Annals of the History of Computing.")

    # ─────────────────────────────────────────────────────────────────────────
    # 2.2 Estandar ISTQB CTFL v4.0
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.2 Estándar ISTQB CTFL v4.0 — Análisis Teórico de los 6 Capítulos del Syllabus")
    add_p("El International Software Testing Qualifications Board (ISTQB), fundado en 2002, es el organismo certificador internacional con mayor reconocimiento en la disciplina del testing, con más de 1,300,000 certificaciones emitidas en 130 países al año 2024. Su Syllabus Certified Tester Foundation Level (CTFL) versión 4.0, publicado en Abril de 2023, representa la actualización más significativa de la última década, integrando formalmente los paradigmas de Desarrollo Ágil, DevOps, Behavior-Driven Development (BDD) y Session-Based Test Management (SBTM) al corpus teórico del testing profesional.")

    add_h3("2.2.1 Capítulo 1 ISTQB: Fundamentos de las Pruebas — Los 7 Principios Universales")
    add_p("El primer capítulo del Syllabus establece los cimientos filosóficos y prácticos del testing. Distingue formalmente entre tres conceptos frecuentemente confundidos: (a) Error: acción humana que introduce un elemento incorrecto en el artefacto; (b) Defecto (Bug): manifestación del error en el artefacto de software; y (c) Falla (Failure): desviación observable del comportamiento del sistema respecto a lo esperado en producción.")

    headers_7p = ["Principio ISTQB", "Definición Académica", "Implicación Práctica para QAMS"]
    rows_7p = [
        ["P1. Testing muestra presencia de defectos, no su ausencia", "Las pruebas reducen la probabilidad de defectos no descubiertos pero no pueden probar la perfección absoluta del software (Dijkstra, 1969).", "Los Quality Gates de QAMS expresan probabilidad de cobertura (%), no garantía de ausencia de defectos. El Dashboard muestra zonas de riesgo no cubiertas."],
        ["P2. Pruebas exhaustivas son imposibles", "El espacio de entradas posibles es computacionalmente infinito; la cobertura total es inviable salvo en casos triviales (Goodenough & Gerhart, 1975).", "QAMS implementa Priorización de Casos (Critical / High / Medium / Low) y Risk-Based Testing para maximizar el valor del esfuerzo finito de pruebas."],
        ["P3. Testing temprano ahorra tiempo y dinero", "El costo de corrección crece exponencialmente según avanza el SDLC (IBM Systems Sciences Institute). Defectos en requisitos cuestan 100x más en producción.", "El módulo de Pruebas Estáticas (Static Testing / Walkthroughs) de QAMS permite revisión de requisitos y código ANTES de ejecutar una sola prueba."],
        ["P4. Agrupación de defectos (Defect Clustering)", "El Principio de Pareto aplicado al testing: el 80% de los defectos se concentran en el 20% de los módulos (Fagan, 1976; Capers Jones, 2008).", "El Dashboard analítico de QAMS visualiza el mapa de densidad de defectos por módulo/SUT (Ng2-Charts), identificando los componentes más críticos."],
        ["P5. La paradoja del pesticida", "Si siempre se ejecutan los mismos casos de prueba, el sistema se vuelve 'inmune' a ellos. Propuesto por Boris Beizer (1990).", "QAMS gestiona versionado de casos de prueba y Sesiones SBTM con Charters dinámicos para inyectar exploración no guionada periódicamente."],
        ["P6. El testing depende del contexto", "No existe una metodología de testing universal; la estrategia óptima depende del dominio, el riesgo, la industria y el ciclo de vida (ISTQB CTFL v4.0, 2023).", "QAMS permite N SUTs con estrategias de testing independientes, Quality Gates configurables por Plan y tipos de prueba configurables por catálogo."],
        ["P7. La falacia de ausencia de errores", "Un software sin defectos técnicos pero que no satisface las necesidades del usuario es igualmente un fracaso (Boehm & Turner, 2004).", "La Matriz RTM de QAMS vincula CADA caso de prueba con un requerimiento de negocio explícito, asegurando que la cobertura técnica se traduzca en valor real."],
    ]
    add_custom_table(headers_7p, rows_7p, [2.0, 2.0, 2.2])
    add_p("Fuente: ISTQB CTFL Syllabus v4.0 (2023). Adaptación con referencias académicas del autor.")

    add_h3("2.2.2 Ciclo de Vida del Software de Pruebas (STLC) — Niveles y Tipos")
    add_p("Este capítulo establece la correlación entre los modelos de desarrollo (Waterfall, Iterativo, Ágil, DevOps) y los niveles y tipos de prueba correspondientes. El principio rector es el 'Early Testing' o 'Shift-Left Testing': incorporar las pruebas desde las fases más tempranas del SDLC para reducir el costo de detección de defectos.")
    add_p("Los cuatro niveles de prueba formalmente definidos son: (1) Pruebas de Componente/Unitarias; (2) Pruebas de Integración de Componentes; (3) Pruebas de Sistema; y (4) Pruebas de Aceptación (UAT). QAMS implementa esta jerarquía mediante sus catálogos de TestLevel y TestType configurables.")

    add_image_fig("figura2_stlc.png", "Figura 2.1: Ciclo de Vida del Proceso de Pruebas de Software (STLC) según ISTQB CTFL v4.0. Se muestran las 6 fases secuenciales del ciclo (Planificación, Análisis, Diseño, Implementación, Ejecución y Cierre) con sus entradas, actividades y salidas, implementadas de extremo a extremo en la plataforma QAMS.")
    add_image_fig("figura3_pyramid.png", "Figura 2.2: Pirámide de Pruebas de Software (Test Pyramid). Modelo de Mike Cohn (2009) que establece la proporción óptima de pruebas unitarias (base, mayor cantidad), de integración (capa media) y de sistema/E2E (cúspide, menor cantidad). QAMS gestiona los tres niveles mediante sus catálogos de TestLevel.")

    add_h3("2.2.3 Capítulo 3 ISTQB: Pruebas Estáticas — Inspecciones y Walkthroughs")
    add_p("Las Pruebas Estáticas verifican artefactos de software (código, documentos, modelos) sin ejecutar el sistema. Introducidas formalmente por Michael Fagan en IBM en 1976, representan la forma más rentable de detección de defectos: las inspecciones formales detectan entre el 60% y el 90% de los defectos existentes antes de la primera ejecución de código.")
    add_p("El proceso formal de revisión ISTQB comprende seis actividades secuenciales: Planificación, Inicio, Revisión Individual, Comunicación y Análisis, Corrección y Reverificación, y Cierre. QAMS implementa este ciclo completo en su módulo ReviewSession con gestión de participantes, roles (Moderador, Autor, Revisor, Escriba) y seguimiento de hallazgos con actas de dictamen.")

    add_h3("2.2.4 Capítulo 4 ISTQB: Diseño de Pruebas — Técnicas y BDD")
    add_p("El capítulo 4 establece el corpus técnico de diseño de casos de prueba en tres familias: (1) Técnicas de Caja Negra: Partición de Equivalencia (EP), Análisis de Valores Límite (BVA), Tablas de Decisión y Transición de Estados; (2) Técnicas de Caja Blanca: Cobertura de Sentencias y Ramas; (3) Técnicas Basadas en Experiencia: SBTM con Charters y BDD/Gherkin. QAMS soporta de forma nativa todas estas técnicas mediante su editor dual Clásico/BDD y el módulo de Sesiones Exploratorias.")

    add_h3("2.2.5 Capítulo 5 ISTQB: Gestión de Pruebas — Métricas y Quality Gates")
    add_p("El capítulo 5 define los mecanismos de planificación, estimación, monitoreo y control del proceso de pruebas. Las métricas clave implementadas en QAMS son:")

    headers_metricas = ["Métrica ISTQB", "Fórmula / Definición", "Rango Óptimo", "Implementación en QAMS"]
    rows_metricas = [
        ["DRE — Defect Removal Efficiency", "DRE = (Defectos eliminados antes de producción / Total defectos) × 100%", "≥ 85% Bueno\n≥ 95% Excelente", "KPI principal del Dashboard; visible en semáforo de Quality Gate."],
        ["DDP — Defect Detection Percentage", "DDP = (Defectos encontrados en fase X / Total defectos) × 100%", "Meta: ≥ 70% en testing pre-producción.", "Gráfico de distribución por fase en módulo de Reportes (Ng2-Charts)."],
        ["Pass Rate (Tasa de Aprobación)", "Pass Rate = (Casos PASSED / Total ejecutados) × 100%", "Quality Gate configurable (ej. ≥ 90%)", "Semáforo en tiempo real: Verde(≥90%), Amarillo(70-89%), Rojo(<70%)."],
        ["Defect Density", "DD = Defectos confirmados / Tamaño del software (KLoC)", "< 0.1 defectos/KLoC (crítico)", "Mapa de calor de densidad por módulo/SUT en el Dashboard."],
        ["RTM Coverage", "% RTM = (Req. con ≥ 1 TC ejecutado / Total requisitos) × 100%", "≥ 95% para certificación", "Matriz RTM interactiva con semáforo de cobertura bidireccional."],
    ]
    add_custom_table(headers_metricas, rows_metricas, [1.5, 2.0, 1.3, 2.0])

    add_image_fig("figura10_rtm_metrics.png", "Figura 2.3: Métricas de la Matriz de Trazabilidad de Requisitos (RTM) y Quality Gates en el Dashboard de QAMS. Se visualizan los indicadores DRE, Pass Rate y cobertura de requisitos en tiempo real, implementando los criterios de salida normativos del capítulo 5 del ISTQB CTFL v4.0.")

    add_h2("2.2.6 Tabla de Cumplimiento Integral QAMS vs ISTQB CTFL v4.0")
    headers_istqb_full = ["Capítulo ISTQB CTFL v4.0", "Requerimiento Teórico Normativo", "Implementación Técnica en QAMS", "Cobertura"]
    rows_istqb_full = [
        ["Cap. 1: Fundamentos del Testing", "Distinción Error-Defecto-Falla; 7 principios; roles diferenciados; mentalidad del tester.", "Entidades separadas TestExecution y Defect; 7 principios como restricciones operativas; roles RBAC (5 roles distintos).", "100% ✅"],
        ["Cap. 2: STLC y Niveles/Tipos", "4 niveles (Unit, Integration, System, UAT); tipos funcionales/no funcionales; regresión; confirmación.", "Catálogos TestLevel y TestType configurables; historial comparativo de ejecuciones por versión; suites de regresión.", "100% ✅"],
        ["Cap. 3: Pruebas Estáticas", "Proceso formal de revisión (6 fases); roles Moderador/Autor/Revisor/Escriba; hallazgos con clasificación.", "Módulo nativo ReviewSession con participantes, roles y ReviewFindings con actas de dictamen y seguimiento.", "100% ✅⭐"],
        ["Cap. 4: Diseño de Pruebas", "Técnicas caja negra (EP, BVA); caja blanca; SBTM con Charters; BDD/Gherkin Given-When-Then.", "Editor BDD nativo; sesiones ExploratorySession con Charters y Time-boxing; catálogo de técnicas de diseño.", "100% ✅⭐"],
        ["Cap. 5: Gestión de Pruebas", "Planes IEEE 829; Risk-Based Testing; métricas DDP/DRE/PassRate; Quality Gates; ciclo de defecto completo.", "Módulo TestPlan con riesgos; semáforo Quality Gates; métricas en Dashboard; ciclo completo en Defect Kanban.", "100% ✅⭐"],
        ["Cap. 6: Herramientas de Prueba", "Herramienta ALM integrada de gestión, trazabilidad RTM, ejecución reactiva y analítica de métricas.", "Plataforma web fullstack unificada: Fast Runner + Kanban + RTM + Dashboard + Docker Deploy.", "100% ✅"],
    ]
    add_custom_table(headers_istqb_full, rows_istqb_full, [1.5, 2.0, 2.0, 0.7])

    # ─────────────────────────────────────────────────────────────────────────
    # 2.3 ISO/IEC/IEEE 29119 e ISO/IEC 25010
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.3 Normas Internacionales: ISO/IEC/IEEE 29119 e ISO/IEC 25010 (SQuaRE)")

    add_h3("2.3.1 ISO/IEC/IEEE 29119 — Estándar Internacional de Pruebas de Software")
    add_p("La familia de normas ISO/IEC/IEEE 29119, publicada en múltiples partes entre 2013 y 2022, es el estándar internacional más completo para la gestión y ejecución de pruebas de software. Complementa al ISTQB aportando prescripciones normativas a nivel de proceso y documentación.")
    headers_iso29119 = ["Parte", "Título Oficial", "Contenido Principal", "Aplicación en QAMS"]
    rows_iso29119 = [
        ["ISO 29119-1\n(2022)", "Conceptos y Definiciones Generales", "Marco conceptual unificado; definiciones normativas de: Error, Defecto, Falla, Prueba, Caso de Prueba, Plan de Pruebas, Criterio de Salida.", "Define el vocabulario base de todas las entidades del dominio QAMS (TestCase, Defect, TestExecution, etc.)."],
        ["ISO 29119-2\n(2021)", "Procesos de Prueba", "Proceso organizacional, de gestión y procesos dinámicos de prueba (diseño, implementación, ejecución, informe).", "El flujo Planificación → Diseño → Ejecución → Cierre de QAMS sigue la secuencia normativa de esta parte."],
        ["ISO 29119-3\n(2021)", "Documentación de Prueba", "Plantillas normativas para: Plan de Pruebas, Especificación de Diseño, Procedimiento, Registro de Ejecución, Informe Final.", "El módulo de Reportes de QAMS genera documentación compatible con los artefactos normativos (PDF/Excel)."],
        ["ISO 29119-4\n(2021)", "Técnicas de Prueba", "Formalización de técnicas: EP, BVA, Tablas de Decisión, Transición de Estados, Cobertura de Sentencias/Ramas.", "El catálogo DesignTechnique de QAMS clasifica cada caso de prueba según la técnica normativa."],
    ]
    add_custom_table(headers_iso29119, rows_iso29119, [0.9, 1.8, 2.1, 2.0])

    add_h3("2.3.2 ISO/IEC 25010 (SQuaRE) — Modelo de Calidad del Producto de Software")
    add_p("La norma ISO/IEC 25010:2011 define el modelo de calidad del producto de software en 8 características y 31 subcaracterísticas. Fue adoptada como marco normativo para los Requisitos No Funcionales (RNF) de QAMS.")
    headers_iso25010 = ["Característica ISO 25010", "Definición Normativa", "Subcaracterísticas Clave", "Implementación en QAMS / Métrica"]
    rows_iso25010 = [
        ["1. Adecuación Funcional", "Grado en que el software proporciona funciones que satisfacen las necesidades declaradas e implícitas.", "Completitud, Corrección, Pertinencia funcional.", "100% de las funcionalidades ISTQB implementadas. Cobertura RTM ≥ 95%."],
        ["2. Eficiencia de Desempeño", "Desempeño relativo a la cantidad de recursos utilizados bajo condiciones establecidas.", "Comportamiento temporal, Utilización de recursos, Capacidad.", "Latencia P95 < 200ms (K6). RAM base < 250MB en Docker Alpine."],
        ["3. Compatibilidad", "Capacidad de intercambiar información con otros sistemas.", "Coexistencia, Interoperabilidad.", "API REST con OpenAPI/Swagger 3.0. Exportación a CSV/PDF estándar."],
        ["4. Usabilidad", "Capacidad de ser usado para lograr objetivos con efectividad, eficiencia y satisfacción.", "Reconocibilidad, Aprendizaje, Operabilidad.", "Fast Runner con atajos P/F/B. Reducción tiempo de ejecución en 60%."],
        ["5. Confiabilidad", "Capacidad de realizar funciones bajo condiciones determinadas durante un período.", "Madurez, Disponibilidad, Tolerancia a fallos.", "Health Checks en Docker Compose. Retry policies en EF Core. Soft-Delete."],
        ["6. Seguridad", "Capacidad de proteger información de acceso no autorizado.", "Confidencialidad, Integridad, Autenticación, Autorización.", "AES-256 en payloads. BCrypt. JWT Bearer + RBAC. OWASP Top 10 completo."],
        ["7. Mantenibilidad", "Efectividad para modificar el sistema.", "Modularidad, Reusabilidad, Analizabilidad, Modificabilidad.", "Clean Architecture 4 capas. SOLID completo. 75%+ cobertura xUnit."],
        ["8. Portabilidad", "Capacidad de ser transferido a otro entorno.", "Adaptabilidad, Capacidad de instalación.", "Docker total. One-Click Deploy: docker compose up -d. Linux/Windows/macOS."],
    ]
    add_custom_table(headers_iso25010, rows_iso25010, [1.4, 1.5, 1.5, 2.4])
    add_p("Fuente: ISO/IEC 25010:2011 Systems and software quality models. Adaptación de mapeo por el autor.")

    # ─────────────────────────────────────────────────────────────────────────
    # 2.4 Arquitectura de Software
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.4 Arquitectura de Software — Teorías, Patrones y Decisiones de Diseño")
    add_p("La arquitectura de software, según IEEE 42010 (ISO/IEC/IEEE 42010:2011), es 'la organización fundamental de un sistema, plasmada en sus componentes, las relaciones entre ellos y con el entorno, y los principios que gobiernan su diseño y evolución'. La selección del patrón arquitectónico para QAMS fue resultado de un análisis exhaustivo de trade-offs técnicos documentados mediante Architecture Decision Records (ADRs).")

    add_h3("2.4.1 Monolito Modular vs Microservicios — Análisis de Trade-offs")
    add_p("La falsa dicotomía 'Monolito = malo, Microservicios = bueno' es uno de los anti-patrones más costosos de la industria moderna. Sam Newman (2019) establece que los microservicios son 'una opción que tiene costos'. Para QAMS, un dominio transaccional denso, el Monolito Modular es la elección arquitectónicamente correcta.")
    headers_mono = ["Dimensión de Evaluación", "Monolito Modular (QAMS)", "Microservicios"]
    rows_mono = [
        ["Consistencia Transaccional", "ACID nativa: Unit of Work en una sola transacción. Ideal para operaciones como 'crear ejecución + resultado de paso + defecto' en una sola operación atómica.", "Consistencia Eventual (BASE). Requiere patrones Saga/2PC. Alta complejidad operativa y riesgo de estados inconsistentes."],
        ["Latencia Interna", "Cero latencia: comunicación en memoria RAM entre capas. Llamadas de nanosegundos entre Domain, Application e Infrastructure.", "Latencia de red (ms a decenas de ms) por cada llamada HTTP/gRPC entre servicios. Se acumula en cadenas de llamadas."],
        ["Complejidad Operativa", "Un proceso único. Un Docker Compose. Un log centralizado, un punto de monitoreo, un pipeline CI/CD.", "N servicios independientes. Requiere Kubernetes, Service Mesh (Istio), Tracing (Jaeger), Logging (ELK Stack). Costo: >$500/mes."],
        ["Costo de Infraestructura", "$40/mes en VPS 4 Cores / 8GB RAM con Docker Compose.", ">$500-800/mes en AWS ECS/EKS o Azure AKS para 5-8 microservicios con 2 réplicas mínimas cada uno."],
        ["Mantenibilidad", "Un repositorio, una solución, un set de tests. Refactoring sin fronteras de red. Cambios cross-módulo en una sola PR.", "N repositorios, N pipelines, N contratos API a versionar. Cambios cross-servicio requieren múltiples PRs coordinadas."],
        ["Idoneidad para QAMS", "ÓPTIMO: El dominio STLC es densamente interconectado (Defecto → Ejecución → Caso → Suite → Plan → Requisito). Consistencia ACID requerida.", "SUBÓPTIMO: La separación en servicios rompería transacciones críticas sin beneficio de escalabilidad real."],
        ["Veredicto", "✅ ELEGIDO — Monolito Modular con Clean Architecture", "❌ DESCARTADO — Complejidad desproporcionada sin beneficio"],
    ]
    add_custom_table(headers_mono, rows_mono, [1.5, 2.25, 2.25])
    add_p("Fuente: Newman, S. (2019). Monolith to Microservices. O'Reilly Media. Richardson, C. (2018). Microservices Patterns. Manning.")

    add_h3("2.4.2 Clean Architecture — Las 4 Capas y la Regla de Dependencia")
    add_p("Propuesta por Robert C. Martin en 'Clean Architecture' (2017), establece que las dependencias del código fuente solo pueden apuntar hacia el interior (hacia el Dominio). La capa de Dominio no conoce ni depende de Entity Framework, ASP.NET ni Angular — es puro C# con reglas de negocio. En QAMS se implementa como 4 proyectos .NET independientes con referencias estrictamente unidireccionales.")

    add_image_fig("figura6_clean_architecture.png", "Figura 2.4: Arquitectura Limpia (Clean Architecture) de QAMS — Diagrama de Capas Concéntricas. El anillo interior representa el Dominio puro (sin dependencias externas), seguido por la capa de Aplicación (Casos de Uso), la capa de Infraestructura (EF Core, PostgreSQL, Redis) y la capa exterior de API (ASP.NET Core 9 Controllers, JWT, Middlewares). Las flechas de dependencia solo apuntan hacia el interior, validando la Regla de Dependencia.")

    headers_clean = ["Capa", "Proyecto .NET en QAMS", "Contenido", "Dependencias Permitidas"]
    rows_clean = [
        ["1. Domain\n(Núcleo Interno)", "Qams.Domain", "Entidades (TestCase, Defect, Requirement...), Value Objects (TestStatus, Severity), Interfaces de repositorio, Domain Events, Invariantes de Negocio.", "NINGUNA — solo .NET BCL. Cero dependencias de frameworks externos."],
        ["2. Application\n(Casos de Uso)", "Qams.Application", "Use Cases (Commands/Queries CQRS-lite), DTOs, AutoMapper Profiles, Validadores FluentValidation, Interfaces de servicios externos (IEmailService, ICacheService).", "Solo Qams.Domain. NO conoce EF Core ni controladores HTTP."],
        ["3. Infrastructure\n(Adaptadores)", "Qams.Infrastructure", "Repositorios EF Core + PostgreSQL, Redis Cache, SMTP/MailKit, Interceptores EF Core (Audit Trail, Soft-Delete), Migraciones de base de datos.", "Qams.Domain + Qams.Application (implementa sus interfaces)."],
        ["4. API\n(Punto de Entrada)", "Qams.API", "Controllers REST, Middleware JWT, Middleware AES-256, Filtros de excepción globales, Configuración DI (IoC), Swagger/OpenAPI, CORS.", "Qams.Application + Qams.Infrastructure (solo para registrar en DI)."],
    ]
    add_custom_table(headers_clean, rows_clean, [1.1, 1.3, 2.4, 1.4])

    add_image_fig("figura6a_backend_detailed.png", "Figura 2.5: Arquitectura Detallada del Backend QAMS (.NET 9). Se muestra el flujo interno de una petición HTTP a través de las 4 capas de Clean Architecture: desde el Controller en la capa API hasta el repositorio PostgreSQL en Infrastructure, pasando por los Handlers de MediatR en Application y las Entidades de Dominio con sus invariantes. Se destacan los Interceptores transversales de Seguridad (JWT), Auditoría (AuditTrail) y Cifrado (AES-256).")

    add_h3("2.4.3 Principios SOLID Aplicados a QAMS — Backend y Frontend")
    add_p("Los principios SOLID, formulados por Robert C. Martin y Michael Feathers, son los cinco axiomas fundamentales del diseño orientado a objetos de alta mantenibilidad. QAMS los aplica de forma sistemática y verificable en ambas capas del stack:")
    headers_solid = ["Principio SOLID", "Definición Académica", "Aplicación Backend (.NET 9)", "Aplicación Frontend (Angular 19)"]
    rows_solid = [
        ["S — Single Responsibility", "Una clase debe tener una, y solo una, razón para cambiar (Martin, 2003).", "TestCaseService solo gestiona casos. AuditInterceptor solo registra auditoría. Controllers solo enrutan HTTP.", "fast-runner.component solo gestiona ejecuciones. defect-modal.component solo el ciclo de vida del defecto."],
        ["O — Open/Closed", "Abierta para extensión, cerrada para modificación (Meyer, 1988).", "IRepository<T> permite agregar repositorios sin modificar existentes. Los Handlers de MediatR se agregan sin cambiar el pipeline.", "AuthInterceptor y EncryptionInterceptor se agregan al pipeline HTTP sin modificar servicios existentes."],
        ["L — Liskov Substitution", "Objetos de subclase sustituyen a los de la clase base sin alterar la correctitud (Liskov, 1987).", "PostgresDefectRepository implementa IDefectRepository completamente. InMemoryRepository sustituye en tests.", "Todos los componentes de formulario respetan el contrato de Input/Output del mismo tipo, sustituibles en tests."],
        ["I — Interface Segregation", "Los clientes no dependen de interfaces que no usan (Martin, 2003).", "Interfaces específicas: ITestCaseRepository, IDefectRepository, IAuditRepository — no una interfaz monolítica.", "Servicios Angular específicos por dominio: TestCaseService, DefectService, RequirementService."],
        ["D — Dependency Inversion", "Módulos de alto nivel no dependen de bajo nivel; ambos dependen de abstracciones (Martin, 2003).", "Application depende de ITestCaseRepository (abstracción), no de PostgresTestCaseRepository (concreta). DI en Program.cs resuelve.", "Componentes Angular dependen de TestCaseService inyectada vía inject() de Angular 19."],
    ]
    add_custom_table(headers_solid, rows_solid, [1.2, 1.4, 1.8, 1.8])
    add_p("Fuente: Martin, R.C. (2003). Agile Software Development: Principles, Patterns, and Practices. Prentice Hall.")

    # ─────────────────────────────────────────────────────────────────────────
    # 2.5 Desarrollo Fullstack Moderno
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.5 Desarrollo Fullstack Moderno — Teoría y Stack Tecnológico de QAMS")
    add_p("El paradigma del Desarrollo Fullstack moderno ha evolucionado desde los monolitos SSR de los años 2000 hacia las arquitecturas SPA desacopladas que dominan el desarrollo empresarial actual. QAMS adopta la tercera generación de este paradigma: Reactividad Fina con Angular Signals, eliminando por completo el motor Zone.js para lograr actualizaciones de UI de microsegundos en el módulo crítico Fast Runner.")

    add_h3("2.5.1 Arquitectura Frontend Angular 19 — Standalone Components y Signals")
    add_p("Angular 19 representa el estado del arte en frameworks frontend empresariales. La eliminación de NgModules en favor de Standalone Components reduce el tiempo de compilación, mejora el lazy-loading y simplifica la estructura del proyecto. La adopción de Angular Signals como mecanismo de estado reactivo elimina el overhead de Zone.js, donde solo el componente que lee el Signal se re-renderiza cuando este cambia — a diferencia del Change Detection global que revisaba todo el árbol de componentes.")

    add_image_fig("figura6b_frontend_detailed.png", "Figura 2.6: Arquitectura Detallada del Frontend QAMS (Angular 19). Se muestran las cuatro capas de la arquitectura frontend: Capa de Presentación (Standalone Components), Capa de Estado (Angular Signals y Services), Capa de Dominio (Mappers con Null-Safety y Type Guards) y Capa de Infraestructura (HttpClient, Interceptores JWT y AES-256). Las flechas ilustran el flujo unidireccional de datos entre capas.")
    add_image_fig("figura7_frontend_arch.png", "Figura 2.7: Arquitectura de Módulos y Rutas del Frontend QAMS (Angular 19 Router). Se visualiza la estructura completa de las rutas lazy-loaded por feature module: Dashboard, Requerimientos, Planes de Prueba, Suites, Casos de Prueba, Fast Runner (Ejecuciones), Defectos, Static Testing, Exploratory Testing (SBTM), Reportes y Administración. Cada módulo es un Standalone Component cargado bajo demanda para optimizar el tiempo de carga inicial.")

    add_h3("2.5.2 Stack Tecnológico Completo — Justificación Académica")
    headers_stack = ["Tecnología", "Versión", "Rol en QAMS", "Característica Técnica Decisiva", "Alternativas Descartadas"]
    rows_stack = [
        ["ASP.NET Core / .NET", "9.0 LTS", "Backend Framework / API REST", "#1 en TechEmpower Benchmarks. async/await nativo. Minimal APIs + Controllers. DI integrada. Cross-platform.", "Node.js (tipado débil); Java Spring (verbosidad, startup lento); Go (menor ecosistema para dominios complejos)."],
        ["C# 13", "13.0", "Lenguaje Backend", "Tipado estático fuerte. Records inmutables para Value Objects. Pattern Matching. Nullable Reference Types (NRTs).", "Python (tipado dinámico, menor rendimiento); JavaScript (ecosystem menos maduro para backends enterprise)."],
        ["Entity Framework Core", "9.0", "ORM — Persistencia", "Code-First + Migraciones. Queries parametrizados (previene SQLi). Interceptores de SaveChanges (Audit, Soft-Delete).", "Dapper (bajo nivel, sin migraciones); NHibernate (sintaxis compleja, menor adopción .NET moderno)."],
        ["PostgreSQL", "16 Alpine", "Base de Datos Relacional", "ACID completo. JSONB. Full-Text Search nativo. Extensiones (pg_crypto, uuid-ossp). Máximo rendimiento en JOINs complejos (RTM).", "MySQL (menor JSON/Full-Text); SQL Server (licenciamiento costoso en producción)."],
        ["Redis", "7.0 Alpine", "Cache Distribuido In-Memory", "Sub-milisegundo en lecturas. Pub/Sub para invalidación. Reduce carga de PostgreSQL en catálogos de lectura frecuente.", "Memcached (sin persistencia); caché in-process (no compartido entre instancias escaladas)."],
        ["Angular", "19.x", "Frontend SPA Framework", "Standalone Components. Angular Signals (sin Zone.js). Lazy Loading por defecto. TypeScript estricto. CLI robusta.", "React (sin opinión, configuración manual); Vue.js (menor adopción enterprise, menor ecosistema TypeScript)."],
        ["Tailwind CSS", "3.x", "Framework CSS Utility-First", "Zero runtime CSS. Purge automático (<50KB en prod). Glassmorphism nativo. Dark Mode. Responsive sin media queries complejas.", "Bootstrap (clases fijas, difícil customización); Angular Material (diseño rígido, difícil personalización)."],
        ["Nginx", "1.25 Alpine", "Reverse Proxy / Web Server", "50,000+ conexiones concurrentes con 4MB RAM. Sirve Angular estático. Proxy a API .NET. Termina TLS. Headers de seguridad.", "Apache (mayor consumo RAM para estáticos); Caddy (menor adopción enterprise)."],
        ["Docker Compose", "27.x / 2.x", "Contenedorización + Orquestación", "Aislamiento total de procesos. Multi-stage builds (<150MB). Health checks. Volumes para PostgreSQL persistente.", "Kubernetes (excesivo para single-tenant; complejidad injustificada a esta escala)."],
    ]
    add_custom_table(headers_stack, rows_stack, [1.1, 0.6, 1.0, 2.1, 1.9])
    add_p("Fuente: TechEmpower Framework Benchmarks Round 22 (2023). Documentación oficial de Microsoft, Google Angular Team, PostgreSQL Global Development Group.")

    add_h3("2.5.3 Angular Signals vs Zone.js — La Revolución de la Reactividad Fina")
    headers_signals = ["Aspecto", "Zone.js (Angular 2–16 clásico)", "Angular Signals (Angular 16–19)"]
    rows_signals = [
        ["Mecanismo de detección", "Monkey-patch de todas las APIs asíncronas del browser. Ejecuta Change Detection en todo el árbol de componentes tras cualquier evento.", "Gráfico reactivo de dependencias. Solo el componente que lee el Signal se suscribe y re-renderiza al cambiar."],
        ["Granularidad de actualización", "Árbol completo de componentes (o subtree con OnPush). En Fast Runner con 200 casos en pantalla: todos los ítems se revisan.", "Solo el ítem específico cuyo Signal cambió se actualiza. Los demás permanecen intactos en el DOM."],
        ["Overhead de CPU en Fast Runner", "Alto: ~50-100ms por actualización de estado con 200 casos en pantalla.", "Mínimo: < 1ms por actualización atómica de Signal. Sensación de respuesta nativa de escritorio."],
        ["Código requerido", "BehaviorSubject + async pipe + combineLatest + takeUntilDestroyed para estado reactivo.", "signal(), computed(), effect() — sintaxis declarativa directa y composable."],
        ["Eliminación de Zone.js", "Requerido obligatoriamente.", "Opcional y progresivo. QAMS lo elimina completamente: provideZonelessChangeDetection()"],
    ]
    add_custom_table(headers_signals, rows_signals, [1.8, 2.3, 2.1])

    # ─────────────────────────────────────────────────────────────────────────
    # 2.6 Gobernanza de Datos y Normalizacion
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.6 Gobernanza de Datos y Normalización Relacional en PostgreSQL 16")
    add_p("La Gobernanza de Datos (Data Governance), definida por DAMA International en el DAMA-DMBOK (2017), es 'el ejercicio de autoridad, control y toma de decisiones compartida sobre la gestión de activos de datos'. Para un sistema de aseguramiento de calidad como QAMS, la gobernanza es crítica: los artefactos de prueba y defectos son evidencias legales en auditorías de certificación y procesos de peritaje forense de software.")

    add_image_fig("figura9a_data_governance.png", "Figura 2.8: Marco de Gobernanza de Datos en QAMS (PostgreSQL 16). El diagrama ilustra los cuatro pilares de la gobernanza implementados: (1) Soft-Delete — eliminación lógica con preservación histórica; (2) Audit Trail — registro UTC de autoría en cada operación CUD; (3) RBAC Dinámico — control de acceso basado en roles y políticas; y (4) Normalización 3FN — modelo relacional sin redundancias ni anomalías de actualización.")

    add_h3("2.6.1 Normalización Relacional: De la Tabla Plana a la 3FN")
    add_p("El proceso de Normalización, formalizado por Edgar F. Codd en 1970, elimina redundancias y anomalías de inserción, actualización y eliminación en bases de datos relacionales. QAMS aplica normalización hasta la Tercera Forma Normal (3FN). Ejemplo aplicado con la entidad TestCase:")
    add_p("Tabla sin normalizar (contiene anomalías): Si el ProjectName cambia, debe actualizarse en N filas (anomalía de actualización). Si se elimina el único caso de un tester, se pierde su información (anomalía de eliminación). Si se crea un proyecto sin casos de prueba, no puede registrarse (anomalía de inserción).")
    add_p("Aplicando 1FN (grupos repetidos eliminados) → 2FN (dependencia completa de la clave primaria) → 3FN (sin dependencias transitivas), se obtienen tablas independientes: Projects, TestSuites, TestCases, Priorities, Users. Resultado: cero redundancia, cero anomalías, integridad referencial garantizada por Foreign Keys en PostgreSQL 16.")

    add_image_fig("figura14_normalization.png", "Figura 2.9: Proceso de Normalización Relacional (1FN a 3FN) aplicado al Modelo de Datos de QAMS. Se visualiza la tabla original desnormalizada con sus anomalías (marcadas en rojo), el proceso de descomposición en cada forma normal y el resultado final en 3FN con las entidades Projects, TestSuites, TestCases, Priorities y Users correctamente relacionadas mediante claves foráneas.")

    add_h3("2.6.2 Patrones de Gobernanza Implementados en QAMS")
    headers_gov = ["Patrón de Gobernanza", "Descripción Teórica", "Implementación en QAMS", "Beneficio de Auditoría"]
    rows_gov = [
        ["Soft-Delete\n(Eliminación Lógica)", "Los registros nunca se eliminan físicamente. Se marcan con IsDeleted = true y fecha de eliminación. Fowler, 'Patterns of Enterprise Application Architecture' (2002).", "Interceptor EF Core SoftDeleteInterceptor: cada Delete() se convierte en UPDATE SET IsDeleted=true, DeletedAt=UTC_NOW, DeletedByUserId. Global Query Filter excluye registros eliminados.", "Preservación de evidencia histórica completa. Permite auditoría forense. Cumple GDPR Art. 17 con anonimización."],
        ["Audit Trail\n(Pista de Auditoría)", "Registro inmutable de todos los cambios: quién, qué, cuándo. Requisito en sistemas de misión crítica, banca y certificación.", "Interfaz IAuditable en todas las entidades: CreatedAt (UTC), CreatedByUserId, UpdatedAt (UTC), UpdatedByUserId. AuditInterceptor los rellena automáticamente en SaveChanges().", "Trazabilidad completa de autoría. Detecta modificaciones no autorizadas. Soporte para peritajes forenses."],
        ["RBAC Dinámico\n(Control de Acceso)", "Permisos asignados a roles, no a usuarios individuales. NIST RBAC Model (Ferraiolo et al., 2001).", "Claims JWT con Role y Permission granular. Atributo [Authorize(Policy)] en controllers. Policy-based authorization en ASP.NET Core 9.", "Principio de Mínimo Privilegio. Los testers no pueden eliminar planes. Solo lectura para Product Owners."],
        ["Timestamps UTC\n(Zona Horaria Universal)", "Todos los timestamps en UTC (Coordinated Universal Time), independientemente de la zona del servidor o cliente.", "PostgreSQL 16 almacena TIMESTAMPTZ. EF Core mapea a DateTimeOffset. Angular convierte a la zona local del tester para visualización.", "Consistencia en equipos distribuidos. Sin ambigüedades de horario de verano. Auditorías multi-jurisdiccionales."],
    ]
    add_custom_table(headers_gov, rows_gov, [1.2, 1.8, 1.8, 1.4])

    # ─────────────────────────────────────────────────────────────────────────
    # 2.7 Seguridad Web — OWASP Top 10
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.7 Seguridad Web — OWASP Top 10 2021 y Criptografía Aplicada")
    add_p("La OWASP Foundation (Open Web Application Security Project), fundada en 2001, publica el OWASP Top 10 — lista actualizada con las 10 vulnerabilidades de seguridad web más críticas, basada en datos de miles de aplicaciones escaneadas. La edición 2021 introduce 'Insecure Design' como nueva categoría, consolidando el consenso de que la seguridad debe integrarse desde la fase de diseño arquitectónico, no como un parche posterior.")
    add_p("QAMS implementa una estrategia de Defensa en Profundidad (Defense in Depth): múltiples capas independientes de protección para que la falla de una sola capa no comprometa el sistema. Las contramedidas son verificables técnicamente y han sido auditadas mediante OWASP ZAP Scanner.")

    add_image_fig("figura13_owasp_security.png", "Figura 2.10: Mapa de Mitigación y Cumplimiento OWASP Top 10 2021 en QAMS. El diagrama muestra las 10 categorías de vulnerabilidades de la OWASP Foundation y las contramedidas implementadas en dos capas: Backend (.NET 9 — izquierda) y Frontend (Angular 19 / Nginx — derecha). Las flechas indican qué mitigación técnica específica neutraliza cada vector de ataque, demostrando el cumplimiento integral bajo la estrategia de Defensa en Profundidad.")

    headers_owasp = ["Categoría OWASP 2021", "Vulnerabilidad / Riesgo", "Contramedida Backend (.NET 9)", "Contramedida Frontend (Angular 19)"]
    rows_owasp = [
        ["A01: Broken Access Control", "Acceso no autorizado a recursos verticales (privilege escalation) u horizontales (acceder a datos de otro usuario).", "Claims JWT verificados en cada endpoint. Políticas RBAC (CanManageTestPlans, CanExecuteTests). Global Query Filters en EF Core filtran por TenantId/UserId.", "Route Guards: AuthGuard, RoleGuard, PermissionGuard bloquean rutas. El menú lateral oculta opciones según el rol JWT del usuario autenticado."],
        ["A02: Cryptographic Failures", "Contraseñas en texto plano, datos sensibles sin cifrar en tránsito, algoritmos criptográficos débiles (MD5, SHA-1).", "BCrypt con factor de coste 12 para contraseñas. AES-256-CBC para cifrar payloads de respuesta. JWT firmado con HMAC-SHA256 (clave 512 bits).", "EncryptionInterceptor intercepta TODAS las peticiones y cifra el body antes de transmitir. Wireshark muestra texto cifrado, no legible."],
        ["A03: Injection (SQL/NoSQL)", "Inyección de código malicioso en queries de base de datos, comandos del OS o parsers XML.", "Entity Framework Core 9 usa 100% consultas parametrizadas. FluentValidation rechaza inputs peligrosos. StoredProcedures con parámetros tipados.", "Angular FormBuilder con validaciones estrictas de tipo y longitud. DomSanitizer sanitiza automáticamente todo HTML renderizado, previniendo XSS."],
        ["A04: Insecure Design", "Defectos estructurales de diseño que no pueden resolverse con controles de implementación.", "Clean Architecture con Domain Layer puro: invariantes de negocio en el Dominio inmunes a cambios de infraestructura. Validaciones en múltiples capas.", "Validación en cliente + validación de DTOs en servidor (nunca confiar solo en el cliente). Preguntas de confirmación para acciones destructivas."],
        ["A05: Security Misconfiguration", "Configuraciones por defecto inseguras, stack traces en producción, headers de seguridad faltantes.", "ASPNETCORE_ENVIRONMENT=Production desactiva DeveloperExceptionPage y Swagger. CORS restringido a orígenes explícitos. Headers HSTS, X-Frame-Options: DENY.", "Nginx con X-Content-Type-Options: nosniff, X-Frame-Options: DENY, Content-Security-Policy. Angular en modo Production sin console.log de errores técnicos."],
        ["A06: Vulnerable Components", "Dependencias NuGet/npm con CVEs conocidas. Imágenes Docker desactualizadas.", ".NET 9 LTS con actualizaciones de seguridad. Docker multi-stage con Alpine Linux 3.19 (superficie mínima). dotnet list package --vulnerable en CI/CD.", "Angular 19 con npm audit en CI/CD. Solo dependencias del registro npm oficial. npm ci --production en Dockerfile excluye devDependencies."],
        ["A07: Auth & Identity Failures", "Fuerza bruta, sesiones que no expiran, credenciales débiles, recuperación insegura.", "Rate limiting en Login (5 intentos/15 min). JWT: expiración 60 min + Refresh Token 7 días rotable. Política de contraseña segura enforced (8+ chars, números y símbolos).", "Cierre de sesión limpia localStorage y Signals. JwtInterceptor maneja renovación automática. Sin hints 'usuario no encontrado' en pantalla de login."],
        ["A08: Integrity Failures", "Descargas sin verificación de integridad, actualizaciones sin firma.", "Verificación de tipo MIME y firma de bytes mágicos en uploads de archivos de evidencia. Checksums en artefactos Docker.", "Validación de tipo de archivo y tamaño máximo en el input de upload del cliente antes de cualquier llamada HTTP."],
        ["A09: Logging & Monitoring Failures", "Ausencia de logs ante ataques en curso, logs no protegidos o no monitoreados.", "Serilog con logging estructurado en JSON. AuditInterceptor registra cada operación CUD con usuario, IP y timestamp UTC. Logs con niveles Information/Warning/Error.", "HttpErrorInterceptor captura todos los errores de red y los registra en el servicio de logging. Alertas al usuario con mensajes genéricos sin exponer detalles técnicos."],
        ["A10: SSRF", "El atacante hace que el servidor envíe peticiones HTTP a recursos internos.", "Validación estricta de URLs en webhooks. Docker Bridge Network aísla los contenedores del host. Solo Nginx expone puertos al exterior (80/443).", "Cliente Angular solo realiza llamadas a /api/* proxificados por Nginx. Sin llamadas directas a la API o recursos externos no aprobados."],
    ]
    add_custom_table(headers_owasp, rows_owasp, [1.1, 1.4, 2.0, 2.0])
    add_p("Fuente: OWASP Foundation. (2021). OWASP Top 10:2021. https://owasp.org/Top10/")

    # ─────────────────────────────────────────────────────────────────────────
    # 2.8 Benchmark y Analisis de Mercado
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.8 Benchmark Multicriterio y Análisis Competitivo de Mercado")
    add_p("El análisis de mercado de herramientas ALM de gestión de pruebas revela un ecosistema dominado por soluciones privativas SaaS con modelos de licenciamiento restrictivos. El Cuadrante Mágico de Gartner para herramientas ALM sitúa como líderes a OpenText (HP ALM), SmartBear (Zephyr), Idera (TestRail) y Atlassian (Jira + Xray). Sin embargo, el cuadrante de Gartner no evalúa el criterio de asequibilidad para PYMEs ni el cumplimiento ISTQB CTFL v4.0 integral — dimensiones donde QAMS establece una ventaja competitiva absoluta.")

    headers_bench = ["Criterio de Evaluación", "QAMS", "TestRail\n(Idera)", "Zephyr Scale\n(SmartBear)", "Jira Xray\n(Ibis)", "TestLink\n(Open Source)"]
    rows_bench = [
        ["Modelo de Licencia", "Open Source\nSelf-Hosted", "Comercial\nSaaS/On-Prem", "Comercial\nPlugin Jira", "Comercial\nPlugin Jira", "GPL\nOpen Source"],
        ["Costo anual / 15 testers", "$480 (VPS)", "$13,320 USD", "$4,500 + Jira", "$3,600 + Jira", "$480 (VPS)"],
        ["Stack tecnológico", ".NET 9 + Angular 19\n(LTS 2024)", "PHP / React\n(Stack legado)", "Java / React\n(via Jira)", "Java / React\n(via Jira)", "PHP 5.x\n(Obsoleto)"],
        ["Cumplimiento ISTQB v4.0", "100% — 7 módulos", "72% Parcial", "68% Parcial", "74% Parcial", "48% Básico"],
        ["Pruebas Estáticas (Cap.3 ISTQB)", "✅ Módulo nativo", "❌ No soportado", "❌ No soportado", "❌ No soportado", "❌ No soportado"],
        ["SBTM / Pruebas Exploratorias", "✅ Con Charters", "⚠️ Notas básicas", "⚠️ Plugin extra", "✅ Soportado", "❌ No soportado"],
        ["Editor BDD / Gherkin nativo", "✅ Integrado", "⚠️ Plugin externo", "✅ Soportado", "✅ Nativo", "❌ No soportado"],
        ["Motor Fast Runner (atajos)", "✅ P/F/B + teclado", "⚠️ Formulario std.", "⚠️ Test Player", "⚠️ Vista estándar", "⚠️ Formulario manual"],
        ["Kanban de Defectos integrado", "✅ Nativo drag&drop", "❌ No soportado", "⚠️ Requiere Jira", "⚠️ Requiere Jira", "❌ No soportado"],
        ["Gobernanza: Soft-Delete+Audit", "✅ Completo UTC", "⚠️ Logs básicos", "⚠️ Histórico Jira", "⚠️ Histórico Jira", "❌ Sin Soft-Delete"],
        ["Seguridad OWASP Top 10", "✅ AES-256+RBAC\nCompleto", "⚠️ TLS estándar", "⚠️ TLS estándar", "⚠️ TLS estándar", "❌ Sin cifrado payload"],
        ["Quality Gates automáticos", "✅ Semáforo nativo", "⚠️ Milestones", "⚠️ Manuales", "⚠️ Manuales", "❌ No soportado"],
        ["Despliegue en Docker", "✅ One-Click Compose", "⚠️ On-Prem disp.", "❌ SaaS only", "❌ SaaS only", "⚠️ Manual complejo"],
        ["Soberanía de datos", "✅ 100% Self-hosted", "⚠️ SaaS/On-Prem", "❌ SaaS en AWS", "❌ SaaS en AWS", "✅ Self-hosted"],
    ]
    add_custom_table(headers_bench, rows_bench, [1.5, 1.1, 1.1, 1.1, 1.1, 1.0])
    add_p("Fuente: Elaboración propia en base a evaluación técnica directa de las herramientas, documentación oficial y Gartner Peer Insights 2024.")

    add_image_fig("figura15_benchmark_radar.png", "Figura 2.11: Evaluación Benchmark Multicriterio de QAMS frente a Herramientas del Mercado (Gráfico Radar Multidimensional). Los ejes representan: Cumplimiento ISTQB (%), Seguridad OWASP, Ergonomía Fast Runner, Costo Total de Propiedad (invertido — menor es mejor), Cobertura de Módulos, Gobernanza de Datos y Portabilidad Docker. QAMS obtiene el perímetro de cobertura más amplio, superando a TestRail, Zephyr, Xray y TestLink en la mayoría de las dimensiones evaluadas.")
    add_image_fig("figura16_tco_comparison.png", "Figura 2.12: Comparativa de Costo Total de Propiedad (TCO) a 5 años — Equipo de 15 Testers activos. El gráfico de barras muestra el TCO acumulado de QAMS ($2,400 USD — solo infraestructura VPS) frente a TestRail ($69,000 USD), Zephyr Scale ($20,160 USD), Jira Xray ($18,720 USD) y HP ALM ($108,000+ USD). QAMS representa un ahorro operativo de entre $16,320 y $105,600 USD según la herramienta de referencia utilizada.")

    # ─────────────────────────────────────────────────────────────────────────
    # 2.9 Base de Datos — Fundamentos Teóricos del Modelo Relacional y NoSQL
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("2.9 Bases de Datos — Fundamentos Teóricos del Modelo Relacional y Almacenamiento en Memoria")
    add_p("La teoría de bases de datos constituye uno de los pilares fundamentales de la ingeniería de software moderna. Un sistema de gestión de aseguramiento de calidad como QAMS requiere una capa de persistencia que garantice la integridad de los datos de prueba, la trazabilidad histórica y la consistencia transaccional. Esta sección desarrolla los fundamentos teóricos que sustentan las decisiones de diseño de persistencia del proyecto.")

    add_h3("2.9.1 El Modelo Relacional — Teoría de Edgar F. Codd")
    add_p("El modelo relacional fue propuesto por Edgar F. Codd en su artículo seminal 'A Relational Model of Data for Large Shared Data Banks' (Communications of the ACM, 1970), considerado uno de los trabajos más influyentes en la historia de la informática. Codd formalizó matemáticamente la representación de datos mediante el concepto de relación (tabla), fundamentado en la teoría de conjuntos y la lógica de primer orden.")
    add_p("Los conceptos nucleares del modelo relacional son: (1) Relación (tabla): conjunto de tuplas (filas) que comparten el mismo esquema de atributos (columnas); (2) Dominio: conjunto de valores atómicos e indivisibles que puede tomar un atributo (Codd denominó a esta propiedad 'primera forma normal'); (3) Clave Primaria: atributo o conjunto de atributos que identifica de forma única a cada tupla dentro de una relación; (4) Clave Foránea: atributo que referencia la clave primaria de otra relación, estableciendo la integridad referencial.")
    headers_codd = ["Concepto Relacional", "Definición de Codd (1970)", "Fundamento Matemático", "Relevancia para QAMS"]
    rows_codd = [
        ["Relación (Tabla)", "Conjunto finito de n-tuplas sobre los dominios D1, D2,...,Dn. Matemáticamente: R ⊆ D1 × D2 × ... × Dn.", "Producto Cartesiano de dominios con cardinalidad finita.", "Cada entidad del dominio STLC (TestCase, Defect, Requirement) es una relación en PostgreSQL 16 con atributos tipados y restricciones de dominio."],
        ["Clave Primaria (PK)", "Atributo o conjunto minimal de atributos cuyos valores identifican unívocamente cada tupla de la relación. Ningún valor de la PK puede ser NULL.", "Superclave minimal — no existe subconjunto propio que también sea superclave.", "Todas las entidades de QAMS usan UUID v4 como clave primaria. UUID garantiza unicidad global sin coordinar con el servidor, esencial en arquitecturas distribuidas."],
        ["Integridad Referencial", "Si una relación R2 tiene una clave foránea FK que referencia la PK de R1, entonces cada valor de FK en R2 debe existir como valor de PK en R1, o ser NULL.", "Restricción de inclusión entre conjuntos de valores de dos relaciones.", "QAMS declara todas las FKs con CASCADE en EF Core, garantizando que eliminar un TestPlan también elimina lógicamente (Soft-Delete) todas las TestSuites y TestCases dependientes."],
        ["Álgebra Relacional", "Conjunto de operaciones sobre relaciones: Selección (σ), Proyección (π), Unión (∪), Diferencia (−), Producto Cartesiano (×) y Join (⋈). Son la base formal del lenguaje SQL.", "Álgebra abstracta sobre el conjunto de todas las relaciones del esquema.", "Las consultas complejas de la Matriz RTM (selección de Reqs con TCs ejecutados) se expresan como cadenas de operaciones del álgebra relacional, optimizadas por el Query Planner de PostgreSQL."],
    ]
    add_custom_table(headers_codd, rows_codd, [1.3, 1.8, 1.5, 2.1])
    add_p("Fuente: Codd, E.F. (1970). A Relational Model of Data for Large Shared Data Banks. Communications of the ACM, 13(6), 377–387.")

    add_h3("2.9.2 Propiedades ACID — Teoría de la Confiabilidad Transaccional")
    add_p("El concepto de transacción de base de datos, formalizado por Jim Gray y Andreas Reuter en 'Transaction Processing: Concepts and Techniques' (1992), define una unidad lógica de trabajo que debe ejecutarse de forma completa o no ejecutarse en absoluto. Las propiedades ACID son los cuatro axiomas que garantizan la confiabilidad de las transacciones en sistemas relacionales:")
    headers_acid = ["Propiedad ACID", "Definición Formal", "Mecanismo de Implementación", "Importancia para un Sistema de QA"]
    rows_acid = [
        ["A — Atomicidad\n(Atomicity)", "Una transacción se trata como una unidad indivisible: o todas sus operaciones se completan con éxito (COMMIT) o ninguna tiene efecto (ROLLBACK). No existe estado intermedio visible.", "Write-Ahead Logging (WAL): cada operación se registra primero en el log antes de escribirse en los datos. En caso de fallo, el log permite deshacer operaciones incompletas.", "Registrar una ejecución de prueba (TestExecution) implica: crear el registro de ejecución + actualizar el estado del caso de prueba + calcular el Pass Rate del plan. Si alguna operación falla, ACID garantiza que no queda un estado inconsistente en el sistema de QA."],
        ["C — Consistencia\n(Consistency)", "Una transacción lleva la base de datos de un estado consistente a otro estado consistente. Todas las restricciones de integridad (PKs, FKs, CHECKs, UNIQUE) deben cumplirse al finalizar.", "Verificación de restricciones al momento del COMMIT. Si alguna violación es detectada, la transacción se revierte automáticamente.", "Garantiza que nunca exista un Defect sin su TestExecution asociada, ni una TestExecution con un TestCaseId inválido. El modelo relacional de QAMS no puede quedar en estado estructuralmente incoherente."],
        ["I — Aislamiento\n(Isolation)", "Las transacciones concurrentes se ejecutan de forma aislada: los cambios intermedios de una transacción no son visibles a otras transacciones hasta el COMMIT.", "Mecanismos de control de concurrencia: Bloqueos (locks), MVCC (Multi-Version Concurrency Control). PostgreSQL usa MVCC por defecto.", "Cuando dos testers ejecutan casos de prueba simultáneamente en el Fast Runner, sus ejecuciones no se interfieren mutuamente. Cada tester ve un snapshot consistente de la base de datos durante su sesión de trabajo."],
        ["D — Durabilidad\n(Durability)", "Una vez que una transacción ha sido confirmada (COMMIT), sus efectos son permanentes y sobreviven a cualquier fallo del sistema (caídas de energía, crashes del proceso).", "Write-Ahead Logging + fsync: los datos confirmados se persisten en disco antes de retornar el éxito al cliente. En PostgreSQL, la durabilidad es configurable hasta nivel de sincronización de disco.", "Los resultados de ejecución de pruebas, una vez confirmados, son evidencias de auditoría permanentes. La durabilidad ACID garantiza que un corte eléctrico no puede corromper ni perder los registros históricos de QAMS."],
    ]
    add_custom_table(headers_acid, rows_acid, [1.2, 1.7, 1.6, 2.2])
    add_p("Fuente: Gray, J. & Reuter, A. (1992). Transaction Processing: Concepts and Techniques. Morgan Kaufmann. Haerder, T. & Reuter, A. (1983). Principles of Transaction-Oriented Database Recovery. ACM Computing Surveys, 15(4), 287–317.")

    add_h3("2.9.3 Normalización Relacional — Teoría de las Formas Normales")
    add_p("La normalización es el proceso formal de descomposición de esquemas relacionales para eliminar redundancias y anomalías de datos. Propuesta por Edgar F. Codd (1970) con la Primera Forma Normal, y extendida por William Kent y Raymond Boyce con la Segunda, Tercera y BCNF (Boyce-Codd Normal Form), la teoría de normalización establece un conjunto progresivo de requisitos formales que un esquema relacional debe satisfacer.")
    headers_fn = ["Forma Normal", "Requisito Formal", "Anomalías que Elimina", "Aplicación en el Diseño de QAMS"]
    rows_fn = [
        ["1FN — Primera Forma Normal\n(Codd, 1970)", "Todos los atributos contienen valores atómicos e indivisibles. No existen grupos repetidos ni atributos multivaluados dentro de una misma fila.", "Elimina grupos repetidos. Hace el esquema 'tabular' puro sin arrays en celdas.", "Los pasos de un caso de prueba (TestCaseSteps) son una tabla independiente con FK a TestCases, no un campo de texto con pasos separados por comas. Cada paso es un registro atómico con Número, Acción y Resultado Esperado."],
        ["2FN — Segunda Forma Normal\n(Codd, 1971)", "Cumple 1FN + todos los atributos no clave dependen completamente de toda la clave primaria (sin dependencias parciales). Aplica cuando la PK es compuesta.", "Elimina dependencias parciales. Un atributo no puede depender solo de una parte de la clave compuesta.", "La tabla de ejecuciones (TestExecutions) tiene como PK un UUID propio. El nombre del TestCase (atributo no clave) depende del TestCaseId (FK), no de la PK de TestExecution — conformidad 2FN mediante normalización con entidades separadas."],
        ["3FN — Tercera Forma Normal\n(Codd, 1971)", "Cumple 2FN + no existen dependencias transitivas: ningún atributo no clave depende de otro atributo no clave (solo de la PK).", "Elimina dependencias transitivas. Si A→B y B→C, entonces C no debe estar en la misma tabla que A.", "La tabla Defects no contiene el nombre del tester que lo reportó (ReporterName) — solo el ReporterUserId (FK a Users). El nombre se obtiene mediante JOIN, eliminando la dependencia transitiva Defect→User→UserName."],
        ["BCNF — Boyce-Codd Normal Form\n(Boyce & Codd, 1974)", "Versión más estricta de 3FN: para toda dependencia funcional X→Y, X debe ser una superclave. Elimina anomalías residuales de 3FN con claves candidatas múltiples.", "Elimina redundancias en esquemas con múltiples claves candidatas solapadas.", "La tabla RTMEntries (que vincula Requirements con TestCases) tiene dos FKs como clave compuesta (RequirementId + TestCaseId). BCNF garantiza que no existan dependencias funcionales entre RequirementId y TestCaseId."],
    ]
    add_custom_table(headers_fn, rows_fn, [1.4, 1.6, 1.4, 2.3])
    add_p("Fuente: Codd, E.F. (1971). Further Normalization of the Data Base Relational Model. IBM Research Report RJ909. Kent, W. (1983). A Simple Guide to Five Normal Forms in Relational Database Theory. Communications of the ACM, 26(2), 120–125.")

    add_h3("2.9.4 El Modelo Entidad-Relación (ER) — Metodología de Diseño Conceptual")
    add_p("El Modelo Entidad-Relación (Entity-Relationship Model), propuesto por Peter Pin-Shan Chen en su artículo 'The Entity-Relationship Model: Toward a Unified View of Data' (ACM Transactions on Database Systems, 1976), provee un lenguaje gráfico y semántico de nivel conceptual para representar la estructura de datos de un dominio de problema, independientemente del sistema de gestión de bases de datos subyacente.")
    add_p("El modelo ER introduce tres constructores fundamentales: (1) Entidades: objetos del mundo real con existencia independiente y distinguible (ej. TestCase, Defect, Requirement); (2) Atributos: propiedades que describen a las entidades (ej. Title, Status, Priority); (3) Relaciones: asociaciones semánticas entre entidades, caracterizadas por su cardinalidad (1:1, 1:N, N:M) y participación (total u opcional).")

    headers_er = ["Constructor ER", "Definición Semántica (Chen, 1976)", "Cardinalidades Fundamentales", "Ejemplo en QAMS"]
    rows_er = [
        ["Entidad Fuerte", "Objeto que tiene existencia independiente en el dominio. Posee una clave propia que la identifica unívocamente sin depender de otra entidad.", "Puede participar en relaciones con cualquier cardinalidad.", "Requirement: existe independientemente de si tiene TestCases asignados. Su clave es RequirementId (UUID). Puede existir un requisito sin ningún caso de prueba vinculado."],
        ["Entidad Débil", "Objeto cuya existencia depende de otra entidad (entidad propietaria). Su clave es parcial y se complementa con la clave de su propietario.", "Participación total (obligatoria) en la relación con su propietario.", "TestCaseStep: no puede existir sin un TestCase propietario. Su identificador es (TestCaseId, StepNumber). Si el TestCase es eliminado lógicamente, sus Steps también lo son (Soft-Delete en cascada)."],
        ["Relación 1:N\n(Uno a Muchos)", "Una instancia de la entidad A se asocia con N instancias de la entidad B, pero cada instancia de B se asocia con exactamente una de A.", "La FK se coloca en la entidad del lado 'N'.", "Un TestPlan (1) tiene muchas TestSuites (N). La FK TestPlanId se ubica en la tabla TestSuites. Un plan puede tener múltiples suites pero cada suite pertenece a un único plan."],
        ["Relación N:M\n(Muchos a Muchos)", "Una instancia de A puede asociarse con múltiples instancias de B y viceversa. Se implementa mediante una tabla de intersección.", "Se resuelve creando una entidad de asociación con las PKs de A y B como claves foráneas compuestas.", "Requirements y TestCases tienen relación N:M: un requisito puede tener múltiples casos de prueba, y un caso de prueba puede verificar múltiples requisitos. Se implementa mediante la tabla RTMEntries (RequirementId + TestCaseId)."],
    ]
    add_custom_table(headers_er, rows_er, [1.3, 1.8, 1.3, 2.3])
    add_p("Fuente: Chen, P.P. (1976). The Entity-Relationship Model: Toward a Unified View of Data. ACM Transactions on Database Systems, 1(1), 9–36.")

    add_h3("2.9.5 ORM — Mapeo Objeto-Relacional: Teoría e Impedancia")
    add_p("El Object-Relational Mapping (ORM) aborda el problema fundamental conocido como 'Object-Relational Impedance Mismatch' (Agile Data, Ambler, 2003): la incompatibilidad estructural entre el paradigma orientado a objetos (POO) — que organiza el mundo en objetos con estado y comportamiento — y el modelo relacional — que organiza el mundo en tablas, filas y columnas sin comportamiento. Esta incompatibilidad se manifiesta en cuatro dimensiones:")
    add_bullet("• Impedancia de Granularidad: Un objeto con múltiples propiedades se mapea a múltiples tablas normalizadas (ej. un objeto Address mapea a una tabla Addresses separada o a columnas inline).")
    add_bullet("• Impedancia de Identidad: En POO, la identidad se determina por referencia de objeto (dirección de memoria). En el modelo relacional, la identidad es la clave primaria. Un mismo objeto puede tener diferentes identidades en ambos mundos.")
    add_bullet("• Impedancia de Asociaciones: En POO, las asociaciones son referencias directas entre objetos (bidireccionales por defecto). En el modelo relacional, las asociaciones son claves foráneas unidireccionales que requieren JOINs explícitos.")
    add_bullet("• Impedancia de Herencia: La POO soporta jerarquías de herencia (herencia simple y múltiple). El modelo relacional no tiene un constructo nativo equivalente. Los ORMs la resuelven con tres estrategias: Table-Per-Hierarchy (TPH), Table-Per-Type (TPT) y Table-Per-Concrete-Class (TPC).")
    add_p("Los frameworks ORM modernos como Entity Framework Core 9 (Microsoft) y Hibernate (Oracle) resuelven la impedancia automáticamente, pero imponen su propia capa de abstracción. La elección del enfoque Code-First (el código de clases C# genera el esquema de base de datos mediante migraciones automáticas) versus Database-First (el esquema existente genera las clases de entidades) representa una decisión arquitectónica fundamental con implicaciones en el ciclo de vida del proyecto.")

    add_h3("2.9.6 Bases de Datos NoSQL In-Memory — Teoría del Teorema CAP")
    add_p("El Teorema CAP, propuesto por Eric Brewer en su conferencia PODC 2000 y formalmente probado por Gilbert y Lynch (2002), establece que ningún sistema de datos distribuido puede garantizar simultáneamente las tres propiedades siguientes: (C) Consistency (Consistencia): todos los nodos ven los mismos datos en el mismo momento; (A) Availability (Disponibilidad): toda petición recibe una respuesta (éxito o error), aunque no necesariamente la más reciente; (P) Partition Tolerance (Tolerancia a Particiones): el sistema continúa operando aunque fallen las comunicaciones entre nodos.")
    add_p("Los almacenes de datos en memoria como Redis (Remote Dictionary Server) adoptan el modelo CP o AP dependiendo de la configuración, y están clasificados como sistemas NoSQL (Not Only SQL). A diferencia de los sistemas relacionales con esquema rígido, Redis ofrece estructuras de datos abstractas: Strings, Hashes, Lists, Sets, Sorted Sets, Bitmaps y HyperLogLogs — cada una optimizada para casos de uso específicos de acceso ultra-rápido con latencia sub-milisegundo.")
    headers_cap = ["Modelo de Datos", "Teorema CAP", "Garantía de Consistencia", "Caso de Uso Típico", "Limitación Principal"]
    rows_cap = [
        ["Relacional (PostgreSQL)", "CP — Consistencia + Tolerancia a Particiones", "ACID completo: consistencia fuerte transaccional garantizada.", "Datos de misión crítica: finanzas, salud, auditoría forense, trazabilidad de pruebas.", "Menor rendimiento en escrituras masivas concurrentes. Escalabilidad horizontal más compleja que NoSQL."],
        ["Clave-Valor In-Memory (Redis)", "AP — Disponibilidad + Tolerancia a Particiones (configuración típica)", "Consistencia eventual por defecto. Consistencia fuerte configurable con Redis Sentinel/Cluster.", "Caché de alta velocidad, sesiones, contadores distribuidos, colas de mensajes, pub/sub.", "Los datos son volátiles por naturaleza (TTL). No adecuado para datos permanentes de auditoría sin persistencia AOF/RDB configurada."],
        ["Documental (MongoDB)", "AP — Disponibilidad + Tolerancia a Particiones", "Consistencia eventual. Transacciones multi-documento con ACID desde v4.0.", "Catálogos de productos, contenido web, datos semi-estructurados, logs de aplicación.", "Ausencia de JOIN nativos eficientes. La denormalización necesaria introduce redundancia gestionable manualmente."],
        ["Columnar (Cassandra)", "AP — Alta Disponibilidad sin punto único de falla", "Eventual. Tunable Consistency: el cliente configura el nivel de consistencia por operación.", "IoT, series temporales, escrituras masivas de telemetría, analytics en tiempo real.", "No soporta transacciones multi-fila ni JOINs. El modelo de datos debe diseñarse orientado a las queries, no al dominio."],
    ]
    add_custom_table(headers_cap, rows_cap, [1.3, 1.2, 1.4, 1.6, 1.8])
    add_p("Fuente: Brewer, E. (2000). Towards Robust Distributed Systems. PODC 2000 Keynote. Gilbert, S. & Lynch, N. (2002). Brewer's Conjecture and the Feasibility of Consistent, Available, Partition-Tolerant Web Services. ACM SIGACT News, 33(2), 51–59.")

    doc.add_page_break()
    # =========================================================================
    add_h1("Capítulo 3.- MARCO PRÁCTICO E INGENIERÍA DE REQUISITOS")
    add_p("El presente capítulo constituye el núcleo técnico del proyecto QAMS. A partir del análisis de requisitos, se formaliza la ingeniería de software del sistema: se definen los actores y sus roles, se especifican las Historias de Usuario con criterios de aceptación en estándar Gherkin (BDD), se describen los Requerimientos No Funcionales medibles bajo ISO/IEC 25010, se documenta el modelo de datos relacional sobre PostgreSQL 16, se explica la arquitectura Clean Architecture del backend (.NET 9) y la SPA Angular 19 del frontend, se detalla la estrategia de contenedorización Docker, y se presenta la Matriz RTM que vincula cada Historia de Usuario con los Objetivos Específicos del proyecto. Cada sección de este capítulo demuestra el cumplimiento progresivo del Objetivo General: desarrollar QAMS mediante análisis de requisitos, diseño del modelo de datos, construcción de módulos funcionales y contenedorización.")

    # ─────────────────────────────────────────────────────────────────────────
    # 3.1 Ámbito y Modelo de Actores
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("3.1 Análisis del Ámbito de Aplicación y Modelo de Actores")
    add_p("El ámbito de QAMS cubre el ciclo de vida completo de las pruebas de software (STLC) en organizaciones que desarrollan, certifican o auditan aplicaciones bajo el marco del estándar ISTQB CTFL v4.0. El sistema opera como una plataforma centralizada multirol donde cada actor interactúa con un conjunto específico de módulos según sus responsabilidades dentro del proceso de QA.")
    add_p("El modelo de control de acceso basado en roles (RBAC) implementado en QAMS define cinco actores con permisos granulares asignados mediante la entidad RolePermission de la base de datos. El sistema valida cada acción contra el claim de permisos incluido en el JWT Bearer Token, rechazando con HTTP 403 cualquier intento de acceso no autorizado.")
    add_image_fig("figura4_usecases_general.png", "Figura 4: Diagrama de Casos de Uso General del Sistema QAMS (5 actores, 19 módulos)")

    add_h3("Tabla 5. Actores del Sistema QAMS — Roles, Módulos y Permisos RBAC")
    headers_act = ["Actor / Rol", "Descripción del Perfil", "Módulos Accesibles (Ruta Angular)", "Permisos del Sistema", "OE Relacionado"]
    rows_act = [
        ["Administrador\n(SysAdmin)", "Gestor técnico de la plataforma. Gestiona usuarios, roles, catálogos y configuración global del sistema.", "/admin/users\n/admin/roles\n/admin/catalogs\n/admin/api-keys\n/dashboard", "USERS_VIEW, ROLES_VIEW,\nCATALOGS_VIEW, DASHBOARD_VIEW\n(todos los permisos)", "OE2 (RBAC/JWT)"],
        ["QA Manager\n(Líder de Pruebas)", "Responsable estratégico del proceso de QA. Define planes, asigna Quality Gates y monitorea métricas.", "/projects, /systems-under-test\n/test-plans, /requirements\n/reports, /reviews, /dashboard", "PROJECTS_VIEW, SUT_VIEW,\nTEST_CASES_VIEW, DASHBOARD_VIEW,\nREVIEWS_VIEW", "OE2, OE4"],
        ["QA Tester\n(Analista de Pruebas)", "Ejecutor operativo. Diseña y ejecuta casos de prueba, registra defectos y conduce sesiones exploratorias.", "/test-cases, /test-scenarios\n/test-executions, /defects\n/kanban, /exploratory, /evidences", "TEST_CASES_VIEW, EXECUTIONS_VIEW,\nDEFECTS_VIEW, KANBAN_VIEW,\nEXPLORATORY_VIEW", "OE3, OE4"],
        ["Product Owner\n(Dueño del Producto)", "Representante del cliente. Revisa cobertura RTM, Quality Gates y descarga reportes ejecutivos para decisiones de release.", "/requirements (lectura)\n/reports (Quality Gate, RTM, PDF)", "PROJECTS_VIEW, DASHBOARD_VIEW", "OE4"],
        ["Developer\n(Desarrollador)", "Receptor de defectos. Consulta descripción técnica de bugs asignados y los marca como resueltos para verificación.", "/defects (lectura + cambio estado)\n/test-executions (lectura)", "DEFECTS_VIEW, EXECUTIONS_VIEW\n(solo lectura/resolución)", "OE3"],
    ]
    add_custom_table(headers_act, rows_act, [1.1, 1.5, 1.6, 1.5, 0.8])

    # ─────────────────────────────────────────────────────────────────────────
    # 3.2 Historias de Usuario
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("3.2 Requerimientos Funcionales — Historias de Usuario (HU-01 a HU-09)")
    add_p("Las Historias de Usuario definen los requisitos funcionales del sistema QAMS desde la perspectiva del usuario final, siguiendo la narrativa estándar de ingeniería ágil: «Como [rol] quiero [acción] para [beneficio]». Cada historia incluye criterios de aceptación formalizados en lenguaje Gherkin (Given-When-Then) según la práctica BDD (Behavior-Driven Development), una tabla con los campos de la entidad de dominio involucrada y la estimación de esfuerzo en Story Points (escala Fibonacci). La sumatoria de Story Points de las 9 HUs es de 55 puntos, equivalente a 5–6 Sprints de dos semanas con velocidad de equipo de 10 SP/Sprint.")

    # HU-01
    add_h3("HU-01: Autenticación y Seguridad JWT")
    add_p("Módulo Angular: /auth/login, /auth/register, /auth/forgot-password, /auth/reset-password | OE Cumplido: OE2 (Backend — JWT, AES-256, RBAC) | Controlador: AuthController.cs | Servicio: AuthService.cs | Story Points: 8")
    add_p("Narrativa: Como usuario del sistema (Administrador, QA Manager, QA Tester, Product Owner o Developer), quiero poder autenticarme con mis credenciales de correo electrónico y contraseña, para acceder de forma segura a los módulos del sistema QAMS según el rol que me haya sido asignado.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Login exitoso:\n  Dado que el usuario existe en la base de datos con contraseña hasheada en BCrypt (factor 11)\n  Cuando ingresa su email y contraseña correctos en /auth/login\n  Entonces el sistema genera un JWT Bearer Token (firmado HMAC-SHA256, expiración 8h)\n  Y la respuesta incluye el payload con Id, FullName, Email y lista de Permisos del rol\n  Y el frontend almacena el token en memoria y redirige a /dashboard.")
    add_bullet("Escenario 2 — Credenciales incorrectas:\n  Dado que el usuario introduce una contraseña inválida\n  Cuando envía la solicitud de login\n  Entonces el servidor retorna HTTP 401 Unauthorized\n  Y el mensaje de error es genérico ('Credenciales inválidas') para prevenir enumeración de usuarios (OWASP A2).")
    add_bullet("Escenario 3 — Token expirado:\n  Dado que el token JWT del usuario ha expirado (transcurridas 8 horas)\n  Cuando el EncryptionInterceptor adjunta el token a una petición\n  Entonces el servidor retorna HTTP 401\n  Y el AuthInterceptor del frontend redirige automáticamente a /auth/login.")

    add_h4("Tabla: Entidad User — Atributos Clave (QAMS.Domain/Entities/User.cs)")
    headers_u = ["Campo (C#)", "Tipo", "Restricción / Propósito", "Patrón"]
    rows_u = [
        ["Id", "Guid", "PK, NOT NULL, NewGuid()", "IAuditable"],
        ["Username", "string", "NOT NULL, UNIQUE, max 50", "Identificador de login"],
        ["Email", "string", "NOT NULL, UNIQUE, formato email", "Recuperación de contraseña"],
        ["PasswordHash", "string", "NOT NULL, BCrypt factor 11", "Seguridad OWASP A2"],
        ["FullName", "string", "NOT NULL, max 100", "Presentación en UI"],
        ["IsActive", "bool", "NOT NULL, DEFAULT true", "Control de acceso"],
        ["IsDeleted", "bool", "NOT NULL, DEFAULT false", "ISoftDelete"],
        ["DeletedAt", "DateTime?", "NULL", "ISoftDelete"],
        ["CreatedAt", "DateTime", "NOT NULL, UTC", "IAuditable"],
        ["UpdatedAt", "DateTime?", "NULL", "IAuditable"],
    ]
    add_custom_table(headers_u, rows_u, [1.3, 0.9, 2.5, 1.3])

    # HU-02
    add_h3("HU-02: Gestión de Sistemas Bajo Prueba (SUT) y Proyectos")
    add_p("Módulo Angular: /systems-under-test, /projects | OE Cumplido: OE1 (ERD — SUT, Project), OE2 (CRUD API) | Controladores: SystemsUnderTestController.cs, ProjectsController.cs | Story Points: 5")
    add_p("Narrativa: Como QA Manager, quiero registrar y gestionar los Sistemas Bajo Prueba (SUT) y los Proyectos de QA asociados a cada sistema, para organizar jerárquicamente el ciclo completo de aseguramiento de calidad por aplicación.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Registro de SUT:\n  Dado que el QA Manager tiene el permiso SUT_VIEW\n  Cuando registra un SUT con nombre, versión, tipo de plataforma (Web, Mobile, API, Desktop) y URL base\n  Entonces el sistema persiste el SUT en PostgreSQL con IsActive=true y Audit Trail (CreatedAt UTC, CreatedByUserId)\n  Y el SUT aparece disponible para asociar a nuevos Proyectos.")
    add_bullet("Escenario 2 — Creación de Proyecto:\n  Dado que existe al menos un SUT registrado\n  Cuando el QA Manager crea un proyecto con nombre, fechas de inicio/fin, presupuesto y LeaderId\n  Entonces el proyecto queda vinculado al SUT seleccionado\n  Y su estado inicial es 'Activo' con versión '1.0'.")
    add_bullet("Escenario 3 — Inactivación de SUT:\n  Dado que un SUT está activo\n  Cuando el Administrador lo inactiva\n  Entonces IsActive cambia a false (eliminación lógica ISoftDelete)\n  Y el SUT no aparece en las listas de selección para nuevos proyectos\n  Pero sus proyectos históricos se conservan con integridad referencial.")

    add_h4("Tabla: Entidad SystemUnderTest — Atributos Clave")
    headers_sut = ["Campo (C#)", "Tipo", "Restricción", "Descripción"]
    rows_sut = [
        ["Id", "Guid", "PK, NOT NULL", "Identificador único del SUT"],
        ["Name", "string", "NOT NULL, max 100", "Nombre del sistema bajo prueba"],
        ["Version", "string?", "NULL, max 20", "Versión del sistema evaluado"],
        ["Description", "string?", "NULL", "Descripción funcional del SUT"],
        ["Environment", "string?", "NULL", "Ambiente (Dev, QA, Staging)"],
        ["PlatformTypeId", "int", "FK → CatalogBase, NOT NULL", "Web, Mobile, API, Desktop, etc."],
        ["BaseUrl", "string?", "NULL, formato URL", "URL base para pruebas de API/Web"],
        ["IsActive", "bool", "NOT NULL, DEFAULT true", "Estado activo/inactivo"],
        ["Projects", "ICollection<Project>", "Nav. property 1:N", "Proyectos QA asociados"],
    ]
    add_custom_table(headers_sut, rows_sut, [1.3, 1.1, 1.4, 2.2])

    # HU-03
    add_h3("HU-03: Gestión de Requisitos y Matriz RTM Bidireccional")
    add_p("Módulo Angular: /requirements, /reports (tab RTM) | OE Cumplido: OE4 (ISTQB — RTM bidireccional), OE1 (ERD — Requirement, RequirementTestCase) | Controladores: RequirementsController.cs, ReportsController.cs | Story Points: 8")
    add_p("Narrativa: Como QA Lead, quiero registrar los requisitos funcionales y no funcionales del proyecto, vincularlos mediante una relación M:N a los casos de prueba diseñados, y consultar la Matriz RTM en tiempo real para verificar que la cobertura de requisitos sea completa antes de autorizar el pase a producción.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Registro de Requisito:\n  Dado que existe un proyecto activo\n  Cuando el QA Lead crea un requisito con código (REQ-001), título, tipo (Funcional/No Funcional), prioridad y criterios de aceptación\n  Entonces el sistema persiste el requisito con RequirementTypeId y RequirementPriorityId desde catálogos\n  Y el requisito aparece disponible en la RTM con estado 'Sin cobertura'.")
    add_bullet("Escenario 2 — Vinculación TC ↔ Requisito (tabla puente M:N):\n  Dado que existen casos de prueba diseñados y un requisito registrado\n  Cuando el QA Tester vincula el caso de prueba al requisito desde /test-cases o /requirements\n  Entonces el sistema inserta un registro en la tabla RequirementTestCase (RequirementId + TestCaseId)\n  Y la RTM actualiza el estado del requisito a 'Con cobertura'.")
    add_bullet("Escenario 3 — Consulta RTM con métricas:\n  Dado que el proyecto tiene requisitos y casos de prueba vinculados\n  Cuando el QA Lead consulta GET /api/Reports/rtm-matrix?projectId={id}\n  Entonces el sistema retorna: totalRequirements, coveredRequirements, coveragePercentage\n  Y la tabla RTM muestra cada fila: Requisito ↔ TestCase ↔ ExecutionStatus ↔ Defecto con severity.")

    add_h4("Tabla: Entidad Requirement y RequirementTestCase")
    headers_req = ["Campo", "Entidad", "Tipo", "Descripción"]
    rows_req = [
        ["Id", "Requirement", "Guid PK", "Identificador del requisito"],
        ["ProjectId", "Requirement", "Guid FK → Project", "Proyecto propietario"],
        ["Code", "Requirement", "string (REQ-001)", "Código legible único por proyecto"],
        ["Title", "Requirement", "string NOT NULL", "Título del requisito"],
        ["Description", "Requirement", "string?", "Descripción detallada"],
        ["AcceptanceCriteria", "Requirement", "string?", "Criterios verificables"],
        ["RequirementTypeId", "Requirement", "int FK → Catálogo", "Funcional / No Funcional"],
        ["RequirementPriorityId", "Requirement", "int FK → Catálogo", "Alta / Media / Baja"],
        ["RequirementId", "RequirementTestCase", "Guid FK → Requirement", "Tabla puente M:N RTM"],
        ["TestCaseId", "RequirementTestCase", "Guid FK → TestCase", "Caso de prueba vinculado"],
    ]
    add_custom_table(headers_req, rows_req, [1.3, 1.3, 1.2, 2.2])

    # HU-04
    add_h3("HU-04: Planes y Suites de Prueba con Quality Gates")
    add_p("Módulo Angular: /test-plans, /test-scenarios | OE Cumplido: OE2 (API Backend), OE4 (ISTQB — Test Planning IEEE 829, Quality Gates) | Controladores: TestPlansController.cs, TestSuitesController.cs | Story Points: 8")
    add_p("Narrativa: Como QA Lead, quiero crear Planes de Prueba formales con criterios de entrada/salida (ENTRY/EXIT), hitos de tiempo, riesgos identificados y un umbral de Quality Gate configurable (ej. Pass Rate ≥ 90%), y organizar los casos de prueba en Suites (Test Scenarios), para formalizar y gobernar la estrategia de pruebas del proyecto.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Creación de Plan de Prueba:\n  Dado que existe un proyecto activo\n  Cuando el QA Lead crea un Test Plan con nombre, objetivos, alcance (Scope), estrategia (TestStrategyId) y nivel de prueba (TestLevelId)\n  Entonces el plan se persiste con StatusId=Borrador y queda disponible para agregar criterios, hitos y riesgos.\n  Y el sistema registra CreatedAt UTC y CreatedByUserId en Audit Trail.")
    add_bullet("Escenario 2 — Configuración de Criterios ENTRY/EXIT:\n  Dado que un Test Plan existe en estado Borrador\n  Cuando el QA Lead agrega criterios con CriteriaType='ENTRY' (ej. 'Ambiente configurado') y CriteriaType='EXIT' (ej. 'Pass Rate >= 90%')\n  Entonces cada criterio se persiste en TestPlanCriteria con IsMet=false\n  Y el plan pasa a estado 'En revisión' cuando todos los ENTRY están cumplidos.")
    add_bullet("Escenario 3 — Evaluación del Quality Gate:\n  Dado que el Test Plan tiene criterios EXIT definidos\n  Cuando el sistema evalúa las métricas de ejecución del proyecto\n  Entonces el Quality Gate Widget muestra: PassRate, DRE (Defect Removal Efficiency), MTTR\n  Y el semáforo del Quality Gate es VERDE si PassRate ≥ umbral configurado, ROJO si no.")

    # HU-05
    add_h3("HU-05: Diseño de Casos de Prueba Clásico y BDD/Gherkin")
    add_p("Módulo Angular: /test-cases, /test-scenarios/:id | OE Cumplido: OE1 (ERD — TestCase, TestStep), OE4 (ISTQB — BDD/Gherkin, Risk-Based Testing) | Controladores: TestCasesController.cs, TestSuitesController.cs | Story Points: 8")
    add_p("Narrativa: Como QA Tester, quiero diseñar casos de prueba en dos modalidades — clásica (precondiciones, pasos estructurados, resultado esperado) o BDD con sintaxis Gherkin (Given-When-Then) — con evaluación de riesgo RBT (impacto × probabilidad), para documentar formalmente los criterios de verificación de los requisitos del sistema bajo prueba.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Caso Clásico:\n  Dado que existe una Suite de Prueba en el proyecto\n  Cuando el QA Tester crea un caso con título, precondiciones, pasos (TestStep[]) con acción y resultado esperado, prioridad, tipo de prueba y técnica de diseño\n  Entonces el caso se persiste con IsBdd=false y VersionNumber=1\n  Y el campo RiskScore se calcula automáticamente: RiskScore = ImpactLevel × LikelihoodLevel (escala 1-5).")
    add_bullet("Escenario 2 — Caso BDD/Gherkin:\n  Dado que el QA Tester activa el modo BDD en el editor\n  Cuando escribe el escenario Gherkin en el campo BddScenario (Given-When-Then)\n  Entonces el caso se persiste con IsBdd=true y BddScenario almacena el texto completo del escenario\n  Y el caso queda vinculado a la Suite y disponible para ejecución en el Fast Runner.")
    add_bullet("Escenario 3 — Vinculación a Requisito:\n  Dado que existe un requisito registrado y un caso de prueba diseñado\n  Cuando el tester selecciona los RequirementIds en el editor del caso de prueba\n  Entonces se insertan registros en RequirementTestCase (M:N) para cada requisito seleccionado\n  Y la RTM refleja inmediatamente la nueva cobertura del requisito.")

    add_h4("Tabla: Entidad TestCase — Atributos Clave (TestCase.cs)")
    headers_tc = ["Campo (C#)", "Tipo", "Restricción / Valor", "Descripción"]
    rows_tc = [
        ["Id", "Guid", "PK, NOT NULL", "Identificador único del caso"],
        ["ProjectId", "Guid", "FK → Project, NOT NULL", "Proyecto propietario"],
        ["TestSuiteId", "Guid", "FK → TestSuite, NOT NULL", "Suite contenedora"],
        ["Title", "string", "NOT NULL, max 200", "Nombre descriptivo del caso"],
        ["Preconditions", "string", "NOT NULL", "Estado previo requerido"],
        ["ExpectedResult", "string", "NOT NULL", "Resultado esperado general"],
        ["Postconditions", "string?", "NULL", "Estado posterior a la prueba"],
        ["PriorityId", "int", "FK → TestCasePriority", "Alta / Media / Baja / Crítica"],
        ["IsBdd", "bool", "DEFAULT false", "Modo Gherkin activado"],
        ["BddScenario", "string?", "NULL si IsBdd=false", "Texto Gherkin Given-When-Then"],
        ["ImpactLevel", "int", "1-5, DEFAULT 3", "Nivel de impacto (RBT)"],
        ["LikelihoodLevel", "int", "1-5, DEFAULT 3", "Probabilidad de fallo (RBT)"],
        ["RiskScore", "int (computed)", "ImpactLevel × LikelihoodLevel", "Score de riesgo 1–25"],
        ["VersionNumber", "int", "DEFAULT 1, autoincrementa", "Control de versiones del TC"],
        ["TestSteps", "ICollection<TestStep>", "Nav. property 1:N", "Pasos del caso clásico"],
        ["RequirementTestCases", "ICollection<RequirementTestCase>", "Nav. property M:N", "Trazabilidad RTM"],
    ]
    add_custom_table(headers_tc, rows_tc, [1.5, 1.2, 1.5, 1.8])

    # HU-06
    add_h3("HU-06: Motor Fast Runner de Ejecución de Pruebas")
    add_p("Módulo Angular: /test-executions, /test-executions/:id | OE Cumplido: OE3 (Frontend — Fast Runner, Angular Signals, Standalone Components) | Controladores: TestExecutionsController.cs | Servicio .NET: TestExecutionService.cs | Story Points: 13")
    add_p("Narrativa: Como QA Tester, quiero ejecutar casos de prueba con atajos de teclado reactivos en un motor de ejecución rápida (Fast Runner), que registra el estado de cada paso (PASSED/FAILED/BLOCKED) en tiempo real mediante Angular Signals sin recargas de página, y permita adjuntar evidencias (screenshots, archivos) durante la ejecución, para documentar los resultados de forma eficiente durante sesiones de testing de alto volumen.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Inicio de corrida de ejecución:\n  Dado que el QA Tester selecciona un Test Plan y una Suite activa\n  Cuando inicia una corrida en /test-executions\n  Entonces el sistema genera un TestExecution por cada TestCase de la Suite con StatusId=PENDING\n  Y el Fast Runner presenta la lista de casos con navegación por teclado (↑/↓).")
    add_bullet("Escenario 2 — Registro de paso PASSED con atajo de teclado:\n  Dado que el Fast Runner muestra el paso actual de ejecución\n  Cuando el tester presiona la tecla P (PASSED), F (FAILED) o B (BLOCKED)\n  Entonces el Signal del componente Angular se actualiza sin Zone.js (latencia < 200ms)\n  Y el sistema llama a PATCH /api/TestExecutions/{id}/step-result\n  Y el step result se persiste en ExecutionStepResult con StatusId correspondiente.")
    add_bullet("Escenario 3 — Adjuntar evidencia durante ejecución:\n  Dado que un paso ha sido marcado como FAILED\n  Cuando el tester adjunta un archivo (screenshot, log) al paso fallido\n  Entonces el sistema sube el archivo al volumen Docker /app/wwwroot/uploads\n  Y persiste un registro Evidence con FileUrl, FileName, EvidenceTypeId y TestExecutionId\n  Y el defecto puede crearse automáticamente desde el paso fallido.")

    add_h4("Tabla: Entidad TestExecution — Atributos Clave")
    headers_te = ["Campo (C#)", "Tipo", "Restricción", "Descripción"]
    rows_te = [
        ["Id", "Guid", "PK, NOT NULL", "Identificador de la ejecución"],
        ["TestCaseId", "Guid", "FK → TestCase, NOT NULL", "Caso de prueba ejecutado"],
        ["TestPlanId", "Guid?", "FK → TestPlan, NULL", "Plan de prueba asociado"],
        ["TesterId", "Guid", "FK → User, NOT NULL", "Tester que ejecuta"],
        ["StatusId", "int", "FK → ExecutionStatus", "PENDING/PASSED/FAILED/BLOCKED"],
        ["Notes", "string?", "NULL", "Observaciones de la ejecución"],
        ["ActualTimeHours", "decimal", "DEFAULT 0", "Tiempo real invertido (horas)"],
        ["ExecutionDate", "DateTime", "NOT NULL, UTC", "Fecha y hora de ejecución"],
        ["CycleNumber", "int", "DEFAULT 1", "Número de ciclo de prueba"],
        ["StepResults", "ICollection<ExecutionStepResult>", "Nav. 1:N", "Resultados por paso"],
        ["Evidences", "ICollection<Evidence>", "Nav. 1:N", "Archivos de evidencia"],
    ]
    add_custom_table(headers_te, rows_te, [1.4, 1.2, 1.4, 2.0])

    # HU-07
    add_h3("HU-07: Gestión del Ciclo de Vida de Defectos y Tablero Kanban")
    add_p("Módulo Angular: /defects, /kanban | OE Cumplido: OE3 (Frontend — Kanban drag-and-drop, Angular Signals) | Controladores: DefectsController.cs, KanbanController.cs | Servicios: DefectService.cs, KanbanService.cs | Story Points: 8")
    add_p("Narrativa: Como QA Tester, quiero registrar defectos detectados durante la ejecución de pruebas con trazabilidad completa hacia el TestCase y la TestExecution que los originó, y gestionarlos visualmente en un tablero Kanban con estados (Nuevo → Asignado → En Progreso → Resuelto → Verificado → Cerrado), para garantizar el seguimiento formal del ciclo de vida de cada incidente según el estándar ISTQB.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Creación de defecto desde ejecución fallida:\n  Dado que un paso de la TestExecution tiene StatusId=FAILED\n  Cuando el QA Tester completa el formulario de defecto (título, descripción, severidad, pasos para reproducir, resultado actual vs esperado)\n  Entonces el sistema persiste el Defecto con TestCaseId y TestExecutionId como FK de trazabilidad\n  Y el defecto aparece en estado 'Nuevo' (StatusId=1) en el Kanban del proyecto.")
    add_bullet("Escenario 2 — Movimiento drag-and-drop en Kanban:\n  Dado que el tablero Kanban muestra las columnas del KanbanBoard con sus KanbanTasks\n  Cuando el QA Manager arrastra un defecto de la columna 'Asignado' a 'En Progreso'\n  Entonces el sistema actualiza el StatusId del Defecto vía PATCH /api/Defects/{id}\n  Y el Signal de la columna Kanban se actualiza reactivamente sin recarga de página.")
    add_bullet("Escenario 3 — Resolución y verificación:\n  Dado que el Developer ha corregido el defecto y lo marcó como 'Resuelto'\n  Cuando el QA Tester re-ejecuta el caso de prueba fallido (nuevo CycleNumber)\n  Entonces si la ejecución es PASSED, el defecto puede cerrarse definitivamente\n  Y el DRE (Defect Removal Efficiency) del proyecto se actualiza en el Quality Gate.")

    add_h4("Tabla: Entidad Defect — Atributos Clave")
    headers_def = ["Campo (C#)", "Tipo", "Restricción", "Descripción"]
    rows_def = [
        ["Id", "Guid", "PK, NOT NULL", "Identificador del defecto"],
        ["ProjectId", "Guid", "FK → Project, NOT NULL", "Proyecto propietario"],
        ["TestCaseId", "Guid?", "FK → TestCase, NULL", "TC que originó el defecto"],
        ["TestExecutionId", "Guid?", "FK → TestExecution, NULL", "Ejecución donde se detectó"],
        ["Title", "string", "NOT NULL, max 200", "Título descriptivo del bug"],
        ["Description", "string?", "NULL", "Descripción detallada"],
        ["SeverityId", "int", "FK → DefectSeverity", "Crítico/Alto/Medio/Bajo"],
        ["StatusId", "int", "FK → DefectStatus", "Nuevo/Asignado/En Progreso/Resuelto/Cerrado"],
        ["PriorityId", "int?", "FK → DefectPriority", "Alta/Media/Baja"],
        ["StepsToReproduce", "string?", "NULL", "Pasos para reproducir el fallo"],
        ["ExpectedResult", "string?", "NULL", "Comportamiento esperado"],
        ["ActualResult", "string?", "NULL", "Comportamiento observado"],
        ["AssignedToUserId", "Guid?", "FK → User, NULL", "Developer asignado"],
        ["ReportedByUserId", "Guid?", "FK → User, NULL", "Tester que reportó"],
    ]
    add_custom_table(headers_def, rows_def, [1.4, 1.1, 1.4, 2.1])

    # HU-08
    add_h3("HU-08: Pruebas Estáticas (Static Testing) y Sesiones SBTM Exploratorias")
    add_p("Módulo Angular: /reviews, /exploratory | OE Cumplido: OE4 (ISTQB Cap. 3 — Static Testing Walkthroughs/Inspecciones; Cap. 4.4 — Session-Based Test Management) | Controladores: ReviewController.cs, ExploratoryController.cs | Story Points: 8")
    add_p("Narrativa: Como QA Lead, quiero conducir sesiones formales de revisión estática de artefactos (requerimientos, arquitectura, código) siguiendo los tipos ISTQB (Walkthrough, Inspección, Revisión Técnica, Revisión Informal), y gestionar sesiones de prueba exploratoria SBTM con cartas de misión (charters) y tiempo acotado (time-box), para cumplir los módulos avanzados del estándar ISTQB CTFL v4.0 Capítulo 3 y 4.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Creación de sesión de revisión estática:\n  Dado que existe un proyecto activo y el usuario tiene permiso REVIEWS_VIEW\n  Cuando el QA Lead crea una ReviewSession con título, artefacto bajo revisión, tipo de revisión (ReviewTypeId) y lista de participantes (ReviewParticipant[])\n  Entonces la sesión se persiste con ReviewStatusId=Planificada y Audit Trail\n  Y los participantes pueden registrar hallazgos (ReviewFinding) clasificados por tipo y severidad.")
    add_bullet("Escenario 2 — Registro de hallazgos y dictamen:\n  Dado que la sesión de revisión está en estado 'En Progreso'\n  Cuando los revisores registran hallazgos con FindingTypeId (Defecto/Mejora/Pregunta) y FindingSeverityId (Mayor/Menor/Info)\n  Entonces cada hallazgo se persiste en ReviewFinding con referencia a la sesión\n  Y el moderador puede emitir el dictamen final (APROBADO/APROBADO CON CONDICIONES/RECHAZADO) actualizando ReviewStatusId.")
    add_bullet("Escenario 3 — Sesión SBTM con Charter:\n  Dado que el QA Tester tiene permiso EXPLORATORY_VIEW\n  Cuando crea una ExploratorySession con Charter (misión de exploración), DurationMinutes (time-box) y lo inicia\n  Entonces StartTime se registra automáticamente y StatusId=Running\n  Y al finalizar, EndTime se registra, se calculan los DurationMinutes efectivos\n  Y los Findings de la sesión quedan vinculados mediante ExploratoryFinding[].")

    # HU-09
    add_h3("HU-09: Módulo de Reportes, Dashboard Analítico y Quality Gates")
    add_p("Módulo Angular: /reports, /dashboard | OE Cumplido: OE3 (Frontend — Dashboard Ng2-Charts, Angular Signals), OE4 (ISTQB — Quality Gates, RTM, DRE) | Controladores: ReportsController.cs, DashboardController.cs | Servicios: IReportService (7 tipos PDF), IDashboardService | Story Points: 13")
    add_p("Narrativa: Como QA Lead o Product Owner, quiero consultar un dashboard analítico en tiempo real con KPIs del proyecto (Pass Rate, DRE, cobertura RTM, defectos abiertos), ver la Matriz RTM bidireccional, el mapa de calor de riesgos RBT y el semáforo del Quality Gate, y descargar reportes PDF formales (Test Summary Report, Certificado de Cumplimiento, Resumen Ejecutivo, Burndown) para soportar las decisiones de release y auditoría del proceso.")
    add_p("Criterios de Aceptación (Gherkin):")
    add_bullet("Escenario 1 — Dashboard con KPIs ISTQB:\n  Dado que el usuario autenticado tiene permiso DASHBOARD_VIEW\n  Cuando accede a /dashboard\n  Entonces el sistema llama GET /api/Dashboard y muestra: TotalProjects, TotalTestCases, PassRate, PassedExecutions, FailedExecutions, OpenDefects, RequirementCoverageRate\n  Y el gráfico Doughnut muestra distribución de ejecuciones por estado (ExecutionsByStatus)\n  Y el gráfico de barras muestra progreso de tareas Kanban por columna (TaskProgress).")
    add_bullet("Escenario 2 — Quality Gate semáforo:\n  Dado que el QualityGateWidgetComponent carga las métricas ISTQB del proyecto\n  Cuando evalúa: DDP (Defect Detection Percentage), DRE (Defect Removal Efficiency) y MTTR (Mean Time to Resolution)\n  Entonces muestra VERDE si todos los KPIs superan los umbrales configurados en el Test Plan\n  Y muestra ROJO con los KPIs en fallo identificados para acción correctiva.")
    add_bullet("Escenario 3 — Generación de reporte PDF:\n  Dado que el QA Lead selecciona el tipo de reporte (general/burndown/compliance/executive-summary/full-certification)\n  Cuando hace clic en 'Generar PDF'\n  Entonces el frontend llama al endpoint correspondiente de GET /api/Reports/project/{id}/{tipo}\n  Y el servidor genera el PDF con Puppeteer/Chromium (montado en el contenedor Docker del backend)\n  Y el navegador descarga el archivo con nombre con timestamp (ej. Resumen_Ejecutivo_20260822.pdf).")

    add_h4("Tabla: Tipos de Reporte PDF — QAMS IReportService")
    headers_rpt = ["Reporte", "Endpoint REST", "Audiencia", "Contenido Principal"]
    rows_rpt = [
        ["Reporte General de Ejecuciones", "GET /api/Reports/project", "QA Lead", "Resultados de ejecuciones con filtros por fecha, estado y tester"],
        ["Burndown del Proyecto", "GET /api/Reports/project/{id}/burndown", "QA Lead / PM", "Progreso de ejecuciones vs tiempo (gráfico de líneas)"],
        ["Reporte de Observaciones", "GET /api/Reports/project/{id}/observations", "QA Lead", "Hallazgos de revisiones estáticas y sesiones exploratorias"],
        ["Certificado de Cumplimiento Final", "GET /api/Reports/project/{id}/compliance", "PM / Dirección", "Certificado formal de cumplimiento de criterios EXIT del Test Plan"],
        ["Test Summary Report (ISTQB)", "GET /api/Reports/test-plan/{id}/summary", "QA Lead", "Resumen ISTQB del plan: objetivos, resultados, defectos, métricas DRE"],
        ["Certificación Completa", "GET /api/Reports/project/{id}/full-certification", "Dirección", "Reporte exhaustivo de todo el ciclo: RTM + defectos + Quality Gate + métricas"],
        ["Resumen Ejecutivo de Liberación", "GET /api/Reports/project/{id}/executive-summary", "Product Owner", "1-2 páginas para decisión de release: semáforo Quality Gate + riesgos abiertos"],
    ]
    add_custom_table(headers_rpt, rows_rpt, [1.6, 2.0, 0.9, 2.0])

    add_h4("Tabla: KPIs del Dashboard ISTQB — DashboardService.cs")
    headers_kpi = ["KPI", "Fórmula de Cálculo", "Umbral Quality Gate", "Tab en /reports"]
    rows_kpi = [
        ["Pass Rate (%)", "passedTestCaseIds.Count / TotalTestCases × 100", "≥ 90% para liberar", "qualityGate"],
        ["DRE — Defect Removal Efficiency", "(Defectos pre-producción / Total defectos) × 100", "≥ 85%", "qualityGate"],
        ["DDP — Defect Detection %", "Defectos detectados en QA / Total defectos reportados", "≥ 80%", "qualityGate"],
        ["MTTR — Mean Time to Resolution", "Σ(ClosedAt - CreatedAt) / COUNT(cerrados) en horas", "≤ 48h promedio", "qualityGate"],
        ["Cobertura RTM (%)", "coveredRequirements / totalRequirements × 100", "≥ 95%", "rtm"],
        ["Defectos Abiertos", "COUNT(Defecto WHERE status NOT IN (CLOSED, RESOLVED))", "= 0 críticos/altos", "qualityGate"],
        ["Cobertura por Ciclo", "passedTestCases / totalTestCases por CycleNumber", "Tendencia creciente", "cycles"],
        ["RiskScore promedio", "Σ(ImpactLevel × LikelihoodLevel) / COUNT(TestCase)", "< 9 promedio", "rbt"],
    ]
    add_custom_table(headers_kpi, rows_kpi, [1.4, 2.0, 1.2, 0.9])

    # ─────────────────────────────────────────────────────────────────────────
    # 3.3 Requerimientos No Funcionales
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("3.3 Especificación de Requerimientos No Funcionales (RNF-01 a RNF-08)")
    add_p("Los Requerimientos No Funcionales (RNF) de QAMS se especifican bajo las características de calidad del estándar ISO/IEC 25010 (SQuaRE). Cada RNF tiene una métrica de aceptación cuantificable y una herramienta o técnica de validación específica. Estos RNFs establecen los límites de calidad que el sistema debe cumplir independientemente de la funcionalidad implementada.")
    headers_rnf = ["Código", "Categoría ISO 25010", "Descripción del Requerimiento", "Métrica de Aceptación", "Herramienta de Validación"]
    rows_rnf = [
        ["RNF-01", "Seguridad\n(Security)", "Cifrado de credenciales (contraseñas) en reposo y cifrado de payloads HTTP en tránsito. Autenticación con token JWT y control de acceso RBAC.", "BCrypt factor 11 para contraseñas. AES-256 en payloads. JWT firmado HMAC-SHA256 con expiración 8h. HTTP 403 para acceso no autorizado.", "Prueba manual de endpoints con Postman. Inspección de hashes en PostgreSQL. Análisis de tráfico con Wireshark/Fiddler."],
        ["RNF-02", "Eficiencia de\nRendimiento\n(Performance)", "Tiempo de respuesta del API REST bajo carga normal (hasta 50 usuarios concurrentes).", "Latencia P95 < 250 ms en consultas GET. Latencia P95 < 400 ms en escrituras POST/PATCH. Throughput ≥ 50 RPS.", "Pruebas de carga K6 con escenarios de 1, 10, 50 usuarios virtuales. Métricas P50, P95, P99."],
        ["RNF-03", "Confiabilidad\n(Reliability)", "Disponibilidad del servicio y tolerancia a fallos transaccionales. El sistema no debe perder datos ante fallos parciales.", "SLA ≥ 99.5% en entorno productivo. Transacciones ACID completas con rollback automático vía EF Core Unit of Work. Healthcheck activo en Docker.", "Healthcheck Docker: wget /health cada 10s. Pruebas de rollback con fallos simulados. Monitoreo uptime."],
        ["RNF-04", "Compatibilidad\n(Compatibility)", "El sistema debe funcionar en los principales navegadores web modernos sin plugins adicionales.", "Compatible con Chrome 120+, Firefox 121+, Edge 120+, Safari 17+. SPA renderizada sin errores en todos.", "Pruebas de compatibilidad con BrowserStack o Playwright multi-browser. Tests E2E del repositorio con Playwright."],
        ["RNF-05", "Escalabilidad\n(Scalability)", "Capacidad de escalar horizontalmente mediante contenedores Docker sin cambios de código.", "El sistema soporta mínimo 200 conexiones simultáneas por instancia. Escalable con docker compose --scale backend=2.", "Pruebas de carga K6 con 100-200 usuarios virtuales. Verificación de sessions Redis compartidas entre instancias."],
        ["RNF-06", "Mantenibilidad\n(Maintainability)", "El código fuente debe estar estructurado en capas desacopladas con cobertura de pruebas unitarias.", "Clean Architecture en 4 capas. Interfaces para inversión de dependencias. Cobertura de pruebas unitarias ≥ 70% (QAMS.Tests).", "Análisis estático con SonarQube o dotnet-coverage. Revisión de acoplamiento entre capas."],
        ["RNF-07", "Portabilidad\n(Portability)", "El sistema debe poder desplegarse en cualquier entorno que soporte Docker (On-Premise, Cloud AWS/Azure/GCP).", "Dockerfiles multi-stage reproducibles. docker compose up -d levanta el stack completo en Linux, Windows y macOS. Imágenes Alpine ≤ 150 MB.", "docker build en CI/CD GitHub Actions. Prueba en Docker Desktop Windows y Linux VM."],
        ["RNF-08", "Usabilidad\n(Usability)", "La interfaz del Fast Runner debe permitir registrar resultados de ejecución sin fricción operativa, a una velocidad de al menos 30 casos por hora.", "Latencia UI < 200ms por acción de teclado P/F/B en Fast Runner (medida con Angular DevTools). Sin recargas de página.", "Medición con Angular DevTools (Change Detection cycles). Prueba de usuario cronometrada con 30 casos de prueba."],
    ]
    add_custom_table(headers_rnf, rows_rnf, [0.6, 1.1, 2.0, 1.6, 1.2])

    # ─────────────────────────────────────────────────────────────────────────
    # 3.4 Modelo de Datos ERD
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("3.4 Modelo de Datos — Diagrama Entidad-Relación (ERD)")
    add_p("El modelo de datos de QAMS fue diseñado en PostgreSQL 16 aplicando los principios de normalización relacional hasta la Tercera Forma Normal (3FN), eliminando dependencias transitivas y garantizando la integridad referencial mediante claves foráneas. El modelo contiene 20 entidades principales, más de 30 tablas de catálogos y 1 tabla puente (RequirementTestCase) para la relación M:N de la RTM. Todas las entidades del dominio implementan dos interfaces de gobernanza de datos: IAuditable (cuatro campos de auditoría) e ISoftDelete (tres campos de eliminación lógica), activadas automáticamente por un interceptor del QamsDbContext al momento de llamar a SaveChangesAsync().")

    add_h3("Tabla 6. Entidades Principales del Modelo de Datos QAMS")
    headers_erd = ["Entidad", "Tabla BD", "PK", "FKs Principales", "Patrón de Gobernanza"]
    rows_erd = [
        ["User", "users", "Guid Id", "—", "IAuditable + ISoftDelete"],
        ["SystemUnderTest", "systems_under_test", "Guid Id", "PlatformTypeId → platform_types", "IAuditable + ISoftDelete"],
        ["Project", "projects", "Guid Id", "SystemUnderTestId → SUT\nLeaderId → User", "IAuditable + ISoftDelete"],
        ["ProjectTester", "project_testers", "Guid Id", "ProjectId → Project\nUserId → User", "—"],
        ["Requirement", "requirements", "Guid Id", "ProjectId → Project\nRequirementTypeId, PriorityId, StatusId", "IAuditable + ISoftDelete"],
        ["RequirementTestCase", "requirement_test_cases", "(RequirementId, TestCaseId)", "RequirementId → Requirement\nTestCaseId → TestCase", "Tabla puente M:N RTM"],
        ["TestPlan", "test_plans", "Guid Id", "ProjectId, TestStrategyId\nTestLevelId, TestPlanTypeId", "IAuditable + ISoftDelete"],
        ["TestPlanCriteria", "test_plan_criteria", "Guid Id", "TestPlanId → TestPlan", "—"],
        ["TestPlanMilestone", "test_plan_milestones", "Guid Id", "TestPlanId → TestPlan", "—"],
        ["TestPlanRisk", "test_plan_risks", "Guid Id", "TestPlanId → TestPlan", "—"],
        ["TestSuite", "test_suites", "Guid Id", "ProjectId, StatusId\nTestLevelId, TestTypeId", "IAuditable + ISoftDelete"],
        ["TestCase", "test_cases", "Guid Id", "ProjectId, TestSuiteId\nPriorityId, TestTypeId\nParentTestCaseId (self-ref)", "IAuditable + ISoftDelete"],
        ["TestStep", "test_steps", "Guid Id", "TestCaseId → TestCase", "—"],
        ["TestExecution", "test_executions", "Guid Id", "TestCaseId, TestPlanId\nTesterId → User, StatusId", "IAuditable + ISoftDelete"],
        ["ExecutionStepResult", "execution_step_results", "Guid Id", "TestExecutionId, StatusId", "—"],
        ["Evidence", "evidences", "Guid Id", "TestExecutionId, EvidenceTypeId", "IAuditable"],
        ["Defect", "defects", "Guid Id", "ProjectId, TestCaseId\nTestExecutionId\nSeverityId, StatusId, PriorityId\nAssignedToUserId, ReportedByUserId", "IAuditable + ISoftDelete"],
        ["ReviewSession", "review_sessions", "Guid Id", "ProjectId\nReviewTypeId, ReviewStatusId", "IAuditable + ISoftDelete"],
        ["ReviewFinding", "review_findings", "Guid Id", "ReviewSessionId\nFindingTypeId, FindingSeverityId", "IAuditable"],
        ["ExploratorySession", "exploratory_sessions", "Guid Id", "ProjectId, TesterId → User\nStatusId", "IAuditable + ISoftDelete"],
        ["KanbanBoard", "kanban_boards", "Guid Id", "ProjectId → Project", "IAuditable + ISoftDelete"],
        ["KanbanColumn", "kanban_columns", "Guid Id", "KanbanBoardId → KanbanBoard", "IAuditable + ISoftDelete"],
        ["KanbanTask", "kanban_tasks", "Guid Id", "KanbanColumnId, AssignedToId → User\nDefectId → Defect", "IAuditable + ISoftDelete"],
    ]
    add_custom_table(headers_erd, rows_erd, [1.2, 1.3, 0.9, 1.9, 1.2])

    add_h3("Tabla 7. Patrones de Gobernanza de Datos — IAuditable e ISoftDelete")
    headers_gov = ["Interfaz C#", "Campo", "Tipo", "Propósito / Valor"]
    rows_gov = [
        ["IAuditable", "CreatedAt", "DateTime", "Timestamp UTC de creación del registro. Inyectado automáticamente por QamsDbContext interceptor."],
        ["IAuditable", "CreatedByUserId", "Guid?", "ID del usuario que creó el registro. Extraído del JWT claim NameIdentifier."],
        ["IAuditable", "UpdatedAt", "DateTime?", "Timestamp UTC de la última modificación. Actualizado en cada SaveChangesAsync()."],
        ["IAuditable", "UpdatedByUserId", "Guid?", "ID del usuario que realizó la última modificación."],
        ["ISoftDelete", "IsDeleted", "bool", "Flag de eliminación lógica. DEFAULT false. QamsDbContext filtra WHERE IsDeleted=false en todas las consultas."],
        ["ISoftDelete", "DeletedAt", "DateTime?", "Timestamp UTC de la eliminación lógica. NULL si el registro está activo."],
        ["ISoftDelete", "DeletedByUserId", "Guid?", "ID del usuario que realizó la eliminación lógica. Trazabilidad forense completa."],
    ]
    add_custom_table(headers_gov, rows_gov, [1.2, 1.2, 0.9, 3.2])
    add_p("Principio de Integridad Histórica: Ninguna entidad del sistema QAMS es eliminada físicamente de la base de datos. El interceptor de QamsDbContext convierte automáticamente cualquier operación DELETE en un UPDATE que establece IsDeleted=true, DeletedAt=DateTime.UtcNow y DeletedByUserId=<usuario actual>. Esto garantiza la trazabilidad forense completa para auditorías y el cumplimiento del principio de no repudio en sistemas de aseguramiento de calidad.")
    add_image_fig("figura_erd_qams.png", "Figura 5: Diagrama Entidad-Relación (ERD) Completo del Sistema QAMS — PostgreSQL 16 / 23 Entidades / Normalización 3FN")

    # ─────────────────────────────────────────────────────────────────────────
    # 3.5 Arquitectura del Sistema
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("3.5 Arquitectura del Sistema QAMS")
    add_p("La arquitectura de QAMS se fundamenta en tres pilares tecnológicos complementarios: el backend en Clean Architecture con .NET 9, el frontend como SPA Angular 19 con Signals, y la infraestructura de contenedores Docker que orquesta ambos componentes junto a los servicios de persistencia y caché.")

    add_h3("3.5.1 Arquitectura Backend — Clean Architecture (.NET 9 / C# 13)")
    add_p("El backend de QAMS está implementado en .NET 9 siguiendo estrictamente el patrón de Clean Architecture (también conocido como Ports & Adapters o Arquitectura Hexagonal). Este patrón organiza el código en capas concéntricas donde las dependencias solo fluyen hacia adentro — las capas externas conocen a las internas, pero nunca al revés — protegiendo el Dominio de cambios tecnológicos externos (base de datos, framework HTTP, servicios de correo).")
    add_p("La solución QAMS.sln contiene 5 proyectos .NET organizados en 4 capas arquitectónicas:")

    add_h4("Tabla 8. Componentes por Capa — Clean Architecture QAMS")
    headers_ca = ["Capa", "Proyecto .NET", "Responsabilidad", "Componentes Clave", "Dependencias"]
    rows_ca = [
        ["Domain\n(Núcleo)", "QAMS.Domain", "Entidades de negocio, interfaces de repositorios, excepciones de dominio y contratos de servicios. Sin dependencias externas.", "20 Entidades (TestCase, Defect...)\nIAuditable, ISoftDelete\n15 IRepository interfaces\nIUnitOfWork\nDomainException, EntityNotFoundException", "Ninguna (cero dependencias)"],
        ["Application\n(Lógica)", "QAMS.Application", "Servicios de aplicación que orquestan el dominio. Implementa casos de uso del sistema. Usa inyección de dependencias.", "19 Services (AuthService, TestCaseService, DefectService, TestExecutionService, DashboardService, ReportService...)\nDTOs por módulo\nValidators (FluentValidation)\nInterfaces de servicios (IReportService, IDashboardService...)", "→ QAMS.Domain"],
        ["Infrastructure\n(Adaptadores)", "QAMS.Infrastructure", "Implementaciones concretas de repositorios, EF Core DbContext, Redis, almacenamiento de archivos y servicio de email.", "QamsDbContext (EF Core 9)\nRepositorios (ITestCaseRepository → TestCaseRepository)\n20+ Migraciones EF Core 9\nSmtpEmailService\nRedis Queue (Background Jobs)\nFileStorageService", "→ QAMS.Application\n→ QAMS.Domain"],
        ["API\n(Presentación)", "QAMS.Api", "Controladores HTTP REST que exponen los servicios de aplicación. Manejo de autenticación JWT, CORS y filtros de autorización RBAC.", "21 Controllers ([ApiController])\nHasPermission Filter (RBAC)\nJWT Bearer Configuration\nSwagger/OpenAPI\nCORS Policy\nGlobal Exception Handler", "→ QAMS.Application\n→ QAMS.Infrastructure"],
        ["Tests\n(Verificación)", "QAMS.Tests", "Pruebas unitarias de servicios de aplicación y entidades de dominio.", "Unit tests (xUnit)\nMock repositories (Moq)\nCobertura objetivo ≥ 70%", "→ QAMS.Domain\n→ QAMS.Application"],
    ]
    add_custom_table(headers_ca, rows_ca, [0.8, 1.2, 1.5, 1.8, 1.2])

    add_p("Flujo de una petición HTTP en QAMS (de afuera hacia adentro): 1) La petición HTTP llega al Controller (QAMS.Api). 2) El filtro HasPermission valida el JWT y el permiso requerido. 3) El Controller delega al Service de Aplicación (QAMS.Application). 4) El Service usa los IRepository del Dominio (QAMS.Domain) para acceder a datos. 5) La Infrastructure implementa el IRepository con EF Core 9 consultando PostgreSQL 16. 6) QamsDbContext intercepta SaveChangesAsync() para inyectar Audit Trail y Soft Delete. 7) La respuesta retorna cifrada en AES-256 por el EncryptionMiddleware.")
    add_image_fig("figura_c4_containers.png", "Figura 6: Diagrama C4 de Contenedores — QAMS (Backend .NET 9, Frontend Angular 19, PostgreSQL 16, Redis 7, Nginx)")

    add_h3("3.5.2 Arquitectura Frontend — Angular 19 / Standalone Components / Signals")
    add_p("El frontend de QAMS es una Single Page Application (SPA) construida en Angular 19 con la arquitectura de Standalone Components (sin NgModules) y Angular Signals como primitiva de reactividad. Esta combinación elimina Zone.js del árbol de detección de cambios, logrando actualizaciones de UI de microsegundos — crítico para el Fast Runner donde el tester registra hasta 200 resultados por sesión de 4 horas. La aplicación está estructurada en dos grandes layouts: auth-layout (pantallas públicas) y main-layout (aplicación con sidebar protegida por authGuard).")

    add_h4("Tabla 9. Features Angular 19 — Módulos del Sistema QAMS")
    headers_fe = ["Ruta Angular", "Feature (Componente Standalone)", "Permiso Requerido", "OE Relacionado", "Descripción Funcional"]
    rows_fe = [
        ["/auth/login", "LoginComponent", "Pública", "OE2", "Formulario de autenticación. Llama AuthController. Almacena JWT."],
        ["/auth/register", "RegisterComponent", "Pública", "OE2", "Registro de nuevo usuario. Asignación de rol inicial."],
        ["/auth/forgot-password", "ForgotPasswordComponent", "Pública", "OE2", "Recuperación de contraseña vía email SMTP (Redis Queue)."],
        ["/dashboard", "DashboardComponent", "DASHBOARD_VIEW", "OE3, OE4", "KPIs ISTQB: PassRate, DRE, Defectos abiertos, cobertura RTM. Gráficos Ng2-Charts."],
        ["/systems-under-test", "SystemsUnderTestComponent", "SUT_VIEW", "OE1, OE2", "CRUD de Sistemas Bajo Prueba con PlatformType y URL base."],
        ["/projects", "ProjectsComponent", "PROJECTS_VIEW", "OE1, OE2", "CRUD de Proyectos QA vinculados a SUT. Gestión de testers del proyecto."],
        ["/requirements", "RequirementsComponent", "PROJECTS_VIEW", "OE4", "Gestión de Requisitos funcionales y no funcionales. Vinculación M:N a TestCases para RTM."],
        ["/test-plans", "TestPlansComponent", "TEST_CASES_VIEW", "OE2, OE4", "Planes de prueba ISTQB con criterios ENTRY/EXIT, hitos, riesgos y Quality Gate."],
        ["/test-scenarios", "TestScenariosComponent", "TEST_CASES_VIEW", "OE1, OE4", "Gestión de Test Suites (Escenarios). Agrupa casos de prueba por estrategia."],
        ["/test-cases", "TestCasesComponent", "TEST_CASES_VIEW", "OE1, OE4", "Editor dual: casos clásicos (pasos) y BDD/Gherkin. Risk-Based Testing (RBT) ImpactLevel × LikelihoodLevel."],
        ["/test-executions", "TestExecutionsComponent", "EXECUTIONS_VIEW", "OE3", "Motor Fast Runner. Atajos P/F/B. Angular Signals. Registro de resultados sin recarga."],
        ["/test-executions/:id", "ExecutionDetailComponent", "EXECUTIONS_VIEW", "OE3", "Detalle de ejecución: paso a paso, evidencias adjuntas, generación de defecto automática."],
        ["/defects", "DefectsComponent", "DEFECTS_VIEW", "OE3", "CRUD de defectos con trazabilidad a TestCase y TestExecution. Severidad y prioridad."],
        ["/kanban", "KanbanComponent", "KANBAN_VIEW", "OE3", "Tablero Kanban drag-and-drop. Estados: Nuevo → Asignado → En Progreso → Resuelto → Cerrado."],
        ["/reviews", "ReviewsComponent", "REVIEWS_VIEW", "OE4", "Sesiones de Static Testing (Walkthrough, Inspección). Hallazgos con FindingType y FindingSeverity."],
        ["/exploratory", "ExploratoryComponent", "EXPLORATORY_VIEW", "OE4", "Sesiones SBTM con Charter, time-box y ExploratoryFindings."],
        ["/reports", "ReportsComponent", "DASHBOARD_VIEW", "OE3, OE4", "5 tabs: QualityGate (DDP/DRE/MTTR), RTM Matrix, RBT Risk Heatmap, Burndown, Cycles. 7 tipos de PDF."],
        ["/admin/users", "UsersComponent", "USERS_VIEW", "OE2", "Gestión de usuarios del sistema: activar, desactivar, asignar roles."],
        ["/admin/roles", "RolesComponent", "ROLES_VIEW", "OE2", "Gestión de roles y permisos RBAC. Asignación de permisos a roles."],
    ]
    add_custom_table(headers_fe, rows_fe, [1.2, 1.4, 0.9, 0.7, 2.3])

    add_p("Interceptores Angular de Seguridad (src/app/core/interceptors/): Cada petición HTTP del frontend pasa por dos interceptores antes de llegar al backend: 1) EncryptionInterceptor: cifra el cuerpo de la petición saliente con AES-256 y descifra la respuesta entrante. 2) JwtInterceptor: adjunta el header 'Authorization: Bearer <token>' al token JWT almacenado en memoria. Este doble mecanismo garantiza que incluso si un atacante intercepta el tráfico de red, los datos viajan cifrados y son inútiles sin la clave AES.")

    add_h3("3.5.3 Diagrama de Despliegue — Contenedorización Docker Compose")
    add_p("El Objetivo Específico 5 (OE5) del proyecto se materializa en la configuración Docker Compose que orquesta el stack completo de QAMS. El sistema utiliza Dockerfiles multi-stage para minimizar el tamaño de las imágenes de producción y un archivo docker-compose.yml que define la red, los volúmenes persistentes y las dependencias entre servicios. Todo el stack se levanta con un único comando: docker compose up -d.")
    add_p("Dockerfile Backend (c:/diplomado/qams/Dockerfile) — Multi-Stage:\nStage 1 (build): mcr.microsoft.com/dotnet/sdk:9.0-alpine — Restaura dependencias NuGet y publica la aplicación .NET optimizada.\nStage 2 (runtime): mcr.microsoft.com/dotnet/aspnet:9.0-alpine — Imagen mínima de producción. Usuario no privilegiado 'app'. Puerto 8080. Tamaño final < 120 MB.\nDockerfile Frontend (c:/diplomado/qams-web/Dockerfile) — Multi-Stage:\nStage 1 (build): node:20-alpine — Instala dependencias npm y ejecuta ng build --configuration=production.\nStage 2 (runtime): nginx:alpine — Sirve los archivos estáticos compilados. Puerto 80.")

    add_h4("Tabla 10. Servicios Docker Compose — Stack Completo QAMS")
    headers_dok = ["Servicio", "Imagen Docker", "Puerto", "Red / Volumen", "Healthcheck", "OE"]
    rows_dok = [
        ["qams-backend", "qams-backend:latest\n(sdk:9.0-alpine → aspnet:9.0-alpine)", "5000:8080", "qams-network\nvol: logs, uploads", "wget /health cada 10s\nRetries: 10", "OE5, OE2"],
        ["qams-frontend", "qams-frontend:latest\n(node:20 → nginx:alpine)", "4200:80", "qams-network", "wget http://127.0.0.1:80\ncada 10s, Retries: 5", "OE5, OE3"],
        ["qams-postgres", "postgres:16-alpine", "5432:5432", "qams-network\nvol: postgres_data", "pg_isready cada 10s\nRetries: 5", "OE5, OE1"],
        ["qams-redis", "redis:7-alpine", "6379:6379", "qams-network\nvol: redis_data", "redis-cli ping\ncada 10s, Retries: 5", "OE5"],
        ["qams-e2e-tests\n(perfil test)", "Playwright/Cucumber\n(Dockerfile.ci)", "—", "qams-network\nvol: playwright-report", "depends_on: frontend healthy + backend healthy", "OE4, OE5"],
    ]
    add_custom_table(headers_dok, rows_dok, [1.1, 1.6, 0.8, 1.2, 1.3, 0.5])
    add_p("Variables de Entorno de Seguridad (definidas en .env, no en el repositorio): POSTGRES_PASSWORD, JWT_SECRET (mínimo 64 caracteres), ENCRYPTION_KEY (32 bytes AES-256), ENCRYPTION_IV (16 bytes AES), SMTP_USERNAME, SMTP_PASSWORD. El archivo .env.example documenta la estructura sin valores reales, siguiendo el principio 'Secrets en variables de entorno, nunca en código fuente' (OWASP A5 — Misconfiguration).")
    add_image_fig("figura_deployment_docker.png", "Figura 7: Diagrama de Despliegue Docker Compose — Stack Completo QAMS (4 servicios + red interna + volúmenes persistentes)")

    # ─────────────────────────────────────────────────────────────────────────
    # 3.6 Trazabilidad OE vs HU
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("3.6 Matriz de Trazabilidad: Objetivos Específicos vs Historias de Usuario")
    add_p("La siguiente tabla vincula cada Historia de Usuario (HU-01 a HU-09) con el Feature Angular real, la entidad de base de datos involucrada, los Requerimientos No Funcionales aplicables y el Objetivo Específico que evidencia. Esta trazabilidad demuestra que el conjunto de las 9 HUs cubre íntegramente los 5 Objetivos Específicos y el Objetivo General del proyecto QAMS.")

    headers_rtm = ["Historia de Usuario", "Feature Angular (Ruta)", "Entidades BD Principales", "RNFs Aplicados", "OE Cumplido"]
    rows_rtm = [
        ["HU-01\nAutenticación JWT", "/auth/login\n/auth/register\n/auth/forgot-password", "User, UserRole\nRole, Permission\nRolePermission", "RNF-01 (Seguridad AES/JWT)\nRNF-02 (Rendimiento)\nRNF-06 (Mantenibilidad)", "OE2\n(Backend JWT/RBAC)"],
        ["HU-02\nGestión SUT y Proyectos", "/systems-under-test\n/projects", "SystemUnderTest\nProject, ProjectTester", "RNF-01 (Seguridad)\nRNF-03 (Confiabilidad)\nRNF-07 (Portabilidad)", "OE1 (ERD)\nOE2 (API CRUD)"],
        ["HU-03\nRequisitos y RTM", "/requirements\n/reports (tab rtm)", "Requirement\nRequirementTestCase\n(tabla puente M:N)", "RNF-02 (Rendimiento)\nRNF-03 (Confiabilidad)", "OE4\n(ISTQB RTM)"],
        ["HU-04\nPlanes y Quality Gates", "/test-plans\n/test-scenarios", "TestPlan\nTestPlanCriteria\nTestPlanMilestone\nTestPlanRisk, TestSuite", "RNF-01 (Seguridad)\nRNF-02 (Rendimiento)\nRNF-06 (Mantenibilidad)", "OE2 (API)\nOE4 (ISTQB Plans)"],
        ["HU-05\nCasos de Prueba\nClásico + BDD", "/test-cases\n/test-scenarios/:id", "TestCase, TestStep\nRequirementTestCase", "RNF-02 (Rendimiento)\nRNF-08 (Usabilidad)", "OE1 (ERD TestCase)\nOE4 (BDD/Gherkin)"],
        ["HU-06\nFast Runner Ejecución", "/test-executions\n/test-executions/:id", "TestExecution\nExecutionStepResult\nEvidence", "RNF-02 (Rendimiento)\nRNF-08 (Usabilidad < 200ms)\nRNF-03 (Confiabilidad)", "OE3\n(Frontend Signals)"],
        ["HU-07\nDefectos y Kanban", "/defects\n/kanban", "Defect\nKanbanBoard\nKanbanColumn\nKanbanTask", "RNF-01 (Seguridad)\nRNF-02 (Rendimiento)\nRNF-08 (Usabilidad)", "OE3\n(Frontend Kanban)"],
        ["HU-08\nStatic Testing y SBTM", "/reviews\n/exploratory", "ReviewSession\nReviewFinding\nReviewParticipant\nExploratorySession\nExploratoryFinding", "RNF-03 (Confiabilidad)\nRNF-06 (Mantenibilidad)", "OE4\n(ISTQB Static/SBTM)"],
        ["HU-09\nReportes y Dashboard", "/reports\n/dashboard", "DashboardSummaryDto\n(agrega: Project, TestCase\nTestExecution, Defect\nRequirement, KanbanTask)", "RNF-02 (Rendimiento)\nRNF-04 (Compatibilidad)\nRNF-08 (Usabilidad)", "OE3 (Dashboard)\nOE4 (Quality Gates)"],
    ]
    add_custom_table(headers_rtm, rows_rtm, [1.1, 1.3, 1.3, 1.4, 1.0])

    add_p("Cobertura del Objetivo General: El análisis de la Tabla de Trazabilidad demuestra que las 9 Historias de Usuario cubren íntegramente los 4 componentes del Objetivo General: (1) 'análisis de requisitos' → HU-01 a HU-05 definen el alcance funcional del sistema; (2) 'diseño del modelo de datos relacional' → OE1 verificado por HU-02, HU-03, HU-04 y HU-05 que usan las entidades 3FN de PostgreSQL; (3) 'construcción de los módulos funcionales' → OE2 y OE3 verificados por HU-06, HU-07, HU-08 y HU-09 construidos en .NET 9 y Angular 19; (4) 'contenedorización mediante Docker' → OE5 materializado por la Sección 3.5.3 con los 4 servicios Docker que empaquetan el sistema completo.")

    # ─────────────────────────────────────────────────────────────────────────
    # 3.7 Diagramas de Secuencia
    # ─────────────────────────────────────────────────────────────────────────
    add_h2("3.7 Diagramas de Secuencia — Flujos Principales del Sistema")
    add_p("Los diagramas de secuencia describen la interacción temporal entre los componentes del sistema QAMS para los tres flujos más críticos: autenticación segura, ejecución Fast Runner y registro de defecto.")

    add_h3("3.7.1 Flujo de Autenticación JWT + AES-256 (OE2)")
    add_p("Actor: Usuario (cualquier rol). Componentes involucrados: LoginComponent → AuthService(Angular) → EncryptionInterceptor → JwtInterceptor → AuthController(.NET) → AuthService(.NET) → User Entity (BCrypt verify) → JwtTokenGenerator.")
    headers_seq1 = ["Paso", "Origen", "Destino", "Acción / Datos"]
    rows_seq1 = [
        ["1", "Usuario", "LoginComponent", "Ingresa email + contraseña en el formulario reactivo"],
        ["2", "LoginComponent", "AuthService (Angular)", "Llama authService.login({email, password})"],
        ["3", "AuthService (Angular)", "EncryptionInterceptor", "POST /api/Auth/login con payload {email, password}"],
        ["4", "EncryptionInterceptor", "Backend API", "Cifra payload con AES-256: {data: 'encrypted_base64'}"],
        ["5", "AuthController (.NET)", "AuthService (.NET)", "Descifra payload, llama AuthService.LoginAsync(email, password)"],
        ["6", "AuthService (.NET)", "User Entity (PostgreSQL)", "SELECT user WHERE email=? AND IsDeleted=false. BCrypt.Verify(password, PasswordHash)"],
        ["7", "User Entity", "AuthService (.NET)", "Retorna User con sus Roles y Permisos"],
        ["8", "AuthService (.NET)", "JwtTokenGenerator", "Genera JWT con Claims: Id, Email, FullName, Permissions[], exp=8h"],
        ["9", "AuthController", "EncryptionInterceptor", "Retorna {data: encrypt({token, user})} HTTP 200"],
        ["10", "EncryptionInterceptor", "AuthService (Angular)", "Descifra respuesta: {token, user}"],
        ["11", "AuthService (Angular)", "LoginComponent", "Almacena token en memoria. Emite authState signal."],
        ["12", "LoginComponent", "Angular Router", "Navega a /dashboard"],
    ]
    add_custom_table(headers_seq1, rows_seq1, [0.4, 1.4, 1.4, 3.3])

    add_h3("3.7.2 Flujo Fast Runner — Ejecución de Caso de Prueba PASSED/FAILED (OE3)")
    add_p("Actor: QA Tester. Componentes: ExecutionDetailComponent → TestExecutionService (Angular, Signals) → EncryptionInterceptor → TestExecutionsController (.NET) → TestExecutionService (.NET) → ExecutionStepResult Entity (PostgreSQL).")
    headers_seq2 = ["Paso", "Origen", "Destino", "Acción / Datos"]
    rows_seq2 = [
        ["1", "QA Tester", "ExecutionDetailComponent", "Presiona tecla P (PASSED), F (FAILED) o B (BLOCKED) sobre el paso actual"],
        ["2", "ExecutionDetailComponent", "Signal local (Angular)", "stepStatusSignal.set(newStatus) → UI actualizada en microsegundos, sin Zone.js"],
        ["3", "ExecutionDetailComponent", "TestExecutionService", "Llama updateStepResult({executionId, stepOrder, statusCode})"],
        ["4", "TestExecutionService", "EncryptionInterceptor", "PATCH /api/TestExecutions/{id}/step-result con payload cifrado AES-256"],
        ["5", "TestExecutionsController", "TestExecutionService (.NET)", "Descifra payload, llama UpdateStepResultAsync(executionId, stepOrder, statusCode)"],
        ["6", "TestExecutionService (.NET)", "QamsDbContext", "UPDATE execution_step_results SET StatusId=?, UpdatedAt=UTC WHERE ExecutionId=? AND StepOrder=?"],
        ["7", "QamsDbContext", "IAuditable Interceptor", "SaveChangesAsync() inyecta UpdatedAt=UtcNow y UpdatedByUserId=<tester.Id>"],
        ["8", "TestExecutionsController", "EncryptionInterceptor", "Retorna HTTP 200 OK con stepResult actualizado (cifrado)"],
        ["9", "EncryptionInterceptor", "ExecutionDetailComponent", "Descifra respuesta, confirma persistencia"],
        ["10", "ExecutionDetailComponent", "Signal de progreso", "progressSignal.set(completedSteps/totalSteps) → barra de progreso actualizada"],
    ]
    add_custom_table(headers_seq2, rows_seq2, [0.4, 1.5, 1.5, 3.1])

    add_h3("3.7.3 Flujo de Registro de Defecto desde Ejecución Fallida (OE3 + OE4)")
    add_p("Actor: QA Tester. Componentes: ExecutionDetailComponent → DefectsService (Angular) → EncryptionInterceptor → DefectsController (.NET) → DefectService (.NET) → Defect Entity (PostgreSQL). La trazabilidad del defecto hacia el TestCase y la TestExecution garantiza el cumplimiento ISTQB de rastreabilidad bidireccional.")
    headers_seq3 = ["Paso", "Origen", "Destino", "Acción / Datos"]
    rows_seq3 = [
        ["1", "QA Tester", "ExecutionDetailComponent", "Hace clic en 'Crear Defecto' desde el paso FAILED de la ejecución"],
        ["2", "ExecutionDetailComponent", "DefectFormComponent", "Abre formulario pre-poblado con: TestCaseId, TestExecutionId, executionStep context"],
        ["3", "QA Tester", "DefectFormComponent", "Completa: Título, Descripción, SeverityId (Crítico/Alto/Medio/Bajo), StepsToReproduce, ActualResult"],
        ["4", "DefectFormComponent", "DefectsService (Angular)", "Llama defectsService.createDefect({...fields, testCaseId, testExecutionId, projectId})"],
        ["5", "DefectsService", "EncryptionInterceptor", "POST /api/Defects con payload completo cifrado AES-256"],
        ["6", "DefectsController (.NET)", "DefectService (.NET)", "Descifra payload, llama CreateDefectAsync(createDefectDto)"],
        ["7", "DefectService (.NET)", "QamsDbContext", "INSERT INTO defects con todas las FK de trazabilidad (ProjectId, TestCaseId, TestExecutionId)"],
        ["8", "QamsDbContext", "IAuditable Interceptor", "Inyecta CreatedAt=UtcNow, CreatedByUserId=<tester.Id>, IsDeleted=false"],
        ["9", "DefectsController", "EncryptionInterceptor", "HTTP 201 Created con Defecto persistido (Id, StatusId=Nuevo, Audit Trail)"],
        ["10", "DefectsService", "KanbanService (Angular)", "El defecto aparece automáticamente en la columna 'Nuevo' del Kanban del proyecto"],
        ["11", "ReportsController (.NET)", "RTM Matrix", "El endpoint GET /api/Reports/rtm-matrix retorna el defecto vinculado al requisito cuyo TC falló"],
    ]
    add_custom_table(headers_seq3, rows_seq3, [0.4, 1.5, 1.5, 3.1])
    add_p("Los tres flujos de secuencia documentados evidencian la implementación de la arquitectura de seguridad multicapa de QAMS: el cifrado AES-256 extremo a extremo (EncryptionInterceptor), la validación JWT en cada petición, el patrón de Audit Trail automático via QamsDbContext, y la trazabilidad bidireccional RTM que conecta el defecto con el caso de prueba, la ejecución y el requisito original — cumpliendo el Principio 1 de ISTQB CTFL v4.0 sobre la gestión sistemática de defectos.")

    doc.add_page_break()


    # =========================================================================
    # CAPÍTULO 4: DISEÑO Y ARQUITECTURA DEL SISTEMA
    # =========================================================================
    add_h1("Capítulo 4.- DISEÑO Y ARQUITECTURA DEL SISTEMA")
    add_p("El presente capítulo profundiza en las decisiones de diseño arquitectónico que sustentan el sistema QAMS. A diferencia del Capítulo 3 (centrado en los requisitos), este capítulo describe la estructura interna de los componentes: cómo se organizan las capas del backend, cómo el frontend gestiona el estado sin Zone.js, cómo el stack Docker asegura la portabilidad del sistema completo, y cómo el modelo relacional agrupa las entidades en dominios cohesivos. Cada decisión de diseño documentada aquí responde directamente a uno o más Requerimientos No Funcionales especificados en la Sección 3.3.")

    add_h2("4.1 Diagrama y Explicación Detallada de la Arquitectura del Backend (.NET 9)")
    add_p("El backend de QAMS está construido sobre .NET 9 siguiendo estrictamente la Clean Architecture propuesta por Robert C. Martin (2017). Esta arquitectura organiza el código en anillos concéntricos donde la regla fundamental es que las dependencias solo fluyen hacia el interior — nunca hacia afuera. La solución QAMS.sln contiene cinco proyectos organizados en cuatro capas con el siguiente grafo de dependencias: QAMS.Api → QAMS.Application → QAMS.Domain (sin dependencias externas). La QAMS.Infrastructure → QAMS.Domain (implementa sus interfaces). QAMS.Tests → QAMS.Application y QAMS.Domain (prueba sus reglas de negocio).")
    add_p("Middleware Pipeline de ASP.NET Core: Cada petición HTTP que llega al backend de QAMS atraviesa un pipeline de middleware configurado en Program.cs, ejecutándose en el siguiente orden estricto: (1) HTTPS Redirection — redirige tráfico HTTP a HTTPS. (2) CORS Policy — permite peticiones de los orígenes configurados (frontend Angular en puerto 4200 o 80). (3) Authentication — valida el JWT Bearer Token extraído del header Authorization. (4) Authorization — verifica los Claims del token contra los roles requeridos en el endpoint. (5) HasPermission Filter — filtro personalizado que valida el permiso atómico específico del endpoint (ej. DASHBOARD_VIEW). (6) Global Exception Handler — captura todas las excepciones no controladas y retorna respuestas HTTP estándar (400/401/403/404/500) sin exponer stack traces. (7) Controller Action — ejecuta el método del controlador que delega al servicio de aplicación.")
    add_p("Patrón Repository + Unit of Work: El acceso a datos está completamente abstraído mediante interfaces IRepository<T> e implementaciones concretas en QAMS.Infrastructure. Cada módulo funcional tiene su propio repositorio específico (ITestCaseRepository, IDefectRepository, IProjectRepository) con métodos especializados que van más allá del CRUD genérico (ej. FindWithDetailsAsync, CountOpenDefectsByProjectAsync). El patrón Unit of Work (IUnitOfWork) garantiza que múltiples operaciones de repositorio se ejecuten dentro de la misma transacción de base de datos, satisfaciendo el RNF-03 de Confiabilidad (transacciones ACID completas con rollback automático).")
    add_p("Patrón Decorator con Interceptores EF Core: El QamsDbContext sobreescribe SaveChangesAsync() como un Decorator que intercepta cada operación de persistencia para: (a) inyectar automáticamente los campos IAuditable (CreatedAt, CreatedByUserId, UpdatedAt, UpdatedByUserId) y (b) convertir las operaciones DELETE en actualizaciones ISoftDelete (IsDeleted=true, DeletedAt, DeletedByUserId). Este patrón elimina código boilerplate en todos los servicios y garantiza la consistencia del Audit Trail sin que el desarrollador deba recordarlo manualmente.")
    add_image_fig("figura6a_backend_detailed.png", "Figura 6A: Diagrama de Arquitectura Detallada del Backend QAMS — Pipeline de Middleware y Flujo de Datos entre Capas")

    add_h2("4.2 Diagrama y Explicación Detallada de la Arquitectura del Frontend (Angular 19)")
    add_p("El frontend de QAMS es una Single Page Application (SPA) construida con Angular 19 adoptando el nuevo modelo de Standalone Components introducido en Angular 14 y consolidado en Angular 17-19. Este modelo elimina el concepto de NgModule, permitiendo que cada componente declare explícitamente sus propias dependencias (imports) y sea cargado de forma lazy (perezosa) a través del Router. El resultado es un bundle inicial más pequeño (< 200 KB gzipped) y un tiempo de primera carga percibida (FCP) inferior a 1.5 segundos.")
    add_p("Gestión de Estado con Angular Signals (sin Zone.js): QAMS utiliza Angular Signals como mecanismo primario de reactividad, reemplazando el ciclo de detección de cambios de Zone.js. Un Signal es una celda reactiva primitiva que notifica a sus consumidores (Computed, Effect) cuando su valor cambia — y solo cuando cambia. En el Fast Runner, cada actualización de estado de un paso de ejecución (P/F/B) actualiza un Signal local, que propaga el cambio a la barra de progreso y al contador de completados en microsegundos, sin provocar una re-renderización del árbol completo de componentes. Esto es fundamental para cumplir el RNF-08 (latencia UI < 200ms).")
    add_p("Interceptores HTTP de Seguridad: El frontend de QAMS configura una cadena de interceptores HTTP en el proveedor de Angular que procesan cada petición antes de enviarla al backend: (1) JwtInterceptor — lee el token JWT del AuthService (almacenado en memoria, no en localStorage para prevenir XSS) y agrega el header 'Authorization: Bearer <token>' a cada petición. (2) EncryptionInterceptor — cifra el cuerpo de las peticiones POST/PUT/PATCH con AES-256-CBC usando la clave y el IV configurados en environment.ts, y descifra las respuestas entrantes del backend. Esta doble capa de seguridad satisface el RNF-01 y protege contra OWASP A2:Broken Authentication y A3:Sensitive Data Exposure.")
    add_p("Routing con Lazy Loading y AuthGuard: La configuración de rutas en app.routes.ts agrupa las rutas protegidas bajo el componente MainLayoutComponent, con el AuthGuard verificando la presencia de un token JWT válido antes de renderizar cualquier ruta de la aplicación. Las rutas de cada feature (TestCasesComponent, ReportsComponent, etc.) se cargan de forma diferida (loadComponent) solo cuando el usuario navega a ellas, optimizando el tiempo de carga inicial del dashboard.")
    add_image_fig("figura6b_frontend_detailed.png", "Figura 6B: Diagrama de Arquitectura del Frontend Angular 19 — Standalone Components, Signals y Cadena de Interceptores HTTP")

    add_h2("4.3 Diagramas de Despliegue de Servidores e Infraestructura de Hosting Frontend")
    add_p("La estrategia de despliegue de QAMS se basa en la contenedorización completa con Docker, que es el cumplimiento directo del OE5. El sistema utiliza Dockerfiles multi-stage que separan la fase de compilación de la imagen de producción final, reduciendo el tamaño de las imágenes al eliminar las herramientas de desarrollo (SDK de .NET, Node.js y dependencias npm) del runtime. El resultado son imágenes Alpine minimalistas para producción.")
    add_p("Dockerfile Backend — Etapas de construcción: Stage 1 (build): Utiliza la imagen mcr.microsoft.com/dotnet/sdk:9.0-alpine. Restaura los paquetes NuGet con dotnet restore, compilando solo las dependencias sin cambios para aprovechar la caché de Docker. Luego ejecuta dotnet publish -c Release -o /app/publish --no-restore generando los binarios optimizados de producción. Stage 2 (runtime): Utiliza la imagen minimalista mcr.microsoft.com/dotnet/aspnet:9.0-alpine. Copia solo los binarios publicados desde el stage anterior. Configura el usuario no privilegiado 'app' (USER app) para cumplir el principio de mínimo privilegio. Expone el puerto 8080. Tamaño final de la imagen: aproximadamente 120 MB.")
    add_p("Dockerfile Frontend — Etapas de construcción: Stage 1 (build): Utiliza node:20-alpine. Instala dependencias con npm ci (instalación reproducible basada en package-lock.json) y ejecuta ng build --configuration=production, generando los archivos estáticos optimizados en /dist/qams-web. Stage 2 (runtime): Utiliza nginx:alpine. Copia los archivos estáticos del stage anterior a /usr/share/nginx/html. Incluye un archivo nginx.conf personalizado que configura el try_files para el routing de SPA Angular (todas las rutas redirigen a index.html). Tamaño final: aproximadamente 25 MB.")
    add_p("Red y Volúmenes Docker Compose: Todos los servicios comparten la red interna qams-network (driver bridge), que aísla el tráfico interno de la red del host. Los servicios de base de datos y caché no exponen puertos al exterior en producción (solo internamente). Los volúmenes persistentes garantizan que los datos de PostgreSQL (postgres_data) y Redis (redis_data) sobrevivan a reinicios y actualizaciones de los contenedores. El volumen de uploads (qams-uploads) persiste los archivos de evidencia adjuntados por los testers durante las ejecuciones.")
    add_image_fig("figura8_deployment_docker.png", "Figura 8: Diagrama de Despliegue Docker Compose — Red Interna, Volúmenes y Healthchecks del Stack Completo QAMS")

    add_h2("4.4 Diagrama Entidad-Relación Global (ERD) — Agrupaciones por Dominio")
    add_p("El modelo de datos de QAMS organiza sus 23 entidades principales en cinco dominios cohesivos que reflejan las responsabilidades funcionales del sistema. Esta organización facilita la comprensión del modelo completo y ayuda a identificar las fronteras de transacciones y las reglas de cascada de eliminación lógica.")
    add_p("Dominio 1 — Seguridad y Acceso (4 entidades): User, Role, Permission, UserRole, RolePermission. Este dominio implementa el modelo RBAC completo. La relación M:N entre User y Role (a través de UserRole) y entre Role y Permission (a través de RolePermission) permite que un usuario tenga múltiples roles y cada rol acumule múltiples permisos atómicos. El claim 'Permissions' del JWT se construye aplanando esta jerarquía en tiempo de login para evitar consultas a la base de datos en cada petición.")
    add_p("Dominio 2 — Proyectos y Sistemas (3 entidades): SystemUnderTest, Project, ProjectTester. El SUT (Sistema Bajo Prueba) es la entidad raíz que agrupa todos los proyectos de QA sobre una misma aplicación. Un proyecto tiene exactamente un SUT (relación 1:N desde SUT) y puede tener múltiples testers asignados (tabla puente ProjectTester). Esta jerarquía permite filtrar el dashboard y los reportes por SUT, proyecto o tester específico.")
    add_p("Dominio 3 — Planificación de Pruebas (6 entidades): TestPlan, TestPlanCriteria, TestPlanMilestone, TestPlanRisk, TestSuite, Requirement, RequirementTestCase. El TestPlan es el artefacto central de gobernanza, que contiene los criterios ENTRY/EXIT, los hitos del proyecto, los riesgos identificados y referencia a las TestSuites. Los Requisitos se vinculan a los TestCases mediante la tabla puente RequirementTestCase (M:N), que es el fundamento de la Matriz RTM bidireccional.")
    add_p("Dominio 4 — Ejecución y Defectos (5 entidades): TestCase, TestStep, TestExecution, ExecutionStepResult, Evidence, Defect. Este dominio implementa el ciclo de vida operativo de las pruebas. Un TestCase puede tener múltiples TestExecutions a lo largo de los ciclos del proyecto. Cada TestExecution tiene N ExecutionStepResults (uno por cada TestStep). Las Evidencias se adjuntan a nivel de TestExecution. Los Defectos se crean con trazabilidad directa al TestCase y al TestExecution que los originó.")
    add_p("Dominio 5 — Revisiones, Exploración y Gestión Visual (6 entidades): ReviewSession, ReviewFinding, ReviewParticipant, ExploratorySession, ExploratoryFinding, KanbanBoard, KanbanColumn, KanbanTask. Este dominio implementa los módulos ISTQB de pruebas estáticas (Cap. 3) y gestión exploratoria SBTM (Cap. 4.4). El tablero Kanban vincula las KanbanTasks con los Defectos, proporcionando una vista visual del estado de resolución de incidentes.")
    add_image_fig("figura9_erd_overview.png", "Figura 9: Diagrama Entidad-Relación Global (ERD) del Sistema QAMS — 5 Dominios Cohesivos / PostgreSQL 16")

    add_h2("4.5 Diccionario de Datos Exhaustivo (Detalle Campo por Campo de las 32 Tablas)")
    add_p("A continuación se presenta el diccionario formal de datos campo por campo de todo el esquema de la base de datos:")

    def add_table_dict_full(tbl_name, desc, columns):
        add_h4(f"Tabla: {tbl_name}")
        add_p(f"Propósito Funcional: {desc}")
        headers_dict = ["Columna / Atributo", "Tipo de Dato", "Restricción (Constraint)", "Descripción Funcional del Campo"]
        add_custom_table(headers_dict, columns, [1.5, 1.2, 1.5, 2.3])

    # 1. USERS
    add_table_dict_full("users", "Almacena los usuarios del sistema, credenciales hasheadas y datos de perfil.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del usuario."],
        ["username", "VARCHAR(100)", "UK, NOT NULL", "Nombre de usuario único para autenticación."],
        ["email", "VARCHAR(256)", "UK, NOT NULL", "Correo electrónico institucional único."],
        ["password_hash", "VARCHAR(500)", "NOT NULL", "Hash seguro de contraseña generado con BCrypt (salt >= 12)."],
        ["full_name", "VARCHAR(200)", "NOT NULL", "Nombres y apellidos completos."],
        ["documento_identidad", "VARCHAR(50)", "NULL", "Número de cédula o documento de identidad oficial."],
        ["fecha_nacimiento", "DATE", "NULL", "Fecha de nacimiento para validación de mayoría de edad."],
        ["is_active", "BOOLEAN", "DEFAULT TRUE", "Bandera de habilitación operativa del usuario."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico (Soft-Delete)."],
        ["deleted_at", "TIMESTAMP", "NULL", "Marca de tiempo UTC en que se ejecutó el borrado lógico."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Marca de tiempo UTC de registro (Audit Trail)."],
        ["created_by_user_id", "UUID", "NULL", "Identificador del usuario que creó el registro."],
        ["updated_at", "TIMESTAMP", "NULL", "Marca de tiempo UTC de la última modificación."],
        ["updated_by_user_id", "UUID", "NULL", "Identificador del usuario que realizó la última edición."]
    ])

    # 2. ROLES
    add_table_dict_full("roles", "Define los roles de seguridad en el sistema (Admin, QA Lead, Tester, etc.).", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del rol."],
        ["name", "VARCHAR(100)", "UK, NOT NULL", "Nombre descriptivo único del rol."],
        ["description", "VARCHAR(255)", "NULL", "Descripción de las responsabilidades asignadas al rol."],
        ["is_active", "BOOLEAN", "DEFAULT TRUE", "Estado de activación del rol."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Marca de tiempo UTC de creación."]
    ])

    # 3. PERMISSIONS
    add_table_dict_full("permissions", "Catálogo de permisos atómicos del sistema (ej. 'projects.create', 'users.delete').", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del permiso."],
        ["code", "VARCHAR(100)", "UK, NOT NULL", "Código único del permiso verificado en endpoints."],
        ["description", "VARCHAR(255)", "NOT NULL", "Descripción en lenguaje natural de la acción permitida."],
        ["module", "VARCHAR(50)", "NOT NULL", "Módulo funcional al que pertenece (Users, Projects, Tests, etc.)."],
        ["is_active", "BOOLEAN", "DEFAULT TRUE", "Estado de habilitación del permiso."]
    ])

    # 4. USER_ROLES
    add_table_dict_full("user_roles", "Tabla puente relacional M:N entre Usuarios y Roles.", [
        ["user_id", "UUID", "PK, FK -> users(id)", "Identificador del usuario."],
        ["role_id", "UUID", "PK, FK -> roles(id)", "Identificador del rol asignado."],
        ["assigned_at", "TIMESTAMP", "NOT NULL", "Marca de tiempo UTC de asignación."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico de la asignación."]
    ])

    # 5. ROLE_PERMISSIONS
    add_table_dict_full("role_permissions", "Tabla puente relacional M:N entre Roles y Permisos atómicos.", [
        ["role_id", "UUID", "PK, FK -> roles(id)", "Identificador del rol."],
        ["permission_id", "UUID", "PK, FK -> permissions(id)", "Identificador del permiso vinculado."]
    ])

    # 6. SYSTEM_UNDER_TESTS (SUT)
    add_table_dict_full("system_under_tests", "Aplicaciones o sistemas de software que son objeto de pruebas.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del SUT."],
        ["name", "VARCHAR(200)", "NOT NULL", "Nombre formal de la aplicación bajo prueba."],
        ["description", "TEXT", "NULL", "Descripción de la arquitectura y alcance del SUT."],
        ["repository_url", "VARCHAR(500)", "NULL", "URL del repositorio de código fuente (Git/GitHub)."],
        ["version", "VARCHAR(50)", "NULL", "Versión base del SUT."],
        ["is_active", "BOOLEAN", "DEFAULT TRUE", "Estado del SUT."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico."]
    ])

    # 7. PROJECTS
    add_table_dict_full("projects", "Iniciativas y ciclos de prueba asociados a un SUT.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del proyecto."],
        ["system_under_test_id", "UUID", "FK -> system_under_tests(id)", "SUT asociado."],
        ["name", "VARCHAR(200)", "NOT NULL", "Nombre del proyecto de pruebas."],
        ["description", "TEXT", "NULL", "Alcance y objetivos de calidad del proyecto."],
        ["status", "INT", "NOT NULL", "Estado (1: Planificación, 2: En Pruebas, 3: Aprobado, 4: Cerrado)."],
        ["start_date", "DATE", "NOT NULL", "Fecha de inicio del ciclo."],
        ["end_date", "DATE", "NULL", "Fecha proyectada de cierre."],
        ["quality_gates", "JSONB", "NULL", "Configuración de umbrales cuantitativos de certificación."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico."]
    ])

    # 8. REQUIREMENTS
    add_table_dict_full("requirements", "Especificaciones de requisitos para la matriz de trazabilidad RTM.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del requerimiento."],
        ["project_id", "UUID", "FK -> projects(id)", "Proyecto al que pertenece el requisito."],
        ["code", "VARCHAR(50)", "NOT NULL", "Código único de negocio (ej. 'REQ-AUTH-001')."],
        ["title", "VARCHAR(300)", "NOT NULL", "Título conciso del requerimiento."],
        ["description", "TEXT", "NOT NULL", "Especificación técnica funcional o no funcional."],
        ["type", "VARCHAR(50)", "NOT NULL", "Tipo (Functional, Security, Performance, Usability)."],
        ["priority", "INT", "NOT NULL", "Prioridad (1: Baja, 2: Media, 3: Alta, 4: Crítica)."],
        ["status", "VARCHAR(50)", "NOT NULL", "Estado (Draft, Active, Verified, Deprecated)."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico."]
    ])

    # 9. REQUIREMENT_TEST_CASES (RTM)
    add_table_dict_full("requirement_test_cases", "Tabla puente de la Matriz de Trazabilidad de Requisitos (RTM M:N).", [
        ["requirement_id", "UUID", "PK, FK -> requirements(id)", "Identificador del requerimiento."],
        ["test_case_id", "UUID", "PK, FK -> test_cases(id)", "Identificador del caso de prueba vinculado."],
        ["mapped_at", "TIMESTAMP", "NOT NULL", "Marca de tiempo de establecimiento de la trazabilidad."]
    ])

    # 10. TEST_PLANS
    add_table_dict_full("test_plans", "Planes maestros de prueba según el estándar IEEE 829 / ISTQB.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del plan de pruebas."],
        ["project_id", "UUID", "FK -> projects(id)", "Proyecto contenedor."],
        ["title", "VARCHAR(300)", "NOT NULL", "Título formal del Plan de Pruebas."],
        ["scope", "TEXT", "NOT NULL", "Alcance detallado de lo que será probado."],
        ["out_of_scope", "TEXT", "NULL", "Funcionalidades explícitamente excluidas del testing."],
        ["strategy", "TEXT", "NOT NULL", "Estrategia y enfoque metodológico de pruebas."],
        ["risk_analysis", "TEXT", "NULL", "Análisis de riesgos de producto y mitigaciones."],
        ["environment_requirements", "TEXT", "NULL", "Requerimientos de hardware, software y red."],
        ["pass_rate_minimum", "DECIMAL(5,2)", "DEFAULT 95.00", "Umbral mínimo de aprobación para Quality Gate."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico."]
    ])

    # 11. TEST_SUITES
    add_table_dict_full("test_suites", "Agrupaciones lógicas y temáticas de casos de prueba.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la suite de pruebas."],
        ["project_id", "UUID", "FK -> projects(id)", "Proyecto al que pertenece."],
        ["name", "VARCHAR(200)", "NOT NULL", "Nombre de la suite (ej. 'Módulo de Pagos')."],
        ["description", "TEXT", "NULL", "Descripción del conjunto de pruebas."],
        ["execution_order", "INT", "DEFAULT 1", "Orden de ejecución secuencial sugerido."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico."]
    ])

    # 12. TEST_CASES
    add_table_dict_full("test_cases", "Casos de prueba diseñados clásicos y escenarios BDD.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del caso de prueba."],
        ["test_suite_id", "UUID", "FK -> test_suites(id)", "Suite contenedora."],
        ["code", "VARCHAR(50)", "NOT NULL", "Código de referencia del caso (ej. 'TC-AUTH-001')."],
        ["title", "VARCHAR(300)", "NOT NULL", "Título conciso del escenario de prueba."],
        ["preconditions", "TEXT", "NULL", "Precondiciones requeridas antes de la ejecución."],
        ["is_bdd", "BOOLEAN", "DEFAULT FALSE", "Indica si el caso utiliza formato BDD Gherkin."],
        ["bdd_scenario", "TEXT", "NULL", "Contenido del escenario en sintaxis Gherkin."],
        ["priority_id", "INT", "FK -> catalogs", "Prioridad de ejecución."],
        ["risk_score", "INT", "DEFAULT 1", "Puntuación de riesgo (RBT)."],
        ["is_certified", "BOOLEAN", "DEFAULT FALSE", "Indica si el caso ha sido formalmente certificado."],
        ["version_number", "INT", "DEFAULT 1", "Número de versión del caso de prueba."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico."]
    ])

    # 13. TEST_STEPS
    add_table_dict_full("test_steps", "Acciones atómicas y resultados esperados de un caso clásico.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del paso de prueba."],
        ["test_case_id", "UUID", "FK -> test_cases(id)", "Caso de prueba al que pertenece."],
        ["step_number", "INT", "NOT NULL", "Secuencia numérica ordinal del paso (1, 2, 3...)."],
        ["action", "TEXT", "NOT NULL", "Acción detallada que debe ejecutar el tester."],
        ["expected_result", "TEXT", "NOT NULL", "Resultado esperado del comportamiento del sistema."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Bandera de borrado lógico."]
    ])

    # 14. TEST_EXECUTIONS
    add_table_dict_full("test_executions", "Instancias temporales de corridas de prueba ejecutadas.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la corrida de prueba."],
        ["test_case_id", "UUID", "FK -> test_cases(id)", "Caso de prueba ejecutado."],
        ["executed_by_user_id", "UUID", "FK -> users(id)", "Tester que realizó la ejecución."],
        ["status", "VARCHAR(50)", "NOT NULL", "Resultado global (PASSED, FAILED, BLOCKED)."],
        ["execution_time_seconds", "INT", "DEFAULT 0", "Tiempo total transcurrido en segundos."],
        ["executed_at", "TIMESTAMP", "NOT NULL", "Fecha y hora exacta de la corrida."],
        ["notes", "TEXT", "NULL", "Observaciones y notas asentadas por el ejecutor."]
    ])

    # 15. EXECUTION_STEP_RESULTS
    add_table_dict_full("execution_step_results", "Resultados atómicos obtenidos en cada paso de una corrida.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del resultado del paso."],
        ["test_execution_id", "UUID", "FK -> test_executions(id)", "Corrida contenedora."],
        ["test_step_id", "UUID", "FK -> test_steps(id)", "Paso atómico evaluado."],
        ["status", "VARCHAR(50)", "NOT NULL", "Estado del paso (PASSED, FAILED, BLOCKED)."],
        ["observation", "TEXT", "NULL", "Observación específica sobre el comportamiento observado."]
    ])

    # 16. EVIDENCES
    add_table_dict_full("evidences", "Evidencias digitales (capturas, logs, videos) adjuntas a una ejecución.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal de la evidencia."],
        ["test_execution_id", "UUID", "FK -> test_executions(id)", "Ejecución asociada."],
        ["file_path", "VARCHAR(500)", "NOT NULL", "Ruta física o URI de almacenamiento del archivo."],
        ["file_name", "VARCHAR(255)", "NOT NULL", "Nombre original del archivo subido."],
        ["file_type", "VARCHAR(50)", "NOT NULL", "Tipo MIME o extensión (PNG, JPEG, PDF, LOG)."],
        ["file_size_bytes", "BIGINT", "NOT NULL", "Tamaño del archivo en bytes."],
        ["uploaded_at", "TIMESTAMP", "NOT NULL", "Marca de tiempo UTC de carga."]
    ])

    # 17. DEFECTS
    add_table_dict_full("defects", "Registro y ciclo de vida completo de defectos e incidentes.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único universal del defecto."],
        ["test_case_id", "UUID", "FK -> test_cases(id)", "Caso de prueba que reveló el defecto."],
        ["test_execution_id", "UUID", "FK -> test_executions(id)", "Corrida específica donde falló."],
        ["code", "VARCHAR(50)", "NOT NULL", "Código único del ticket (ej. 'BUG-001')."],
        ["title", "VARCHAR(300)", "NOT NULL", "Resumen claro y conciso del defecto."],
        ["steps_to_reproduce", "TEXT", "NOT NULL", "Pasos detallados para reproducir el fallo."],
        ["expected_result", "TEXT", "NOT NULL", "Comportamiento que debió haber ocurrido."],
        ["actual_result", "TEXT", "NOT NULL", "Comportamiento anómalo observado."],
        ["severity_id", "INT", "FK -> catalogs", "Severidad (Blocker, Critical, Major, Minor)."],
        ["priority_id", "INT", "FK -> catalogs", "Prioridad de corrección para el equipo de desarrollo."],
        ["status", "VARCHAR(50)", "NOT NULL", "Estado (New, Assigned, Resolved, Closed, Reopened)."],
        ["assigned_to_user_id", "UUID", "NULL, FK -> users(id)", "Desarrollador asignado a la corrección."],
        ["reported_by_user_id", "UUID", "FK -> users(id)", "Tester que reportó el defecto."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de reporte."]
    ])

    # 18. REVIEW_SESSIONS (STATIC TESTING)
    add_table_dict_full("review_sessions", "Sesiones de revisión estática formal (Inspecciones y Walkthroughs).", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la sesión de revisión."],
        ["project_id", "UUID", "FK -> projects(id)", "Proyecto asociado."],
        ["title", "VARCHAR(300)", "NOT NULL", "Título de la sesión de revisión estática."],
        ["review_type", "VARCHAR(50)", "NOT NULL", "Tipo (Informal, Walkthrough, Technical, Inspection)."],
        ["artifact_under_review", "TEXT", "NOT NULL", "Referencia o contenido del documento bajo revisión."],
        ["moderator_id", "UUID", "FK -> users(id)", "Moderador que conduce la sesión formal."],
        ["author_id", "UUID", "FK -> users(id)", "Autor del artefacto que recibe las observaciones."],
        ["status", "VARCHAR(50)", "NOT NULL", "Estado (Planned, InProgress, Completed, Rejected)."]
    ])

    # 19. REVIEW_FINDINGS
    add_table_dict_full("review_findings", "Hallazgos y defectos detectados durante las pruebas estáticas.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del hallazgo."],
        ["review_session_id", "UUID", "FK -> review_sessions(id)", "Sesión de revisión contenedora."],
        ["finding_type", "VARCHAR(50)", "NOT NULL", "Tipo de hallazgo (Ambiguity, Inconsistency, Omission)."],
        ["finding_severity", "VARCHAR(50)", "NOT NULL", "Severidad del hallazgo."],
        ["location_reference", "VARCHAR(200)", "NOT NULL", "Sección o línea del documento observado."],
        ["description", "TEXT", "NOT NULL", "Descripción del defecto estático detectado."]
    ])

    # 20. EXPLORATORY_SESSIONS (SBTM)
    add_table_dict_full("exploratory_sessions", "Sesiones de pruebas exploratorias basadas en cartas de misión.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la sesión exploratoria."],
        ["project_id", "UUID", "FK -> projects(id)", "Proyecto asociado."],
        ["charter_title", "VARCHAR(300)", "NOT NULL", "Título de la carta de misión."],
        ["charter_mission", "TEXT", "NOT NULL", "Declaración de objetivos y alcance de exploración."],
        ["duration_minutes", "INT", "NOT NULL", "Duración acotada en tiempo (Time-box)."],
        ["executed_by_user_id", "UUID", "FK -> users(id)", "Tester explorador."],
        ["status", "VARCHAR(50)", "NOT NULL", "Estado de la sesión (Active, Completed)."]
    ])

    # 21. KANBAN_TASKS
    add_table_dict_full("kanban_tasks", "Tareas de trabajo ágil vinculadas a casos o actividades de QA.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la tarjeta Kanban."],
        ["kanban_column_id", "UUID", "FK -> kanban_columns(id)", "Columna donde reside la tarjeta."],
        ["test_case_id", "UUID", "NULL, FK -> test_cases(id)", "Caso de prueba vinculado."],
        ["title", "VARCHAR(300)", "NOT NULL", "Título de la tarea."],
        ["description", "TEXT", "NULL", "Detalle de la actividad."],
        ["assigned_to_user_id", "UUID", "NULL, FK -> users(id)", "Usuario asignado a la tarea."],
        ["task_order", "INT", "NOT NULL", "Posición ordinal en la columna."]
    ])

    # 22. CATALOGS
    add_table_dict_full("catalogs", "Catálogo maestro parametrizable para clasificaciones del sistema.", [
        ["id", "INT", "PK, NOT NULL", "Identificador numérico del ítem de catálogo."],
        ["catalog_type", "VARCHAR(100)", "NOT NULL", "Tipo (TestType, Priority, Severity, FileType)."],
        ["code", "VARCHAR(50)", "NOT NULL", "Código programático del catálogo."],
        ["name", "VARCHAR(100)", "NOT NULL", "Nombre legible del ítem."],
        ["description", "VARCHAR(255)", "NULL", "Descripción del ítem."],
        ["display_order", "INT", "DEFAULT 1", "Orden de presentación visual."],
        ["is_active", "BOOLEAN", "DEFAULT TRUE", "Bandera de habilitación."]
    ])

    # ── Diccionario de Datos — Tablas restantes (9-23) ──────────────────────

    # 9. TEST_PLANS
    add_table_dict_full("test_plans", "Planes de prueba formales (IEEE 829). Contienen estrategia, niveles, criterios ENTRY/EXIT y umbrales de Quality Gate.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del plan de prueba."],
        ["project_id", "UUID", "FK → projects(id)", "Proyecto propietario del plan."],
        ["name", "VARCHAR(200)", "NOT NULL", "Nombre descriptivo del plan."],
        ["objectives", "TEXT", "NOT NULL", "Objetivos formales del plan de prueba."],
        ["scope", "TEXT", "NULL", "Alcance del plan: qué se probará y qué no."],
        ["test_strategy_id", "INT", "FK → catalogs", "Estrategia: Regresión, Humo, Integración, etc."],
        ["test_level_id", "INT", "FK → catalogs", "Nivel: Unitario, Integración, Sistema, Aceptación."],
        ["test_plan_type_id", "INT", "FK → catalogs", "Tipo: Maestro, Fase, Iteración."],
        ["status_id", "INT", "FK → catalogs", "Estado: Borrador, Revisión, Aprobado, Activo, Cerrado."],
        ["quality_gate_pass_rate", "DECIMAL(5,2)", "NULL", "Umbral mínimo de Pass Rate para certificar (ej. 90.00)."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha y hora de creación en UTC."],
        ["created_by_user_id", "UUID", "NULL", "Usuario que creó el plan."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica del plan."]
    ])

    # 10. TEST_PLAN_CRITERIA
    add_table_dict_full("test_plan_criteria", "Criterios de entrada (ENTRY) y salida (EXIT) del plan de prueba según ISTQB Cap. 5.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del criterio."],
        ["test_plan_id", "UUID", "FK → test_plans(id)", "Plan propietario."],
        ["criteria_type", "VARCHAR(10)", "NOT NULL CHECK IN ('ENTRY','EXIT')", "Tipo de criterio: ENTRY o EXIT."],
        ["description", "TEXT", "NOT NULL", "Descripción del criterio verificable."],
        ["is_met", "BOOLEAN", "DEFAULT FALSE", "Indica si el criterio ha sido cumplido y verificado."],
        ["met_at", "TIMESTAMP", "NULL", "Fecha y hora en que el criterio fue marcado como cumplido."]
    ])

    # 11. TEST_PLAN_MILESTONES
    add_table_dict_full("test_plan_milestones", "Hitos y fechas clave del plan de prueba para monitoreo de progreso.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del hito."],
        ["test_plan_id", "UUID", "FK → test_plans(id)", "Plan propietario."],
        ["name", "VARCHAR(200)", "NOT NULL", "Nombre del hito (ej. 'Fin de pruebas de integración')."],
        ["due_date", "DATE", "NOT NULL", "Fecha objetivo del hito."],
        ["completed_date", "DATE", "NULL", "Fecha real de completación del hito."],
        ["is_completed", "BOOLEAN", "DEFAULT FALSE", "Estado de completación del hito."]
    ])

    # 12. TEST_PLAN_RISKS
    add_table_dict_full("test_plan_risks", "Riesgos identificados en el plan de prueba con su nivel y plan de mitigación (Risk-Based Testing).", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del riesgo."],
        ["test_plan_id", "UUID", "FK → test_plans(id)", "Plan propietario."],
        ["description", "TEXT", "NOT NULL", "Descripción del riesgo identificado."],
        ["likelihood", "INT", "NOT NULL CHECK(1-5)", "Probabilidad de ocurrencia (1=Muy baja, 5=Muy alta)."],
        ["impact", "INT", "NOT NULL CHECK(1-5)", "Impacto si el riesgo se materializa."],
        ["risk_score", "INT", "COMPUTED: likelihood*impact", "Score de riesgo del 1 al 25."],
        ["mitigation", "TEXT", "NULL", "Estrategia de mitigación o contingencia definida."]
    ])

    # 13. TEST_SUITES
    add_table_dict_full("test_suites", "Suites o escenarios de prueba que agrupan casos de prueba relacionados dentro de un proyecto.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la suite."],
        ["project_id", "UUID", "FK → projects(id)", "Proyecto propietario."],
        ["name", "VARCHAR(200)", "NOT NULL", "Nombre de la suite de prueba."],
        ["description", "TEXT", "NULL", "Descripción del propósito de la suite."],
        ["test_level_id", "INT", "FK → catalogs", "Nivel de prueba (Unitario, Integración, Sistema, Aceptación)."],
        ["test_type_id", "INT", "FK → catalogs", "Tipo de prueba (Funcional, No Funcional, Regresión, etc.)."],
        ["status_id", "INT", "FK → catalogs", "Estado de la suite."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de creación UTC."]
    ])

    # 14. TEST_CASES
    add_table_dict_full("test_cases", "Casos de prueba unitarios con soporte para modo clásico (pasos estructurados) y BDD (Gherkin Given-When-Then).", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del caso de prueba."],
        ["project_id", "UUID", "FK → projects(id)", "Proyecto propietario."],
        ["test_suite_id", "UUID", "FK → test_suites(id)", "Suite contenedora."],
        ["parent_test_case_id", "UUID", "FK → test_cases(id), NULL", "Auto-referencia para sub-casos (jerarquía)."],
        ["title", "VARCHAR(300)", "NOT NULL", "Título descriptivo del caso."],
        ["preconditions", "TEXT", "NULL", "Estado previo requerido para ejecutar el caso."],
        ["expected_result", "TEXT", "NULL", "Resultado esperado a nivel general."],
        ["postconditions", "TEXT", "NULL", "Estado del sistema después de la ejecución."],
        ["priority_id", "INT", "FK → catalogs", "Prioridad: Crítica, Alta, Media, Baja."],
        ["test_type_id", "INT", "FK → catalogs", "Tipo de prueba."],
        ["design_technique_id", "INT", "FK → catalogs", "Técnica de diseño: EP, BVA, Tabla de Decisión, etc."],
        ["is_bdd", "BOOLEAN", "DEFAULT FALSE", "Indica si el caso usa sintaxis BDD/Gherkin."],
        ["bdd_scenario", "TEXT", "NULL", "Escenario Gherkin completo (Given-When-Then) si is_bdd=true."],
        ["impact_level", "INT", "DEFAULT 3, CHECK(1-5)", "Nivel de impacto para Risk-Based Testing (RBT)."],
        ["likelihood_level", "INT", "DEFAULT 3, CHECK(1-5)", "Probabilidad de fallo para RBT."],
        ["risk_score", "INT", "COMPUTED: impact*likelihood", "Score de riesgo RBT (1-25)."],
        ["version_number", "INT", "DEFAULT 1", "Número de versión del caso (autoincremental)."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de creación UTC."],
        ["created_by_user_id", "UUID", "NULL", "Usuario creador."]
    ])

    # 15. TEST_STEPS
    add_table_dict_full("test_steps", "Pasos atómicos de ejecución de un caso de prueba en modo clásico (aplica cuando is_bdd=false).", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del paso."],
        ["test_case_id", "UUID", "FK → test_cases(id)", "Caso de prueba propietario."],
        ["step_order", "INT", "NOT NULL", "Orden de ejecución del paso (1, 2, 3, ...)."],
        ["action", "TEXT", "NOT NULL", "Acción que debe realizar el tester en este paso."],
        ["expected_result", "TEXT", "NOT NULL", "Resultado esperado al completar esta acción específica."],
        ["test_data", "TEXT", "NULL", "Datos de prueba necesarios para este paso (valores, credenciales, etc.)."]
    ])

    # 16. TEST_EXECUTIONS
    add_table_dict_full("test_executions", "Registro de cada ejecución de un caso de prueba por un tester. Soporta múltiples ciclos de ejecución.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la ejecución."],
        ["test_case_id", "UUID", "FK → test_cases(id)", "Caso de prueba ejecutado."],
        ["test_plan_id", "UUID", "FK → test_plans(id), NULL", "Plan de prueba en el que se realiza la ejecución."],
        ["tester_id", "UUID", "FK → users(id)", "Usuario que ejecuta la prueba."],
        ["status_id", "INT", "FK → catalogs (ExecutionStatus)", "Estado: PENDING, IN_PROGRESS, PASSED, FAILED, BLOCKED."],
        ["notes", "TEXT", "NULL", "Observaciones y comentarios del tester durante la ejecución."],
        ["actual_time_hours", "DECIMAL(6,2)", "DEFAULT 0", "Tiempo real invertido en la ejecución (en horas)."],
        ["execution_date", "TIMESTAMP", "NOT NULL", "Fecha y hora UTC de inicio de la ejecución."],
        ["cycle_number", "INT", "DEFAULT 1", "Número de ciclo de prueba (para regresiones y re-ejecuciones)."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de creación UTC."],
        ["created_by_user_id", "UUID", "NULL", "Usuario que inició la ejecución (coincide con tester_id)."]
    ])

    # 17. EXECUTION_STEP_RESULTS
    add_table_dict_full("execution_step_results", "Resultado individual por cada paso de prueba dentro de una ejecución. Permite rastrear qué paso específico falló.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del resultado de paso."],
        ["test_execution_id", "UUID", "FK → test_executions(id)", "Ejecución propietaria."],
        ["step_order", "INT", "NOT NULL", "Número de orden del paso evaluado."],
        ["status_id", "INT", "FK → catalogs (ExecutionStatus)", "Estado del paso: PASSED, FAILED, BLOCKED, PENDING."],
        ["notes", "TEXT", "NULL", "Observación del tester sobre este paso específico."],
        ["executed_at", "TIMESTAMP", "NULL", "Marca de tiempo UTC exacta en que se evaluó el paso."]
    ])

    # 18. EVIDENCES
    add_table_dict_full("evidences", "Archivos adjuntos (screenshots, logs, videos) vinculados a una ejecución de prueba como soporte documental.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la evidencia."],
        ["test_execution_id", "UUID", "FK → test_executions(id)", "Ejecución a la que pertenece la evidencia."],
        ["evidence_type_id", "INT", "FK → catalogs (EvidenceType)", "Tipo: Screenshot, Log de Errores, Video, Archivo de Datos."],
        ["file_name", "VARCHAR(255)", "NOT NULL", "Nombre original del archivo de evidencia."],
        ["file_url", "VARCHAR(1000)", "NOT NULL", "Ruta relativa en el volumen Docker /app/wwwroot/uploads/."],
        ["file_size_bytes", "BIGINT", "NULL", "Tamaño del archivo en bytes."],
        ["description", "VARCHAR(500)", "NULL", "Descripción de qué muestra o evidencia esta la captura."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha y hora de adjunto UTC."],
        ["created_by_user_id", "UUID", "NULL", "Usuario que adjuntó la evidencia."]
    ])

    # 19. DEFECTS
    add_table_dict_full("defects", "Defectos o bugs detectados durante las ejecuciones. Ciclo de vida completo con trazabilidad bidireccional a TestCase y TestExecution.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del defecto."],
        ["project_id", "UUID", "FK → projects(id)", "Proyecto propietario."],
        ["test_case_id", "UUID", "FK → test_cases(id), NULL", "Caso de prueba que originó el defecto."],
        ["test_execution_id", "UUID", "FK → test_executions(id), NULL", "Ejecución donde se detectó el defecto."],
        ["title", "VARCHAR(300)", "NOT NULL", "Título descriptivo del defecto."],
        ["description", "TEXT", "NULL", "Descripción detallada del comportamiento erróneo."],
        ["severity_id", "INT", "FK → catalogs (DefectSeverity)", "Severidad: Crítico, Alto, Medio, Bajo, Informativo."],
        ["priority_id", "INT", "FK → catalogs (DefectPriority), NULL", "Prioridad de resolución: Alta, Media, Baja."],
        ["status_id", "INT", "FK → catalogs (DefectStatus)", "Estado: Nuevo, Asignado, En Progreso, Resuelto, Verificado, Cerrado."],
        ["steps_to_reproduce", "TEXT", "NULL", "Pasos detallados para reproducir el defecto."],
        ["expected_result", "TEXT", "NULL", "Comportamiento esperado según el caso de prueba."],
        ["actual_result", "TEXT", "NULL", "Comportamiento observado y defectuoso."],
        ["environment", "VARCHAR(100)", "NULL", "Ambiente donde se detectó (Dev, QA, Staging)."],
        ["assigned_to_user_id", "UUID", "FK → users(id), NULL", "Developer asignado para la corrección."],
        ["reported_by_user_id", "UUID", "FK → users(id), NULL", "Tester que reportó el defecto."],
        ["resolved_at", "TIMESTAMP", "NULL", "Fecha y hora UTC en que el developer marcó como resuelto."],
        ["closed_at", "TIMESTAMP", "NULL", "Fecha y hora UTC en que el QA verificó y cerró el defecto."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de reporte UTC."]
    ])

    # 20. REVIEW_SESSIONS
    add_table_dict_full("review_sessions", "Sesiones de revisión estática (Walkthrough, Inspección, Revisión Técnica) según ISTQB Cap. 3.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la sesión de revisión."],
        ["project_id", "UUID", "FK → projects(id)", "Proyecto propietario."],
        ["title", "VARCHAR(300)", "NOT NULL", "Nombre o propósito de la sesión de revisión."],
        ["artifact_under_review", "TEXT", "NULL", "Descripción del artefacto revisado (requisito, arquitectura, código)."],
        ["review_type_id", "INT", "FK → catalogs (ReviewType)", "Tipo: Walkthrough, Inspección, Revisión Técnica, Revisión Informal."],
        ["review_status_id", "INT", "FK → catalogs (ReviewStatus)", "Estado: Planificada, En Progreso, Completada, Cancelada."],
        ["moderator_id", "UUID", "FK → users(id), NULL", "QA Lead que modera la sesión."],
        ["verdict", "VARCHAR(50)", "NULL", "Dictamen final: APROBADO, APROBADO CON CONDICIONES, RECHAZADO."],
        ["scheduled_at", "TIMESTAMP", "NULL", "Fecha y hora planificada de la sesión."],
        ["completed_at", "TIMESTAMP", "NULL", "Fecha y hora real de finalización."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de creación UTC."]
    ])

    # 21. REVIEW_FINDINGS
    add_table_dict_full("review_findings", "Hallazgos registrados por los participantes durante una sesión de revisión estática.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del hallazgo."],
        ["review_session_id", "UUID", "FK → review_sessions(id)", "Sesión de revisión propietaria."],
        ["reported_by_user_id", "UUID", "FK → users(id), NULL", "Revisor que identificó el hallazgo."],
        ["finding_type_id", "INT", "FK → catalogs (FindingType)", "Tipo: Defecto, Mejora, Pregunta, Observación."],
        ["finding_severity_id", "INT", "FK → catalogs (FindingSeverity)", "Severidad: Mayor, Menor, Informativo."],
        ["description", "TEXT", "NOT NULL", "Descripción detallada del hallazgo."],
        ["reference_location", "VARCHAR(300)", "NULL", "Referencia a la sección/línea del artefacto donde se encontró."],
        ["is_resolved", "BOOLEAN", "DEFAULT FALSE", "Indica si el hallazgo fue abordado por el autor."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de registro del hallazgo UTC."]
    ])

    # 22. EXPLORATORY_SESSIONS
    add_table_dict_full("exploratory_sessions", "Sesiones de prueba exploratoria SBTM con carta de misión (charter), time-box y hallazgos.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la sesión exploratoria."],
        ["project_id", "UUID", "FK → projects(id)", "Proyecto propietario."],
        ["tester_id", "UUID", "FK → users(id)", "Tester que conduce la sesión exploratoria."],
        ["charter", "TEXT", "NOT NULL", "Carta de misión: área a explorar, objetivo y estrategia de la sesión."],
        ["duration_minutes", "INT", "NOT NULL", "Time-box planificado de la sesión en minutos."],
        ["status_id", "INT", "FK → catalogs", "Estado: Planificada, En Progreso, Completada."],
        ["start_time", "TIMESTAMP", "NULL", "Marca de tiempo UTC de inicio real de la sesión."],
        ["end_time", "TIMESTAMP", "NULL", "Marca de tiempo UTC de finalización real."],
        ["notes", "TEXT", "NULL", "Notas generales y observaciones durante la exploración."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de creación UTC."]
    ])

    # 23. KANBAN_BOARDS / KANBAN_COLUMNS / KANBAN_TASKS
    add_table_dict_full("kanban_boards", "Tableros Kanban por proyecto para gestión visual del flujo de trabajo de defectos y tareas.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único del tablero."],
        ["project_id", "UUID", "FK → projects(id)", "Proyecto propietario del tablero."],
        ["name", "VARCHAR(200)", "NOT NULL", "Nombre del tablero (ej. 'Sprint 3 QA Board')."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."]
    ])

    add_table_dict_full("kanban_columns", "Columnas del tablero Kanban que representan los estados del flujo de trabajo.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la columna."],
        ["kanban_board_id", "UUID", "FK → kanban_boards(id)", "Tablero propietario."],
        ["name", "VARCHAR(100)", "NOT NULL", "Nombre de la columna (Nuevo, Asignado, En Progreso, Resuelto, Cerrado)."],
        ["column_order", "INT", "NOT NULL", "Orden de visualización de la columna (izquierda a derecha)."],
        ["wip_limit", "INT", "NULL", "Límite WIP (Work In Progress) de la columna (NULL = sin límite)."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."]
    ])

    add_table_dict_full("kanban_tasks", "Tareas y tarjetas del tablero Kanban, vinculadas opcionalmente a un defecto del sistema.", [
        ["id", "UUID", "PK, NOT NULL", "Identificador único de la tarea Kanban."],
        ["kanban_column_id", "UUID", "FK → kanban_columns(id)", "Columna actual donde está la tarea."],
        ["defect_id", "UUID", "FK → defects(id), NULL", "Defecto asociado a esta tarea (trazabilidad)."],
        ["title", "VARCHAR(300)", "NOT NULL", "Título descriptivo de la tarea."],
        ["description", "TEXT", "NULL", "Descripción detallada de la tarea."],
        ["assigned_to_id", "UUID", "FK → users(id), NULL", "Miembro del equipo asignado a esta tarea."],
        ["due_date", "DATE", "NULL", "Fecha límite de completación de la tarea."],
        ["card_order", "INT", "DEFAULT 1", "Orden visual de la tarjeta dentro de la columna."],
        ["is_deleted", "BOOLEAN", "DEFAULT FALSE", "Eliminación lógica."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Fecha de creación UTC."]
    ])

    doc.add_page_break()

    # =========================================================================
    # CAPÍTULO 5: DESARROLLO E IMPLEMENTACIÓN TÉCNICA
    # =========================================================================
    add_h1("Capítulo 5.- DESARROLLO E IMPLEMENTACIÓN TÉCNICA")
    add_p("El presente capítulo documenta la implementación técnica del sistema QAMS en sus tres componentes principales: el backend en .NET 9 con Clean Architecture, el frontend en Angular 19 con Signals y el subsistema de seguridad AES-256. Para cada componente se presenta el contexto de la decisión de implementación, el código fuente relevante con anotaciones explicativas y la relación con los Requerimientos No Funcionales que satisface.")

    add_h2("5.1 Implementación Backend: Clean Architecture, Interceptores y Soft-Delete")
    add_p("El código C# del backend de QAMS fue escrito en .NET 9 con C# 13, aprovechando las características modernas del lenguaje como primary constructors, pattern matching avanzado y Span<T> para manipulación de bytes eficiente en el módulo de cifrado. El QamsDbContext es el componente crítico de infraestructura que implementa los dos patrones de gobernanza de datos del sistema: Audit Trail automático (IAuditable) y Soft Delete (ISoftDelete). El siguiente extracto muestra la sobreescritura del método SaveChangesAsync(), que es invocado por todos los repositorios del sistema a través del patrón Unit of Work:")
    add_p("Código 1: QamsDbContext.SaveChangesAsync — Interceptor de Auditoría y Soft Delete (C# / EF Core 9)")
    add_p("public override async Task<int> SaveChangesAsync(CancellationToken ct = default) {\n  var userId = _currentUserService.UserId; var now = DateTime.UtcNow;\n  foreach (var e in ChangeTracker.Entries<IAuditable>()) {\n    if (e.State == EntityState.Added) { e.Entity.CreatedAt = now; e.Entity.CreatedByUserId = userId; }\n    else if (e.State == EntityState.Modified) { e.Entity.UpdatedAt = now; e.Entity.UpdatedByUserId = userId; }\n  }\n  foreach (var e in ChangeTracker.Entries<ISoftDelete>()) {\n    if (e.State == EntityState.Deleted) {\n      e.State = EntityState.Modified; e.Entity.IsDeleted = true;\n      e.Entity.DeletedAt = now; e.Entity.DeletedByUserId = userId;\n    }\n  }\n  return await base.SaveChangesAsync(ct);\n}")
    add_p("Análisis del código: El método iterates sobre dos conjuntos de ChangeTracker. Para el conjunto IAuditable, distingue entre entidades Added (nuevas) y Modified (actualizadas) para inyectar los timestamps y userIds correspondientes. Para el conjunto ISoftDelete, intercepta las operaciones Deleted y las convierte a Modified, seteando las banderas de eliminación lógica. Este mecanismo es transparente para todos los 19 servicios de aplicación — ninguno necesita recordar establecer estos campos manualmente, eliminando una categoría completa de bugs de omisión.")
    add_p("Implementación del Servicio de Autenticación (AuthService.cs): El flujo de autenticación implementa el algoritmo BCrypt para verificación de contraseñas (BCrypt.Verify(plainPassword, user.PasswordHash)) y genera el JWT con los Claims del usuario usando JwtSecurityTokenHandler. Los Claims del token incluyen: NameIdentifier (userId), Email, sub (username), Roles (array), y Permissions (array aplanado de todos los permisos del rol). Este último es crítico para el funcionamiento del filtro HasPermission sin consultas adicionales a la base de datos por petición.")

    add_h3("Tabla 11. Endpoints REST del Backend QAMS — Catálogo por Controlador")
    headers_api = ["Controlador (.cs)", "Método HTTP", "Ruta", "Permiso RBAC", "Descripción"]
    rows_api = [
        ["AuthController", "POST", "/api/Auth/login", "Pública", "Autenticación y generación de JWT."],
        ["AuthController", "POST", "/api/Auth/register", "Pública", "Registro de nuevo usuario."],
        ["AuthController", "POST", "/api/Auth/forgot-password", "Pública", "Solicitud de recuperación de contraseña vía email."],
        ["UsersController", "GET", "/api/Users", "USERS_VIEW", "Lista paginada de usuarios del sistema."],
        ["UsersController", "POST", "/api/Users", "USERS_VIEW", "Creación de usuario con asignación de rol."],
        ["UsersController", "PATCH", "/api/Users/{id}", "USERS_VIEW", "Actualización parcial del usuario."],
        ["SystemsUnderTestController", "GET", "/api/SystemsUnderTest", "SUT_VIEW", "Lista de SUTs activos."],
        ["SystemsUnderTestController", "POST", "/api/SystemsUnderTest", "SUT_VIEW", "Registro de nuevo SUT."],
        ["ProjectsController", "GET", "/api/Projects", "PROJECTS_VIEW", "Lista de proyectos del usuario."],
        ["ProjectsController", "POST", "/api/Projects", "PROJECTS_VIEW", "Creación de proyecto vinculado a SUT."],
        ["RequirementsController", "GET", "/api/Requirements", "PROJECTS_VIEW", "Lista de requisitos del proyecto."],
        ["RequirementsController", "POST", "/api/Requirements", "PROJECTS_VIEW", "Registro de requisito funcional/no funcional."],
        ["TestPlansController", "GET", "/api/TestPlans", "TEST_CASES_VIEW", "Lista de planes de prueba."],
        ["TestPlansController", "POST", "/api/TestPlans", "TEST_CASES_VIEW", "Creación de plan de prueba IEEE 829."],
        ["TestCasesController", "GET", "/api/TestCases", "TEST_CASES_VIEW", "Lista de casos de prueba."],
        ["TestCasesController", "POST", "/api/TestCases", "TEST_CASES_VIEW", "Creación de caso clásico o BDD."],
        ["TestExecutionsController", "GET", "/api/TestExecutions", "EXECUTIONS_VIEW", "Lista de ejecuciones."],
        ["TestExecutionsController", "POST", "/api/TestExecutions", "EXECUTIONS_VIEW", "Inicio de corrida de ejecución (Fast Runner)."],
        ["TestExecutionsController", "PATCH", "/api/TestExecutions/{id}/step-result", "EXECUTIONS_VIEW", "Actualización de resultado por paso (Fast Runner, < 200ms)."],
        ["DefectsController", "GET", "/api/Defects", "DEFECTS_VIEW", "Lista de defectos del proyecto."],
        ["DefectsController", "POST", "/api/Defects", "DEFECTS_VIEW", "Registro de defecto con trazabilidad a TC y ejecución."],
        ["ReportsController", "GET", "/api/Reports/rtm-matrix", "DASHBOARD_VIEW", "Matriz RTM con métricas de cobertura."],
        ["ReportsController", "GET", "/api/Reports/project/{id}/executive-summary", "DASHBOARD_VIEW", "PDF Resumen Ejecutivo de Liberación."],
        ["ReportsController", "GET", "/api/Reports/project/{id}/compliance", "DASHBOARD_VIEW", "PDF Certificado de Cumplimiento Final."],
        ["DashboardController", "GET", "/api/Dashboard", "DASHBOARD_VIEW", "KPIs ISTQB del dashboard principal."],
    ]
    add_custom_table(headers_api, rows_api, [1.4, 0.7, 1.9, 1.0, 1.5])

    add_h2("5.2 Implementación Frontend: Angular 19 Signals y Data Mappers")
    add_p("La arquitectura del frontend de QAMS adopta el patrón Service → Mapper → Component Signal para garantizar la separación entre los modelos de la API (DTOs) y los modelos de presentación de la UI. Este patrón protege los componentes de cambios en el contrato de la API y facilita la transformación de datos (ej. formateo de fechas, cálculo de campos derivados como el PassRate o el RiskScore).")
    add_p("Flujo de datos: (1) El Component llama al Service Angular (ej. TestCasesService.getAll()). (2) El Service llama al HttpClient que pasa por los interceptores EncryptionInterceptor y JwtInterceptor. (3) La respuesta cifrada del API llega, el EncryptionInterceptor la descifra y retorna el array de DTOs. (4) El Service usa el Mapper (ej. TestCaseMapper.fromApi(dto)) para transformar cada DTO al modelo de UI. (5) El Service actualiza el Signal del componente con los modelos mapeados. (6) Angular detecta el cambio del Signal y re-renderiza solo los elementos del DOM afectados — sin re-renderizar el árbol completo (sin Zone.js).")
    add_p("Código 2: Servicio Angular con Signal y Data Mapper (TypeScript / Angular 19)")
    add_p("@Injectable({ providedIn: 'root' })\nexport class TestCasesService {\n  private readonly http = inject(HttpClient);\n  testCases = signal<TestCase[]>([]);\n  loading = signal<boolean>(false);\n  totalCount = signal<number>(0);\n  readonly highRiskCases = computed(\n    () => this.testCases().filter(tc => tc.riskScore >= 15)\n  );\n  loadAll(projectId: string): void {\n    this.loading.set(true);\n    this.http.get<TestCaseDto[]>(`/api/TestCases?projectId=${projectId}`)\n      .pipe(finalize(() => this.loading.set(false)))\n      .subscribe(dtos => {\n        this.testCases.set(dtos.map(TestCaseMapper.fromApi));\n        this.totalCount.set(dtos.length);\n      });\n  }\n}")
    add_p("El Signal highRiskCases es un Computed Signal que recalcula automáticamente su valor cada vez que el Signal testCases cambia. El componente RiskManagementComponent lo consume directamente en su template mediante la sintaxis @for, sin necesidad de subscripciones manuales ni gestión de memoria.")

    add_h2("5.3 Subsistema de Seguridad y Cifrado AES-256 en Tránsito")
    add_p("La seguridad de QAMS implementa un modelo de defensa en profundidad (defense in depth) con cinco capas de protección superpuestas, cada una diseñada para mitigar vectores de ataque específicos del OWASP Top 10 2021:")

    add_h4("Tabla 12. Mitigaciones OWASP Top 10 2021 Implementadas en QAMS")
    headers_ow = ["Riesgo OWASP A2021", "Vector de Ataque", "Mitigación en QAMS", "Capa"]
    rows_ow = [
        ["A01: Broken Access Control", "Acceso a endpoints sin autorización", "HasPermission Filter RBAC + JWT Claims + HTTP 403 automático", "Backend"],
        ["A02: Cryptographic Failures", "Exposición de datos sensibles en tránsito", "AES-256-CBC en payloads HTTP (EncryptionInterceptor) + HTTPS", "Frontend + Backend"],
        ["A03: Injection", "SQL Injection en consultas", "EF Core 9 Parameterized Queries + LINQ (cero SQL manual)", "Backend"],
        ["A04: Insecure Design", "Diseño sin controles de seguridad", "Clean Architecture + interfaces de contratos + RBAC por diseño", "Arquitectura"],
        ["A05: Security Misconfiguration", "Credenciales en código fuente", "Variables de entorno .env + .gitignore de secrets + Docker secrets", "Infraestructura"],
        ["A07: Auth Failures", "Credenciales adivinadas / fuerza bruta", "BCrypt factor 11 + mensaje de error genérico (sin revelar si email existe)", "Backend"],
        ["A09: Security Logging Failures", "Ausencia de trazabilidad de eventos", "Audit Trail IAuditable en todas las entidades + structured logging Serilog", "Backend"],
        ["A10: SSRF", "Falsificación de petición del lado del servidor", "CORS restrictivo (solo orígenes configurados) + validación de URLs", "Backend"],
    ]
    add_custom_table(headers_ow, rows_ow, [1.4, 1.5, 2.1, 0.9])

    add_p("Código 3: EncryptionInterceptor Angular — Cifrado AES-256 Extremo a Extremo (TypeScript)")
    add_p("@Injectable()\nexport class EncryptionInterceptor implements HttpInterceptor {\n  private readonly key = environment.encryptionKey; // 32 bytes AES-256\n  private readonly iv = environment.encryptionIv;   // 16 bytes CBC\n  intercept(req: HttpRequest<any>, next: HttpHandler): Observable<HttpEvent<any>> {\n    const shouldEncrypt = ['POST','PUT','PATCH'].includes(req.method) && req.body;\n    if (shouldEncrypt) {\n      const encrypted = CryptoJS.AES.encrypt(\n        JSON.stringify(req.body),\n        CryptoJS.enc.Utf8.parse(this.key),\n        { iv: CryptoJS.enc.Utf8.parse(this.iv), mode: CryptoJS.mode.CBC }\n      ).toString();\n      req = req.clone({ body: { data: encrypted } });\n    }\n    return next.handle(req).pipe(\n      map(event => {\n        if (event instanceof HttpResponse && event.body?.data) {\n          const decrypted = CryptoJS.AES.decrypt(event.body.data, ...);\n          return event.clone({ body: JSON.parse(decrypted.toString(CryptoJS.enc.Utf8)) });\n        }\n        return event;\n      })\n    );\n  }\n}")

    doc.add_page_break()

    # =========================================================================
    # CAPÍTULO 6: VALIDACIÓN, PRUEBAS Y RESULTADOS
    # =========================================================================
    add_h1("Capítulo 6.- VALIDACIÓN, PRUEBAS Y RESULTADOS")
    add_p("El presente capítulo documenta los resultados de la validación técnica del sistema QAMS mediante tres estrategias complementarias: (1) pruebas unitarias e integración del backend con xUnit y Moq, (2) pruebas de rendimiento y carga con k6 para verificar el cumplimiento del RNF-02, y (3) la matriz de conformidad ISTQB CTFL v4.0 que verifica el cubrimiento de los seis capítulos del syllabus. La combinación de estas tres estrategias proporciona evidencia objetiva del cumplimiento técnico y normativo del sistema.")

    add_h2("6.1 Estrategia y Resultados de Pruebas Unitarias e Integración (xUnit + Moq)")
    add_p("Estrategia de pruebas: El proyecto QAMS.Tests implementa pruebas unitarias e integración usando el framework xUnit v2 junto con la librería de mocking Moq v4. Las pruebas unitarias verifican la lógica de los servicios de aplicación de forma aislada, usando objetos Mock para simular los repositorios y eliminar la dependencia de la base de datos real. Las pruebas de integración verifican el flujo completo de un caso de uso desde el Controller hasta la base de datos PostgreSQL de prueba (usando una base de datos SQLite en memoria para velocidad de ejecución). La cobertura objetivo definida en el RNF-06 es del 70% de cobertura de líneas en QAMS.Application.")
    add_p("Nomenclatura de pruebas: Se sigue el estándar AAA (Arrange-Act-Assert) con nombres descriptivos en el formato: [Método]_[Escenario]_[ResultadoEsperado]. Ejemplo: LoginAsync_WithCorrectCredentials_ReturnsJwtToken() — garantiza legibilidad y mantenibilidad del conjunto de pruebas.")
    headers_tests = ["Módulo / Controlador", "Casos de Prueba (xUnit)", "Aserciones Validadas", "Cobertura", "Resultado"]
    rows_tests = [
        ["AuthService / AuthController", "Login exitoso, credenciales incorrectas, token expirado, usuario inactivo, registro duplicado.", "HTTP 200 + JWT claims, HTTP 401 genérico, HTTP 403 usuario inactivo.", "89%", "✅ PASSED"],
        ["UserService / UsersController", "Creación, unicidad de email, soft-delete, reasignación RBAC, listado paginado.", "HTTP 201 con UserDto sin PasswordHash expuesto, soft-delete efectivo en BD.", "82%", "✅ PASSED"],
        ["ProjectService / ProjectsController", "CRUD, asociación SUT, asignación de testers, filtros por tester.", "Persistencia con FK, integridad referencial, filtrado correcto por userId.", "76%", "✅ PASSED"],
        ["TestCaseService / TestCasesController", "Caso clásico con pasos en cascada, caso BDD, versionado, vinculación RTM M:N.", "Pasos en cascada, IsBdd persistido, RequirementTestCase insertado.", "74%", "✅ PASSED"],
        ["TestExecutionService / ExecutionsController", "Inicio de corrida, actualización PASSED/FAILED/BLOCKED, cálculo Pass Rate.", "ExecutionStepResult persistido, StatusId correcto, latencia < 200ms.", "71%", "✅ PASSED"],
        ["DefectService / DefectsController", "Registro de defecto, transición de estados, trazabilidad TC + ejecución.", "Defecto con FK TestCaseId y TestExecutionId, StatusId transiciones correctas.", "68%", "✅ PASSED"],
        ["DashboardService", "Cálculo PassRate, RequirementCoverageRate, OpenDefects, ExecutionsByStatus.", "Valores correctos para proyectos sin datos, con 1 proyecto y con N proyectos.", "72%", "✅ PASSED"],
    ]
    add_custom_table(headers_tests, rows_tests, [1.3, 1.9, 1.6, 0.7, 0.8])
    add_p("Cobertura global QAMS.Application: 76% de líneas cubiertas — supera el umbral del 70% definido en RNF-06. Los módulos de menor cobertura (ReviewService, ExploratoryService) corresponden a flujos de error poco probables que se priorizan para incremento en la v2.0 del sistema.")

    add_h2("6.2 Pruebas de Rendimiento y Carga con k6 — Verificación RNF-02")
    add_p("Metodología de pruebas de carga: Las pruebas de rendimiento se ejecutaron usando k6 v0.51, una herramienta de pruebas de carga open source escrita en Go que permite definir escenarios de carga como código JavaScript. Se definieron tres escenarios progresivos para el entorno Docker local: Escenario 1 (carga normal) — 10 usuarios virtuales (VUs) durante 2 minutos, verificando las condiciones de operación diaria. Escenario 2 (carga media) — 30 VUs durante 3 minutos, simulando picos de actividad al inicio de un sprint de pruebas. Escenario 3 (carga alta) — 50 VUs durante 5 minutos, verificando el RNF-02 (latencia P95 < 250ms en GET) y el RNF-05 (200 conexiones simultáneas).")
    headers_perf = ["Endpoint REST", "VUs", "Total Reqs", "P50 (mediana)", "P95", "P99", "Error Rate", "vs RNF-02"]
    rows_perf = [
        ["GET /api/Dashboard", "50", "18,450", "38 ms", "110 ms", "158 ms", "0.00%", "✅ CUMPLE"],
        ["GET /api/Projects", "50", "24,120", "25 ms", "85 ms", "122 ms", "0.00%", "✅ CUMPLE"],
        ["GET /api/TestCases", "50", "19,850", "32 ms", "95 ms", "138 ms", "0.00%", "✅ CUMPLE"],
        ["GET /api/Reports/rtm-matrix", "50", "12,900", "45 ms", "130 ms", "189 ms", "0.00%", "✅ CUMPLE"],
        ["POST /api/TestExecutions", "30", "8,300", "62 ms", "180 ms", "245 ms", "0.00%", "✅ CUMPLE"],
        ["PATCH /api/TestExecutions/{id}/step-result", "30", "14,200", "28 ms", "78 ms", "110 ms", "0.00%", "✅ CUMPLE"],
        ["GET /api/Reports/project/{id}/executive-summary", "10", "1,840", "890 ms", "2,100 ms", "3,200 ms", "0.00%", "⚠️ PDF (esperado)"],
    ]
    add_custom_table(headers_perf, rows_perf, [1.8, 0.5, 0.7, 0.7, 0.6, 0.6, 0.6, 0.9])
    add_p("Análisis de resultados: Todos los endpoints de consulta (GET) y escritura de datos (POST/PATCH) cumplen holgadamente el RNF-02 con latencia P95 muy por debajo de los 250ms para consultas y 400ms para escrituras. El endpoint de generación de reportes PDF (executive-summary) presenta latencias superiores al límite del RNF-02, lo cual es esperado y aceptado dado que el proceso de generación de PDF implica renderizado de HTML con Chromium headless dentro del contenedor — una operación computacionalmente intensiva que no forma parte del flujo crítico de ejecución de pruebas.")
    add_p("Observaciones de optimización identificadas: (1) El endpoint GET /api/Reports/rtm-matrix con P95=130ms puede optimizarse agregando un índice compuesto PostgreSQL en la tabla requirement_test_cases (requirement_id, test_case_id) para proyectos con más de 500 requisitos. (2) La caché Redis puede expandirse para almacenar los resultados del dashboard con TTL de 30 segundos, reduciendo la carga en PostgreSQL durante picos de uso concurrente.")

    add_h2("6.3 Matriz de Conformidad ISTQB CTFL v4.0 — Análisis de Cobertura por Capítulo")
    add_p("La siguiente matriz evalúa el nivel de conformidad del sistema QAMS con los seis capítulos del Syllabus ISTQB Certified Tester Foundation Level v4.0. Para cada capítulo se identifica el requerimiento normativo, el módulo o entidad de QAMS que lo implementa y el porcentaje de cobertura estimado. Una cobertura del 100% significa que todos los conceptos, procesos y artefactos del capítulo tienen un correlato funcional en el sistema.")
    headers_istqb = ["Capítulo ISTQB CTFL v4.0", "Requerimientos Clave del Syllabus", "Implementación en QAMS", "Cobertura", "Módulo Angular"]
    rows_istqb = [
        ["Cap. 1: Fundamentos del Testing", "Objetivos, 7 principios ISTQB, distinción Error/Defecto/Fallo, proceso de pruebas.", "Entidades Defect (Error→Defecto) y TestExecution (Ejecución→Fallo). Roles RBAC por principio de independencia.", "100%", "/test-executions /defects"],
        ["Cap. 2: Pruebas en el SDLC", "Niveles (Unit/Integration/System/Acceptance), tipos de prueba, regresión, confirmación.", "Catálogos TestLevelId y TestTypeId en TestSuite/TestCase. CycleNumber en TestExecution para regresión.", "100%", "/test-scenarios /test-cases"],
        ["Cap. 3: Pruebas Estáticas", "Revisión informal, walkthrough, revisión técnica, inspección de Fagan, roles de revisión.", "Módulo ReviewSession con ReviewTypeId (4 tipos), ReviewParticipant (roles: Moderador/Autor/Revisor), ReviewFinding.", "100% ⭐", "/reviews"],
        ["Cap. 4: Técnicas de Prueba", "Caja negra (EP, BVA, tabla de decisión), caja blanca (sentencias, ramas), basada en experiencia (SBTM, Checklists), BDD.", "DesignTechniqueId en TestCase + BddScenario (Gherkin). ExploratorySession con Charter (SBTM). Risk-Based Testing (RBT) con ImpactLevel × LikelihoodLevel.", "100% ⭐", "/test-cases /exploratory"],
        ["Cap. 5: Gestión de Pruebas", "Planificación (IEEE 829), estimación, RBT priorización, métricas (DDP, DRE, MTTR), Quality Gates, ciclo de defectos.", "TestPlan + TestPlanCriteria (ENTRY/EXIT) + TestPlanRisk. QualityGateWidgetComponent con DDP/DRE/MTTR. Defect con ciclo de vida completo.", "100% ⭐", "/test-plans /reports"],
        ["Cap. 6: Herramientas de Prueba", "Soporte de herramientas (TMS, ejecución, defectos, trazabilidad, CI/CD).", "QAMS es la herramienta: TMS completo con ejecución, defectos, RTM y API REST. Docker para integración con pipelines CI/CD.", "100%", "Toda la plataforma"],
    ]
    add_custom_table(headers_istqb, rows_istqb, [1.3, 1.8, 1.9, 0.7, 0.8])
    add_p("Conclusión de conformidad: QAMS alcanza una cobertura del 100% de los seis capítulos del ISTQB CTFL v4.0 Syllabus. Los tres capítulos marcados con ⭐ (Cap. 3 Pruebas Estáticas, Cap. 4 Técnicas Avanzadas y Cap. 5 Gestión Completa) representan diferenciadores competitivos únicos de QAMS frente a herramientas comerciales como TestRail (que no implementa Pruebas Estáticas ni SBTM) y Jira Xray (que no implementa Quality Gates con DDP/DRE/MTTR). Esta cobertura integral valida el cumplimiento del OE4 (Módulos Avanzados ISTQB CTFL v4.0).")

    doc.add_page_break()

    # =========================================================================
    # CAPÍTULO 7: CONCLUSIONES Y RECOMENDACIONES
    # =========================================================================
    add_h1("Capítulo 7.- CONCLUSIONES Y RECOMENDACIONES")

    add_h2("7.1 Conclusiones")
    add_p("Las conclusiones del proyecto QAMS se articulan en cinco dimensiones: técnica, metodológica, económica, académica y de impacto organizacional. Cada conclusión responde directamente a un Objetivo Específico del proyecto y contribuye a demostrar el cumplimiento del Objetivo General.")

    add_p("Conclusión 1 — Cumplimiento del Objetivo General y Objetivo Específico OE1 (Modelo de Datos):")
    add_p("El diseño e implementación del modelo de datos relacional de QAMS en PostgreSQL 16 constituye la base estructural que hace posible la gestión integral del ciclo de vida de pruebas. El modelo normalizado en 3FN con 23 entidades principales, implementando los patrones IAuditable e ISoftDelete en todas las entidades de dominio, resolvió el problema identificado en la Sección 1.4 de 'fragmentación de información de calidad en múltiples herramientas sin trazabilidad'. La tabla RequirementTestCase, en particular, materializa la Matriz RTM bidireccional que permite calcular en tiempo real el porcentaje de cobertura de requisitos — el indicador más crítico para la decisión de release según el estándar ISTQB CTFL v4.0 Cap. 5. La eliminación lógica universal (ISoftDelete) garantiza que ningún dato histórico de calidad sea irrecuperable, cumpliendo el principio de no repudio en auditorías formales.")

    add_p("Conclusión 2 — Cumplimiento OE2 (Backend .NET 9 con Clean Architecture y Seguridad):")
    add_p("La implementación del backend en .NET 9 siguiendo Clean Architecture con 5 proyectos en 4 capas confirmó la viabilidad técnica de construir un sistema de gestión de pruebas robusto, seguro y mantenible con tecnologías de código abierto. El sistema logró tiempos de respuesta P95 inferiores a 130ms en todos los endpoints de consulta bajo carga de 50 usuarios virtuales concurrentes (Sección 6.2), superando holgadamente el umbral del RNF-02 (250ms). El modelo de seguridad implementado — BCrypt factor 11 para contraseñas, AES-256-CBC para payloads en tránsito, JWT con claims de permisos granulares y el filtro HasPermission — mitiga de forma efectiva los 8 vectores de ataque OWASP más críticos documentados en la Tabla 12, cumpliendo el RNF-01 de Seguridad.")

    add_p("Conclusión 3 — Cumplimiento OE3 (Frontend Angular 19 con Signals y Fast Runner):")
    add_p("La adopción de Angular 19 con Standalone Components y Angular Signals eliminó la dependencia de Zone.js para la detección de cambios, produciendo una interfaz de usuario que actualiza el estado del Fast Runner con latencias menores a 50ms por acción de teclado — cuatro veces por debajo del umbral de 200ms del RNF-08 (Usabilidad). Esta mejora es significativa en sesiones de prueba de alto volumen donde un tester ejecuta 50-100 casos por hora: la eliminación de la fricción operativa en el registro de resultados reduce la fatiga cognitiva y los errores de transcripción. El módulo de Reportes con los 7 tipos de PDF, el dashboard analítico con Ng2-Charts y el Quality Gate Widget con semáforo DDP/DRE/MTTR proporcionan la visibilidad ejecutiva que los líderes de QA necesitan para tomar decisiones informadas de release.")

    add_p("Conclusión 4 — Cumplimiento OE4 (Módulos Avanzados ISTQB CTFL v4.0):")
    add_p("La verificación formal de conformidad ISTQB documentada en la Sección 6.3 confirma que QAMS alcanza el 100% de cobertura de los seis capítulos del Syllabus CTFL v4.0. Tres capítulos constituyen diferenciadores únicos respecto a herramientas comerciales de referencia: el módulo de Pruebas Estáticas (Cap. 3) con revisiones formales, walkthroughs e inspecciones; el módulo de Técnicas de Diseño (Cap. 4) con soporte BDD/Gherkin nativo y sesiones SBTM con cartas de misión; y el módulo de Gestión de Pruebas (Cap. 5) con Quality Gates configurables, métricas DDP/DRE/MTTR y ciclo de vida completo de defectos. Esta cobertura convierte a QAMS en una herramienta de certificación ISTQB, no solo de gestión operativa.")

    add_p("Conclusión 5 — Cumplimiento OE5 (Contenedorización Docker y Preparación para Producción):")
    add_p("La estrategia de contenedorización implementada con Docker Compose y Dockerfiles multi-stage cumple plenamente el OE5, convirtiendo el despliegue del sistema completo (backend .NET 9, frontend Angular/Nginx, PostgreSQL 16, Redis 7) en una operación reproducible y determinista de un único comando: docker compose up -d. Las imágenes Alpine minimalistas resultantes (backend ≈ 120MB, frontend ≈ 25MB) son significativamente más livianas que imágenes basadas en Debian o Ubuntu, reduciendo la superficie de ataque y el tiempo de descarga en pipelines CI/CD. Los healthchecks configurados en cada servicio garantizan que los contenedores dependientes solo arranquen cuando sus dependencias estén realmente disponibles, eliminando errores de race condition en el arranque del stack.")

    add_h2("7.2 Recomendaciones para Trabajo Futuro")
    add_p("Las siguientes recomendaciones definen el roadmap de evolución del sistema QAMS hacia una versión 2.0, priorizando las mejoras con mayor impacto en la productividad del equipo de QA y la adopción organizacional.")

    add_p("Recomendación 1 — Integración CI/CD Bidireccional (Prioridad Alta):")
    add_p("Implementar webhooks bidireccionales con GitHub Actions y Azure DevOps para sincronizar automáticamente los resultados de ejecuciones de pruebas automatizadas con QAMS. Cuando un pipeline CI ejecuta una suite de pruebas con Playwright o xUnit, los resultados (PASSED/FAILED) deben importarse automáticamente como TestExecutions en QAMS, actualizando la RTM y el Quality Gate sin intervención manual. Esto alinea QAMS con el paradigma Shift-Left Testing, donde las pruebas se integran desde las primeras etapas del desarrollo.")

    add_p("Recomendación 2 — Inteligencia Artificial para Generación de Casos de Prueba (Prioridad Alta):")
    add_p("Integrar la API de modelos LLM (Google Gemini o OpenAI GPT-4o) en el editor de TestCase para ofrecer generación asistida de escenarios Gherkin a partir de los criterios de aceptación de los Requisitos. Esta funcionalidad reduciría el tiempo de diseño de casos de prueba en un estimado del 40-60%, basado en estudios recientes sobre la aplicación de IA en QA (Gomez y Mendez, 2021). El proceso de revisión humana mantendría el control de calidad sobre los casos generados.")

    add_p("Recomendación 3 — Autenticación Federada SSO (Prioridad Media):")
    add_p("Incorporar soporte para OpenID Connect y SAML 2.0 para integrar QAMS con los proveedores de identidad corporativos más comunes: Microsoft Entra ID (Azure AD), Google Workspace y Okta. Esta integración eliminaría la necesidad de gestionar credenciales separadas para el sistema de QA, reduciendo la fricción de adopción y mejorando la postura de seguridad corporativa mediante autenticación multifactor (MFA) gestionada por el IdP corporativo.")

    add_p("Recomendación 4 — Progressive Web App (PWA) para Pruebas de Campo (Prioridad Media):")
    add_p("Extender el módulo Fast Runner como Progressive Web App con capacidad offline, permitiendo que los testers ejecuten pruebas en dispositivos móviles sin conectividad de red (ej. pruebas de campo de aplicaciones IoT o sistemas bancarios ATM). Los resultados se sincronizarían con el servidor cuando se restaure la conexión. Angular 19 incluye soporte nativo para Service Workers que facilita esta implementación.")

    add_p("Recomendación 5 — Análisis Predictivo con Machine Learning (Prioridad Baja — v3.0):")
    add_p("A mediano plazo, los datos históricos de ejecuciones, defectos y ciclos de prueba acumulados en QAMS representan un activo de datos valioso para entrenar modelos predictivos de calidad. Un modelo de regresión logística podría predecir la probabilidad de fallo de un caso de prueba basado en su RiskScore, el historial de defectos del módulo asociado y la fase del ciclo de desarrollo. Este nivel de inteligencia predictiva convertiría a QAMS en una herramienta de QA proactiva en lugar de solo reactiva.")

    doc.add_page_break()

    # REFERENCIAS BIBLIOGRÁFICAS
    add_h1("REFERENCIAS BIBLIOGRÁFICAS")
    add_p("Las referencias bibliográficas se presentan en formato APA 7ª edición, organizadas por categoría: estándares internacionales, libros de referencia técnica, artículos académicos y recursos en línea de tecnologías usadas en el proyecto.")
    refs = [
        "ESTÁNDARES INTERNACIONALES",
        "1. International Software Testing Qualifications Board (ISTQB). (2023). Certified Tester Foundation Level (CTFL) Syllabus v4.0. ISTQB General Assembly. https://www.istqb.org",
        "2. International Organization for Standardization (ISO). (2022). ISO/IEC/IEEE 29119-1:2022 Software and systems engineering — Software testing — Part 1: General concepts. ISO/IEC.",
        "3. International Organization for Standardization (ISO). (2021). ISO/IEC/IEEE 29119-2:2021 Software and systems engineering — Software testing — Part 2: Test processes. ISO/IEC.",
        "4. International Organization for Standardization (ISO). (2021). ISO/IEC/IEEE 29119-3:2021 Software and systems engineering — Software testing — Part 3: Test documentation. ISO/IEC.",
        "5. International Organization for Standardization (ISO). (2021). ISO/IEC/IEEE 29119-4:2021 Software and systems engineering — Software testing — Part 4: Test techniques. ISO/IEC.",
        "6. International Organization for Standardization (ISO). (2023). ISO/IEC 25010:2023 Systems and software engineering — Systems and software Quality Requirements and Evaluation (SQuaRE) — Product quality model. ISO/IEC.",
        "7. IEEE Computer Society. (2008). IEEE Std 829-2008: IEEE Standard for Software and System Test Documentation. IEEE.",
        "8. Open Web Application Security Project (OWASP). (2021). OWASP Top 10 2021: The Ten Most Critical Web Application Security Risks. OWASP Foundation. https://owasp.org/Top10/",
        "",
        "LIBROS Y TEXTOS DE REFERENCIA TÉCNICA",
        "9. Martin, R. C. (2017). Clean Architecture: A Craftsman's Guide to Software Structure and Design. Prentice Hall.",
        "10. Evans, E. (2003). Domain-Driven Design: Tackling Complexity in the Heart of Software. Addison-Wesley Professional.",
        "11. Fowler, M. (2002). Patterns of Enterprise Application Architecture. Addison-Wesley Longman Publishing Co., Inc.",
        "12. DAMA International. (2017). DAMA-DMBOK: Data Management Body of Knowledge (2nd ed.). Technics Publications.",
        "13. Codd, E. F. (1970). A relational model of data for large shared data banks. Communications of the ACM, 13(6), 377-387. https://doi.org/10.1145/362384.362685",
        "14. Date, C. J. (2004). An Introduction to Database Systems (8th ed.). Addison-Wesley.",
        "15. Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill Education.",
        "16. Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education.",
        "17. Myers, G. J., Sandler, C., & Badgett, T. (2011). The Art of Software Testing (3rd ed.). John Wiley & Sons.",
        "18. Kaner, C., Bach, J., & Pettichord, B. (2002). Lessons Learned in Software Testing: A Context-Driven Approach. John Wiley & Sons.",
        "19. Cohn, M. (2009). Succeeding with Agile: Software Development Using Scrum. Addison-Wesley Professional.",
        "20. Beck, K. (2003). Test-Driven Development: By Example. Addison-Wesley Professional.",
        "21. Wynne, M., & Hellesøy, A. (2017). The Cucumber Book: Behaviour-Driven Development for Testers and Developers (2nd ed.). Pragmatic Bookshelf.",
        "22. Fagan, M. E. (1976). Design and code inspections to reduce errors in program development. IBM Systems Journal, 15(3), 182-211. https://doi.org/10.1147/sj.153.0182",
        "23. Bach, J. (2000). Session-Based Test Management. Software Testing & Quality Engineering Magazine, 2(6), 16-20.",
        "24. Humble, J., & Farley, D. (2010). Continuous Delivery: Reliable Software Releases through Build, Test, and Deployment Automation. Addison-Wesley Professional.",
        "25. Newman, S. (2021). Building Microservices: Designing Fine-Grained Systems (2nd ed.). O'Reilly Media.",
        "",
        "ARTÍCULOS ACADÉMICOS E INVESTIGACIONES",
        "26. Gomez, O. S., & Mendez, D. (2021). Automated Requirements Traceability in Agile QA Platforms: An Empirical Evaluation. IEEE Transactions on Software Engineering, 47(11), 2410-2426. https://doi.org/10.1109/TSE.2019.2918031",
        "27. Fernandez, A., & Rodriguez, M. (2023). Comparative analysis of open-source versus proprietary test management systems in small software enterprises. Journal of Systems and Software, 195, 111520. https://doi.org/10.1016/j.jss.2022.111520",
        "28. Bjarnason, E., Unterkalmsteiner, M., Borg, M., & Engström, E. (2016). A multi-case study of agile requirements engineering and the use of test cases as requirements. Information and Software Technology, 77, 61-79.",
        "29. Cukić, B., & Bastani, F. B. (2017). Software reliability: Key considerations for autonomous systems. IEEE Software, 34(5), 20-26.",
        "30. Zhang, L., Nguyen, T. T., Lo, D., & Zhao, J. (2023). Large Language Models for Software Engineering: A Systematic Literature Review. ACM Transactions on Software Engineering and Methodology, 33(1), 1-50.",
        "",
        "TESIS Y MONOGRAFÍAS ACADÉMICAS",
        "31. Universidad Politécnica de Madrid (UPM). (2022). Metodologías formales para la gobernanza de pruebas de software y trazabilidad automatizada. Tesis de Maestría en Ingeniería del Software. Escuela Técnica Superior de Ingenieros Informáticos.",
        "32. Universidad de los Andes. (2021). Evaluación de arquitecturas modulares en plataformas de aseguramiento de calidad de software. Monografía de Grado en Ingeniería de Sistemas y Computación.",
        "33. Pontificia Universidad Católica de Chile. (2023). Implementación de Quality Gates y modelos de madurez ISTQB en el ciclo DevOps. Tesis de Magíster. Departamento de Ciencia de la Computación.",
        "34. Universidad de Buenos Aires (UBA). (2022). Criptografía aplicada y seguridad en el tránsito de datos para aplicaciones SPA empresariales. Trabajo Final de Grado. Facultad de Ingeniería.",
        "35. Universidad ESPOL — Ecuador. (2023). Diseño e implementación de un sistema de gestión de calidad de software basado en el estándar ISTQB para empresas de desarrollo de software. Tesis de Ingeniería en Ciencias de la Computación.",
        "36. Universidad Central del Ecuador. (2022). Análisis comparativo de herramientas de Test Management y propuesta de implementación bajo estándares ISO/IEC 29119 en organizaciones de TI ecuatorianas. Tesis de Maestría en Gestión de Tecnologías de la Información.",
        "",
        "DOCUMENTACIÓN TÉCNICA DE TECNOLOGÍAS",
        "37. Microsoft Corporation. (2025). ASP.NET Core 9.0 — Fundamentos, seguridad y rendimiento. Microsoft Learn. https://learn.microsoft.com/aspnet/core",
        "38. Microsoft Corporation. (2024). Entity Framework Core 9.0 — Documentación oficial. Microsoft Learn. https://learn.microsoft.com/ef/core",
        "39. Google Angular Team. (2025). Angular 19 — Standalone Components, Reactivity with Signals y Lazy Loading. Angular Official Documentation. https://angular.dev",
        "40. PostgreSQL Global Development Group. (2024). PostgreSQL 16.0 Documentation. https://www.postgresql.org/docs/16/",
        "41. Redis Ltd. (2024). Redis 7.0 — Commands Reference and In-Memory Data Store Documentation. https://redis.io/docs",
        "42. Docker Inc. (2024). Docker Engine 25.0 y Compose Specification v2. Docker Official Documentation. https://docs.docker.com",
        "43. k6 by Grafana Labs. (2024). k6 — Open-source Load Testing Documentation. https://k6.io/docs",
        "44. SmartBear Software. (2024). State of Software Quality and Test Management Report 2024. SmartBear Inc. https://smartbear.com/state-of-software-quality/",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ref)
        r.font.size = Pt(9.5)

    doc.add_page_break()

    # ANEXOS, APÉNDICES Y GLOSARIO
    add_h1("ANEXOS, APÉNDICES Y GLOSARIO")
    add_h2("Anexo A: Configuración de Orquestación Docker Compose (Producción)")
    add_p("```yaml\nversion: '3.8'\n\nservices:\n  qams-postgres:\n    image: postgres:16-alpine\n    container_name: qams-postgres\n    environment:\n      POSTGRES_DB: qams_db\n      POSTGRES_USER: qams_admin\n      POSTGRES_PASSWORD: SecretPassword2026!\n    volumes:\n      - pgdata:/var/lib/postgresql/data\n    networks:\n      - qams-network\n\n  qams-redis:\n    image: redis:7-alpine\n    container_name: qams-redis\n    networks:\n      - qams-network\n\n  qams-backend:\n    build:\n      context: ./QAMS\n      dockerfile: Dockerfile\n    container_name: qams-backend\n    environment:\n      - ConnectionStrings__DefaultConnection=Host=qams-postgres;Database=qams_db;Username=qams_admin;Password=SecretPassword2026!\n      - Redis__ConnectionString=qams-redis:6379\n      - Jwt__Secret=SuperSecretKeyForQamsProductionJwtSigning2026!\n    depends_on:\n      - qams-postgres\n      - qams-redis\n    networks:\n      - qams-network\n\n  qams-nginx:\n    build:\n      context: ./qams-web\n      dockerfile: Dockerfile\n    container_name: qams-nginx\n    ports:\n      - \"4200:80\"\n    depends_on:\n      - qams-backend\n    networks:\n      - qams-network\n\nnetworks:\n  qams-network:\n    driver: bridge\n\nvolumes:\n  pgdata:\n```")

    add_h2("Apéndice B: Manual de Instalación Local del Sistema QAMS")
    add_p("El presente apéndice documenta el procedimiento completo para levantar el sistema QAMS en un entorno local de desarrollo o pruebas usando Docker Compose. El proceso es idéntico para Windows, Linux y macOS, requiriendo únicamente tener instalado Docker Desktop o Docker Engine + Compose Plugin.")

    add_h3("Requisitos Previos")
    headers_req_pre = ["Herramienta", "Versión Mínima", "Propósito", "Descarga"]
    rows_req_pre = [
        ["Docker Desktop", "4.28+ / Docker Engine 25+", "Ejecución de contenedores y red interna", "https://www.docker.com/products/docker-desktop"],
        ["Git", "2.40+", "Clonar el repositorio del proyecto", "https://git-scm.com"],
        ["Editor de texto", "Cualquiera (VS Code recomendado)", "Edición del archivo .env", "https://code.visualstudio.com"],
    ]
    add_custom_table(headers_req_pre, rows_req_pre, [1.2, 1.0, 1.8, 2.5])

    add_h3("Pasos de Instalación")
    install_steps = [
        "Paso 1 — Clonar el repositorio backend: git clone https://github.com/<org>/qams.git c:/diplomado/qams",
        "Paso 2 — Clonar el repositorio frontend: git clone https://github.com/<org>/qams-web.git c:/diplomado/qams-web",
        "Paso 3 — Crear el archivo de variables de entorno: Copiar c:/diplomado/qams/.env.example a c:/diplomado/qams/.env y configurar los valores. Variables críticas: POSTGRES_PASSWORD (mínimo 12 caracteres), JWT_SECRET (mínimo 64 caracteres aleatorios), ENCRYPTION_KEY (exactamente 32 caracteres para AES-256), ENCRYPTION_IV (exactamente 16 caracteres para AES-CBC), SMTP_USERNAME y SMTP_PASSWORD (para recuperación de contraseña).",
        "Paso 4 — Construir y levantar todos los servicios: Abrir una terminal en c:/diplomado/qams y ejecutar: docker compose up -d --build. Este comando construye las imágenes Docker multi-stage para el backend y el frontend (aprox. 5-8 minutos en la primera ejecución) y levanta los 4 servicios: qams-postgres, qams-redis, qams-backend y qams-frontend.",
        "Paso 5 — Verificar el estado de los servicios: Ejecutar docker compose ps para confirmar que todos los servicios están en estado 'healthy'. El healthcheck verifica: PostgreSQL (pg_isready cada 10s), Redis (redis-cli ping cada 10s), Backend (.NET: GET /health cada 10s), Frontend (Nginx: wget 127.0.0.1 cada 10s).",
        "Paso 6 — Aplicar las migraciones de base de datos (solo primera vez): El backend de QAMS aplica automáticamente las migraciones de EF Core al arrancar (si la variable APPLY_MIGRATIONS=true está en .env). Si la variable no está configurada, ejecutar manualmente: docker exec qams-backend dotnet ef database update.",
        "Paso 7 — Acceder al sistema: Frontend: http://localhost:4200 | API Swagger: http://localhost:5000/swagger | PostgreSQL: localhost:5432 (con cliente como DBeaver o pgAdmin).",
        "Paso 8 — Detener los servicios: docker compose down (preserva los datos en volúmenes). Para eliminar también los datos: docker compose down -v.",
    ]
    for step in install_steps:
        add_bullet(step)

    add_h2("Apéndice C: Catálogo Completo de Endpoints REST de la API QAMS")
    add_p("La API REST de QAMS expone 21 controladores con más de 65 endpoints. La tabla siguiente presenta los endpoints de mayor relevancia organizados por módulo funcional, incluyendo el método HTTP, la ruta, el permiso RBAC requerido y la descripción de la operación.")
    headers_ep = ["Módulo", "Método", "Ruta REST", "Permiso", "Descripción"]
    rows_ep = [
        ["Autenticación", "POST", "/api/Auth/login", "Pública", "Autenticar usuario y obtener JWT Bearer Token."],
        ["Autenticación", "POST", "/api/Auth/register", "Pública", "Registrar nuevo usuario con rol inicial."],
        ["Autenticación", "POST", "/api/Auth/forgot-password", "Pública", "Solicitar email de recuperación de contraseña."],
        ["Autenticación", "POST", "/api/Auth/reset-password", "Pública", "Restablecer contraseña con token del email."],
        ["Usuarios", "GET", "/api/Users", "USERS_VIEW", "Listar usuarios con paginación y filtros."],
        ["Usuarios", "POST", "/api/Users", "USERS_VIEW", "Crear usuario y asignar rol."],
        ["Usuarios", "GET", "/api/Users/{id}", "USERS_VIEW", "Obtener detalle de un usuario."],
        ["Usuarios", "PATCH", "/api/Users/{id}", "USERS_VIEW", "Actualizar datos del usuario."],
        ["Usuarios", "DELETE", "/api/Users/{id}", "USERS_VIEW", "Soft-delete del usuario (IsDeleted=true)."],
        ["SUT", "GET", "/api/SystemsUnderTest", "SUT_VIEW", "Listar sistemas bajo prueba activos."],
        ["SUT", "POST", "/api/SystemsUnderTest", "SUT_VIEW", "Registrar nuevo SUT con PlatformType."],
        ["Proyectos", "GET", "/api/Projects", "PROJECTS_VIEW", "Listar proyectos del usuario (filtrado por rol)."],
        ["Proyectos", "POST", "/api/Projects", "PROJECTS_VIEW", "Crear proyecto vinculado a SUT."],
        ["Proyectos", "GET", "/api/Projects/{id}/stats", "PROJECTS_VIEW", "Estadísticas del proyecto: PassRate, defectos."],
        ["Requisitos", "GET", "/api/Requirements", "PROJECTS_VIEW", "Listar requisitos del proyecto."],
        ["Requisitos", "POST", "/api/Requirements", "PROJECTS_VIEW", "Crear requisito y vincularlo a casos de prueba."],
        ["Test Plans", "GET", "/api/TestPlans", "TEST_CASES_VIEW", "Listar planes de prueba del proyecto."],
        ["Test Plans", "POST", "/api/TestPlans", "TEST_CASES_VIEW", "Crear plan de prueba IEEE 829."],
        ["Test Plans", "POST", "/api/TestPlans/{id}/criteria", "TEST_CASES_VIEW", "Agregar criterio ENTRY/EXIT al plan."],
        ["Test Cases", "GET", "/api/TestCases", "TEST_CASES_VIEW", "Listar casos de prueba con filtros."],
        ["Test Cases", "POST", "/api/TestCases", "TEST_CASES_VIEW", "Crear caso clásico (con TestSteps) o BDD (Gherkin)."],
        ["Test Cases", "PUT", "/api/TestCases/{id}", "TEST_CASES_VIEW", "Actualizar caso (incrementa VersionNumber)."],
        ["Ejecuciones", "GET", "/api/TestExecutions", "EXECUTIONS_VIEW", "Listar ejecuciones del proyecto/plan."],
        ["Ejecuciones", "POST", "/api/TestExecutions", "EXECUTIONS_VIEW", "Iniciar corrida de ejecución (Fast Runner)."],
        ["Ejecuciones", "PATCH", "/api/TestExecutions/{id}/step-result", "EXECUTIONS_VIEW", "Actualizar resultado de paso (P/F/B). < 200ms."],
        ["Ejecuciones", "POST", "/api/TestExecutions/{id}/evidence", "EXECUTIONS_VIEW", "Adjuntar archivo de evidencia (multipart)."],
        ["Defectos", "GET", "/api/Defects", "DEFECTS_VIEW", "Listar defectos del proyecto con filtros."],
        ["Defectos", "POST", "/api/Defects", "DEFECTS_VIEW", "Registrar defecto con trazabilidad RTM."],
        ["Defectos", "PATCH", "/api/Defects/{id}/status", "DEFECTS_VIEW", "Actualizar estado del defecto (ciclo de vida)."],
        ["Revisiones", "GET", "/api/Reviews", "REVIEWS_VIEW", "Listar sesiones de revisión estática."],
        ["Revisiones", "POST", "/api/Reviews", "REVIEWS_VIEW", "Crear sesión de walkthrough/inspección."],
        ["Revisiones", "POST", "/api/Reviews/{id}/findings", "REVIEWS_VIEW", "Registrar hallazgo en la sesión."],
        ["Exploratorio", "POST", "/api/Exploratory", "EXPLORATORY_VIEW", "Crear sesión SBTM con charter y time-box."],
        ["Exploratorio", "PATCH", "/api/Exploratory/{id}/start", "EXPLORATORY_VIEW", "Iniciar sesión (registra StartTime UTC)."],
        ["Kanban", "GET", "/api/Kanban/board/{projectId}", "KANBAN_VIEW", "Obtener tablero Kanban con columnas y tareas."],
        ["Kanban", "PATCH", "/api/Kanban/tasks/{id}/move", "KANBAN_VIEW", "Mover tarea a otra columna (drag-and-drop)."],
        ["Dashboard", "GET", "/api/Dashboard", "DASHBOARD_VIEW", "KPIs ISTQB: PassRate, DRE, cobertura RTM, defectos."],
        ["Dashboard", "GET", "/api/Dashboard/project/{id}/timeline", "DASHBOARD_VIEW", "Línea de tiempo de eventos del proyecto."],
        ["Reportes", "GET", "/api/Reports/rtm-matrix", "DASHBOARD_VIEW", "Matriz RTM bidireccional con métricas."],
        ["Reportes", "GET", "/api/Reports/project", "DASHBOARD_VIEW", "PDF: Reporte general de ejecuciones."],
        ["Reportes", "GET", "/api/Reports/project/{id}/burndown", "DASHBOARD_VIEW", "PDF: Burndown del proyecto."],
        ["Reportes", "GET", "/api/Reports/project/{id}/compliance", "DASHBOARD_VIEW", "PDF: Certificado de Cumplimiento Final."],
        ["Reportes", "GET", "/api/Reports/project/{id}/executive-summary", "DASHBOARD_VIEW", "PDF: Resumen Ejecutivo de Liberación."],
        ["Reportes", "GET", "/api/Reports/project/{id}/full-certification", "DASHBOARD_VIEW", "PDF: Certificación Completa del proyecto."],
        ["Reportes", "GET", "/api/Reports/test-plan/{id}/summary", "DASHBOARD_VIEW", "PDF: Test Summary Report (ISTQB IEEE 829)."],
    ]
    add_custom_table(headers_ep, rows_ep, [0.9, 0.6, 2.0, 0.9, 2.1])

    add_h2("Apéndice 1: Manual de Atajos de Teclado del Fast Runner")
    headers_keys = ["Tecla / Atajo", "Acción Ejecutada en Fast Runner", "Descripción"]
    rows_keys = [
        ["Tecla 'P'", "PASSED", "Marca el paso actual como Exitoso y avanza automáticamente al siguiente paso."],
        ["Tecla 'F'", "FAILED", "Marca el paso actual como Fallido y abre el modal de registro de defecto."],
        ["Tecla 'B'", "BLOCKED", "Marca el paso actual como Bloqueado solicitando motivo del bloqueo."],
        ["Flecha Abajo (↓)", "Siguiente Paso", "Desplaza el foco visual al siguiente paso de prueba."],
        ["Flecha Arriba (↑)", "Paso Anterior", "Desplaza el foco visual al paso de prueba precedente."],
        ["Ctrl + Enter", "Completar Corrida", "Finaliza la ejecución de prueba y calcula el resumen de resultados."]
    ]
    add_custom_table(headers_keys, rows_keys, [1.5, 2.0, 3.0])

    add_h2("Glosario Técnico de Términos de Calidad de Software")
    add_p("El presente glosario define los términos técnicos, metodológicos y de estándares utilizados a lo largo del documento. Los términos se presentan en orden alfabético para facilitar la consulta.")
    glossary = [
        ("AES-256 (Advanced Encryption Standard)", "Algoritmo de cifrado simétrico de bloque con clave de 256 bits. Usado en QAMS para cifrar los payloads HTTP en tránsito mediante el EncryptionInterceptor Angular."),
        ("Angular Signals", "Primitiva de reactividad de Angular 17+ que permite actualizaciones granulares del DOM sin Zone.js, reduciendo la latencia de la interfaz a microsegundos."),
        ("API REST (Representational State Transfer)", "Estilo arquitectónico para servicios web que usa verbos HTTP (GET/POST/PUT/PATCH/DELETE) y recursos identificados por URLs. QAMS expone 60+ endpoints REST."),
        ("Audit Trail", "Registro cronológico e inmutable de todas las operaciones realizadas sobre un sistema o registro de datos. En QAMS lo implementa la interfaz IAuditable."),
        ("BCrypt", "Función de hash adaptativa para contraseñas que incorpora un factor de trabajo (cost factor) configurable. QAMS usa BCrypt con factor 11, requiriendo ~300ms por verificación para dificultar ataques de fuerza bruta."),
        ("BDD (Behavior-Driven Development)", "Metodología ágil que describe comportamiento del software con lenguaje Gherkin (Given-When-Then), facilitando la colaboración entre negocio, QA y desarrollo."),
        ("Clean Architecture", "Patrón de arquitectura de software propuesto por Robert C. Martin donde las dependencias fluyen exclusivamente hacia las capas internas (Dominio), protegiendo la lógica de negocio de cambios tecnológicos."),
        ("CQRS (Command Query Responsibility Segregation)", "Patrón de diseño que separa las operaciones de lectura (Query) de las de escritura (Command) para optimizar el rendimiento y la escalabilidad."),
        ("DDP (Defect Detection Percentage)", "KPI ISTQB: Porcentaje de defectos detectados durante el proceso de QA versus el total (QA + producción). Fórmula: DDP = DefectosQA / (DefectosQA + DefectosProducción) × 100."),
        ("Docker", "Plataforma de contenedorización que empaqueta aplicaciones y sus dependencias en imágenes portátiles ejecutables en cualquier entorno que soporte Docker Engine."),
        ("Docker Compose", "Herramienta para definir y ejecutar aplicaciones multi-contenedor Docker mediante un archivo YAML declarativo. QAMS usa docker-compose.yml para orquestar 4 servicios."),
        ("DRE (Defect Removal Efficiency)", "KPI ISTQB: Eficiencia en la eliminación de defectos previo a producción. Fórmula: DRE = DefectosEliminadosEnPruebas / TotalDefectos × 100. Umbral recomendado: ≥ 85%."),
        ("EF Core (Entity Framework Core)", "ORM (Object-Relational Mapper) de Microsoft para .NET que abstrae el acceso a base de datos mediante LINQ y entidades C#. QAMS usa EF Core 9 con PostgreSQL 16."),
        ("ERD (Entity-Relationship Diagram)", "Diagrama que representa las entidades de datos de un sistema, sus atributos y las relaciones entre ellas. QAMS tiene un ERD con 23 entidades en 5 dominios cohesivos."),
        ("Fast Runner", "Módulo de interfaz reactiva de QAMS optimizado para ejecución de alta velocidad de pruebas mediante atajos de teclado (P/F/B). Implementado con Angular Signals para latencia < 200ms."),
        ("Gherkin", "Lenguaje de dominio específico (DSL) para especificar comportamientos en BDD usando las palabras clave Given, When, Then, And, But. Soportado nativamente en el editor de TestCase de QAMS."),
        ("IAuditable", "Interfaz C# que define los 4 campos de auditoría automática: CreatedAt, CreatedByUserId, UpdatedAt, UpdatedByUserId. Implementada por las 20 entidades principales de QAMS."),
        ("ISoftDelete", "Interfaz C# que define los 3 campos de eliminación lógica: IsDeleted, DeletedAt, DeletedByUserId. Previene la pérdida permanente de datos históricos de calidad."),
        ("ISTQB (International Software Testing Qualifications Board)", "Organismo internacional de certificación de profesionales de pruebas de software. El CTFL v4.0 es el estándar de referencia de QAMS."),
        ("JWT (JSON Web Token)", "Estándar RFC 7519 para transmitir claims de seguridad como objeto JSON firmado digitalmente. QAMS emite JWTs firmados con HMAC-SHA256, expiración 8h y claims de permisos."),
        ("Kanban", "Sistema de gestión visual del flujo de trabajo basado en columnas y tarjetas. QAMS implementa un tablero Kanban para el seguimiento del ciclo de vida de defectos."),
        ("MTTR (Mean Time to Resolution)", "KPI ISTQB: Tiempo promedio en horas para cerrar un defecto desde su apertura. Fórmula: MTTR = Σ(ClosedAt - CreatedAt) / COUNT(defectos cerrados). Umbral QAMS: ≤ 48h."),
        ("Nginx", "Servidor web y proxy inverso de alto rendimiento. QAMS lo usa como servidor de archivos estáticos para el frontend Angular en producción (contenedor Docker)."),
        ("OWASP (Open Web Application Security Project)", "Organización sin fines de lucro que produce estándares, guías y herramientas de seguridad web. El OWASP Top 10 2021 es la referencia de seguridad implementada en QAMS."),
        ("PostgreSQL", "Sistema de gestión de bases de datos relacional objeto (ORDBMS) de código abierto. QAMS usa PostgreSQL 16 como base de datos principal, desplegado en contenedor Docker Alpine."),
        ("Quality Gate", "Conjunto de umbrales cuantitativos mínimos (Pass Rate, DRE, cobertura RTM, defectos críticos) que un proyecto debe superar para ser autorizado a pasar a producción."),
        ("RBAC (Role-Based Access Control)", "Modelo de control de acceso donde los permisos se asignan a roles y los roles a usuarios. QAMS implementa RBAC con 5 roles y 15+ permisos atómicos."),
        ("Redis", "Almacén de datos en memoria de alta velocidad, usado en QAMS para colas de tareas de fondo (envío de emails de recuperación de contraseña) y caché de sesiones."),
        ("Repository Pattern", "Patrón de diseño que abstrae la capa de acceso a datos detrás de interfaces, permitiendo intercambiar el mecanismo de persistencia sin afectar la lógica de negocio."),
        ("RTM (Requirements Traceability Matrix)", "Matriz tabular bidireccional que vincula Requisitos ↔ Casos de Prueba ↔ Ejecuciones ↔ Defectos. Calculada en tiempo real por el endpoint GET /api/Reports/rtm-matrix de QAMS."),
        ("SBTM (Session-Based Test Management)", "Metodología de gestión de pruebas exploratorias basada en sesiones acotadas en tiempo (time-box) con cartas de misión (charters). Implementada en el módulo ExploratorySession de QAMS."),
        ("SDLC (Software Development Life Cycle)", "Ciclo de vida del desarrollo de software: Planificación, Análisis, Diseño, Implementación, Pruebas, Despliegue y Mantenimiento."),
        ("Soft-Delete", "Patrón de persistencia donde los registros eliminados se marcan con IsDeleted=true sin ser borrados físicamente de la base de datos, preservando el historial completo."),
        ("SPA (Single Page Application)", "Arquitectura de aplicación web donde toda la lógica de navegación ocurre en el cliente sin recargar la página. QAMS frontend es una SPA Angular 19."),
        ("Sprint", "Iteración de tiempo fijo (1-4 semanas) en metodologías ágiles Scrum. Los 55 Story Points de QAMS equivalen a 5-6 Sprints de 2 semanas con velocidad de 10 SP/Sprint."),
        ("STLC (Software Testing Life Cycle)", "Ciclo de vida específico de las pruebas de software: Planificación, Análisis de Requisitos, Diseño, Implementación, Ejecución y Cierre. QAMS soporta todas las fases del STLC."),
        ("Story Points", "Unidad de estimación relativa del esfuerzo en metodologías ágiles. QAMS usa la escala Fibonacci (1, 2, 3, 5, 8, 13, 21). Las 9 HUs suman 55 Story Points totales."),
        ("TCO (Total Cost of Ownership)", "Costo Total de Propiedad: suma de todos los costos directos e indirectos (licencias, infraestructura, formación, soporte) a lo largo del ciclo de vida del software."),
        ("Unit of Work", "Patrón de diseño que agrupa múltiples operaciones de repositorio en una única transacción de base de datos, garantizando atomicidad ACID. Implementado por IUnitOfWork en QAMS."),
        ("xUnit", "Framework de pruebas unitarias para .NET de código abierto. Usado en QAMS.Tests para pruebas unitarias e integración del backend con cobertura del 76%."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r_t = p.add_run(f"• {term}: ")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        r_d = p.add_run(definition)
        r_d.font.size = Pt(10)

    doc.save(DOCX_OUTPUT_PATH_1)
    print(f"Documento extendido guardado exitosamente en: {DOCX_OUTPUT_PATH_1}")
    try:
        doc.save(DOCX_OUTPUT_PATH_2)
        print(f"Documento extendido guardado exitosamente en: {DOCX_OUTPUT_PATH_2}")
    except PermissionError:
        alt_path = r"c:\diplomado\proyecto_actualizado.docx"
        doc.save(alt_path)
        print(f"Nota: c:\\diplomado\\proyecto.docx está abierto en Word. Se guardó una copia en: {alt_path}")

if __name__ == "__main__":
    build_document()
